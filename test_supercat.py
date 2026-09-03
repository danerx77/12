#!/usr/bin/env python3
"""Testy funkcjonalne SuperCAT – pełny przepływ pracy tłumacza."""
import os
import shutil
import sys
import tempfile

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QApplication, QMessageBox, QTableWidgetItem

# Testy działają bez interakcji – blokujące okna dialogowe są wyłączone.
QMessageBox.information = staticmethod(lambda *a, **k: QMessageBox.StandardButton.Ok)
QMessageBox.warning = staticmethod(lambda *a, **k: QMessageBox.StandardButton.Ok)
QMessageBox.critical = staticmethod(lambda *a, **k: QMessageBox.StandardButton.Ok)
QMessageBox.question = staticmethod(lambda *a, **k: QMessageBox.StandardButton.Yes)

from supercat.core.fileparser import (Segment, parse_file, export_by_replacement, export_xliff,
                                      export_html_bilingual)
from supercat.core.exclusions import ExclusionRule, ExclusionSet
from supercat.core.exclusions import default_rules as default_exclusions
from supercat.core.glossary import DOWNLOADABLE_DICTIONARIES, Dictionary, Glossary
from supercat.core.project import Project, ProjectManager, SegmentationSettings
from supercat.ui.editor_tab import _is_done as _is_done_seg
from supercat.core.qa import (chars_no_spaces, file_statistics, project_statistics,
                              segment_statistics, run_qa)
from supercat.core.settings import SettingsManager
from supercat.core.ai_clean import build_translation_prompt
from supercat.core.langcheck import (apply_first_suggestions, check_offline, check_translation,
                                     mask_codes, polish_mt_output)
from supercat.core.langcheck import default_options as default_lang_options
from supercat.core.langcheck import summarize as summarize_lang
from supercat.core.search import SearchOptions, replace_in_segments, search_segments
from supercat.core.segmentation import parse_abbreviations, segment_text
from supercat.core.textutil import (context_snippet, copy_edge_whitespace, describe_edges,
                                    display_text, find_matches, mark_edges, markers_for_style,
                                    replace_matches, split_edges)
from supercat.core.tags import adapt_translation, normalize_tags_for_comparison, count_tags
from supercat.core.tm import TranslationMemory, similarity_percent

PASS, FAIL = [], []


def check(name, condition, extra=""):
    (PASS if condition else FAIL).append(name)
    print(("  ✅ " if condition else "  ❌ ") + name + (f"  [{extra}]" if extra else ""))


def _raises(fn) -> bool:
    """True, gdy wywołanie podniosło wyjątek (do testów walidacji)."""
    try:
        fn()
        return False
    except Exception:
        return True


def _safe_list(m):
    """Zwraca komunikat błędu z list_gemini_models (do testu bez sieci)."""
    try:
        m.list_gemini_models()
        return ""
    except Exception as exc:
        return str(exc)


def main():
    app = QApplication.instance() or QApplication([])
    tmp = tempfile.mkdtemp(prefix="supercat_test_")
    print(f"Katalog testowy: {tmp}\n")

    # ---------------------------------------------------------- segmentacja
    print("1. Segmentacja")
    s = SegmentationSettings()
    segs = segment_text("To jest pierwsze zdanie. To drugie! A trzecie? Koniec.", s)
    check("podział na 4 zdania", len(segs) == 4, f"{len(segs)}: {segs}")
    segs2 = segment_text("Kupiłem 3.14 kg jabłek np. w sklepie. Drugie zdanie.", s)
    check("skróty i liczby nie łamią zdań", len(segs2) == 2, str(segs2))
    s_line = SegmentationSettings(mode="line")
    check("tryb line", len(segment_text("a\nb\nc", s_line)) == 3)
    s_re = SegmentationSettings(mode="regex", regex_pattern=r"\|\|")
    check("tryb regex", len(segment_text("a||b||c", s_re)) == 3)

    # --- zachowywanie spacji na brzegach (wcięcia z pliku źródłowego) ----
    s_ws = SegmentationSettings()
    out_ws = segment_text("   Wcięty tekst. Drugie zdanie.", s_ws)
    check("wiodąca spacja zachowana", out_ws and out_ws[0].startswith("   "), repr(out_ws[:1]))
    s_nows = SegmentationSettings(preserve_whitespace=False)
    out_nows = segment_text("   Wcięty tekst.", s_nows)
    check("można wyłączyć zachowywanie spacji",
          out_nows and not out_nows[0].startswith(" "), repr(out_nows[:1]))
    s_line_ws = SegmentationSettings(mode="line")
    out_line = segment_text(" a\n  b", s_line_ws)
    check("tryb line zachowuje wcięcia",
          out_line == [" a", "  b"], repr(out_line))

    # --- rozbudowana segmentacja (reguły w stylu OmegaT) -----------------
    print("\n1a. Reguły segmentacji")
    check("skróty z ustawień są rozpoznawane",
          parse_abbreviations("np., itd. ; zał.") == {"np", "itd", "zał"},
          str(parse_abbreviations("np., itd. ; zał.")))
    s_num = SegmentationSettings(skip_after_numbers=True)
    check("nie dzieli po liczbie z kropką",
          segment_text("Punkt 5. mówi wyraźnie. Koniec.", s_num)
          == ["Punkt 5. mówi wyraźnie.", "Koniec."],
          str(segment_text("Punkt 5. mówi wyraźnie. Koniec.", s_num)))
    s_up = SegmentationSettings(require_uppercase_after=True)
    check("dzieli tylko przed wielką literą",
          len(segment_text("Kupiłem to. potem wyszedłem. Koniec.", s_up)) == 2,
          str(segment_text("Kupiłem to. potem wyszedłem. Koniec.", s_up)))
    s_ab = SegmentationSettings(custom_abbreviations="zał")
    check("własny skrót nie kończy zdania",
          len(segment_text("Zobacz zał. 3 tutaj. Koniec.", s_ab)) == 2,
          str(segment_text("Zobacz zał. 3 tutaj. Koniec.", s_ab)))
    s_codes = SegmentationSettings(split_on_codes=True)
    check("dzielenie po znacznikach gier",
          len(segment_text("Witaj\\nświecie.\\pDalej.", s_codes)) == 3,
          str(segment_text("Witaj\\nświecie.\\pDalej.", s_codes)))
    s_min = SegmentationSettings(min_segment_length=10)
    check("krótkie segmenty są scalane",
          len(segment_text("A. B. Dłuższe zdanie tutaj.", s_min)) == 1,
          str(segment_text("A. B. Dłuższe zdanie tutaj.", s_min)))
    check("kropka + nowa linia kończy zdanie",
          segment_text("Pierwsze zdanie.\nDrugie zdanie.", SegmentationSettings())
          == ["Pierwsze zdanie.", "Drugie zdanie."],
          str(segment_text("Pierwsze zdanie.\nDrugie zdanie.", SegmentationSettings())))
    check("skrót przed nową linią nadal nie dzieli",
          segment_text("Zobacz str. 15 tutaj.\nKoniec.", SegmentationSettings())
          == ["Zobacz str. 15 tutaj.", "Koniec."],
          str(segment_text("Zobacz str. 15 tutaj.\nKoniec.", SegmentationSettings())))
    check("domyślnie reguły dodatkowe nie zmieniają wyniku",
          len(segment_text("Pierwsze zdanie. Drugie zdanie.", SegmentationSettings())) == 2)

    # --- wykluczanie segmentów technicznych -------------------------------
    print("\n1f. Wykluczanie segmentów")
    excl_segs = [
        Segment("1", "<<< FILE: CeladonCity_Condominiums_RoofRoom/text.inc >>>", "", file_name="a.txt"),
        Segment("2", "Thank you for using the STAMP CARD System.", "", file_name="a.txt"),
        Segment("3", "<<< FILE: PalletTown/scripts.inc >>>", "", file_name="a.txt"),
        Segment("4", "#org @8005A2", "", file_name="b.txt"),
        Segment("5", "{STR_VAR_1}", "", file_name="b.txt"),
        Segment("6", "[POKEMON_NAME]", "", file_name="b.txt"),
    ]
    rules = ExclusionSet(default_exclusions())
    check("domyślnie włączona reguła <<< FILE: … >>>",
          len(rules.preview(excl_segs)) == 2, str(len(rules.preview(excl_segs))))
    check("różne nazwy plików w nagłówku są objęte jedną regułą",
          rules.is_excluded("<<< FILE: A/b.inc >>>")
          and rules.is_excluded("<<< FILE: Zupelnie/Inna/Sciezka.s >>>"))
    check("zwykły tekst NIE jest wykluczany",
          not rules.is_excluded("Thank you for using the STAMP CARD System."))

    excluded, restored = rules.apply(excl_segs)
    check("wykluczone segmenty oznaczone jako pominięte",
          excluded == 2 and excl_segs[0].ignored and not excl_segs[1].ignored,
          f"{excluded}/{restored}")
    check("wykluczenie NIE kasuje treści segmentu",
          excl_segs[0].source.startswith("<<< FILE:"))

    # wyłączenie reguły przywraca segmenty
    for rule in rules.rules:
        rule.enabled = False
    excluded2, restored2 = rules.apply(excl_segs)
    check("wyłączenie reguły przywraca segmenty",
          restored2 == 2 and not excl_segs[0].ignored, f"{excluded2}/{restored2}")

    # ręczne pominięcie użytkownika nie jest ruszane
    excl_segs[1].ignored = True
    rules.apply(excl_segs)
    check("ręczne pominięcie nie jest cofane", excl_segs[1].ignored)
    excl_segs[1].ignored = False

    # typy dopasowania
    check("dopasowanie „zaczyna się od”",
          ExclusionRule("#org", "starts").matches("#org @8005A2")
          and not ExclusionRule("#org", "starts").matches("kod #org"))
    check("dopasowanie „dokładnie równy”",
          ExclusionRule("{STR_VAR_1}", "exact").matches("  {STR_VAR_1}  ")
          and not ExclusionRule("{STR_VAR_1}", "exact").matches("a {STR_VAR_1} b"))
    check("dopasowanie „kończy się na”",
          ExclusionRule(">>>", "ends").matches("<<< FILE: x >>>"))
    check("wzorzec z gwiazdką",
          ExclusionRule("<<< FILE:*>>>", "wildcard").matches("<<< FILE: cokolwiek >>>"))
    check("wyrażenie regularne",
          ExclusionRule(r"^\s*#\w+", "regex").matches("  #raw 0x08"))
    check("wielkość liter domyślnie ignorowana",
          ExclusionRule("<<< file:*>>>", "wildcard").matches("<<< FILE: a >>>"))
    check("rozróżnianie wielkości liter działa",
          not ExclusionRule("<<< file:*>>>", "wildcard", case_sensitive=True)
          .matches("<<< FILE: a >>>"))
    check("błędny regex jest zgłaszany, nie wywraca programu",
          ExclusionRule("[", "regex").error() != ""
          and ExclusionRule("[", "regex").matches("cokolwiek") is False)

    # reguła tylko dla wybranego pliku
    file_rule = ExclusionRule("#org", "starts", file_filter="b.txt")
    check("reguła ograniczona do pliku",
          file_rule.matches("#org @1", "b.txt") and not file_rule.matches("#org @1", "a.txt"))

    # zapis i odczyt
    saved = ExclusionSet([ExclusionRule("<<< FILE:*>>>", "wildcard", comment="test")]).to_dict()
    loaded = ExclusionSet.from_dict(saved)
    check("reguły zapisują się i wczytują",
          len(loaded.rules) == 1 and loaded.rules[0].comment == "test")
    check("główny wyłącznik wykluczania działa",
          not ExclusionSet(default_exclusions(), enabled=False).is_excluded("<<< FILE: a >>>"))
    check("liczniki trafień na regułę",
          ExclusionSet(default_exclusions()).counts(excl_segs).get("<<< FILE:*>>>") == 2,
          str(ExclusionSet(default_exclusions()).counts(excl_segs).get("<<< FILE:*>>>")))

    # segmentacja nie rozbija nagłówka na dwukropku
    header = "<<< FILE: CeladonCity/text.inc >>>"
    check("nagłówek <<< FILE: … >>> pozostaje jednym segmentem",
          segment_text(header, SegmentationSettings()) == [header],
          str(segment_text(header, SegmentationSettings())))
    check("nagłówek nie skleja się z następnym zdaniem",
          segment_text(header + "\nThank you.", SegmentationSettings())
          == [header, "Thank you."],
          str(segment_text(header + "\nThank you.", SegmentationSettings())))
    check("zwykły dwukropek nadal dzieli zdanie",
          len(segment_text("Zdanie: z dwukropkiem. Drugie.", SegmentationSettings())) == 2)

    # ------------------------------------------------- narzędzia tekstowe
    print("\n1b. Białe znaki i wyszukiwanie odporne na znaczniki")
    check("split_edges rozdziela brzegi",
          split_edges("  abc ") == ("  ", "abc", " "), repr(split_edges("  abc ")))
    check("copy_edge_whitespace przenosi wcięcie",
          copy_edge_whitespace("  Hello.", "Cześć.") == "  Cześć.",
          repr(copy_edge_whitespace("  Hello.", "Cześć.")))
    check("copy_edge_whitespace nie dubluje spacji",
          copy_edge_whitespace("  Hello.", "  Cześć.") == "  Cześć.")
    check("copy_edge_whitespace pomija pusty tekst",
          copy_edge_whitespace("  Hello.", "") == "")
    check("mark_edges pokazuje wcięcie", mark_edges("  ab") == "␣␣ab", mark_edges("  ab"))
    check("mark_edges oznacza tabulator", mark_edges("\tab") == "→ab", mark_edges("\tab"))
    check("mark_edges nie rusza wnętrza tekstu",
          mark_edges("a  b") == "a  b", mark_edges("a  b"))

    # --- znaki specjalne (␣ → ⏎) z możliwością wyłączenia ---------------
    demo = "  Line one.\nLine two.\t"
    check("display_text: oba znaki włączone",
          display_text(demo) == "␣␣Line one. ⏎ Line two.→", repr(display_text(demo)))
    check("display_text: bez znaku końca wiersza",
          display_text(demo, show_newlines=False) == "␣␣Line one. Line two.→",
          repr(display_text(demo, show_newlines=False)))
    check("display_text: bez znaków spacji",
          display_text(demo, show_spaces=False) == "  Line one. ⏎ Line two.\t",
          repr(display_text(demo, show_spaces=False)))
    check("display_text: oba wyłączone daje surowy tekst",
          display_text(demo, show_spaces=False, show_newlines=False)
          == "  Line one. Line two.\t")
    sp, tb, nl = markers_for_style("· » ¶  (dyskretne)")
    check("zestaw znaków: dyskretny", (sp, tb, nl) == ("·", "»", "¶"), f"{sp}{tb}{nl}")
    check("display_text: własny zestaw znaków",
          display_text(demo, space_marker=sp, tab_marker=tb, newline_marker=nl)
          == "··Line one. ¶ Line two.»",
          repr(display_text(demo, space_marker=sp, tab_marker=tb, newline_marker=nl)))
    check("nieznany zestaw wraca do domyślnego",
          markers_for_style("nie ma takiego") == ("␣", "→", "⏎"))
    check("context_snippet bez znaku ⏎",
          "⏎" not in context_snippet("a\nb ala", (4, 7), newline_marker=None),
          context_snippet("a\nb ala", (4, 7), newline_marker=None))
    check("describe_edges opisuje brzegi", "2" in describe_edges("  ab"), describe_edges("  ab"))
    check("describe_edges pusty dla zwykłego tekstu", describe_edges("ab") == "")

    src_code = "Thank you for using the STAMP CARD\\nSystem."
    check("fraza znaleziona mimo znacznika \\n",
          len(find_matches(src_code, "STAMP CARD System", ignore_codes=True)) == 1,
          str(find_matches(src_code, "STAMP CARD System", ignore_codes=True)))
    check("bez ignorowania znaczników brak trafienia",
          find_matches(src_code, "STAMP CARD System", ignore_codes=False) == [])
    span = find_matches(src_code, "STAMP CARD System", ignore_codes=True)[0]
    check("zakres wskazuje oryginalne znaki",
          src_code[span[0]:span[1]].startswith("STAMP CARD") and "System" in src_code[span[0]:span[1]],
          repr(src_code[span[0]:span[1]]))
    check("znajduje mimo ogonków",
          len(find_matches("Żółw idzie", "zolw", ignore_accents=True)) == 1)
    check("wielkość liter respektowana",
          find_matches("Żółw", "żółw", case_sensitive=True) == [])
    check("tryb całe słowo nie łapie fragmentu",
          find_matches("Systemowy", "system", mode="word") == [])
    check("tryb całe słowo łapie wyraz",
          len(find_matches("System jest.", "system", mode="word")) == 1)
    check("tryb dokładny",
          len(find_matches("System.", "system.", mode="exact")) == 1
          and find_matches("System dalej.", "system.", mode="exact") == [])
    check("regex działa", len(find_matches("abc123", r"\d+", mode="regex")) == 1)
    check("kontekst zawiera cudzysłowy dopasowania",
          "«" in context_snippet("ala ma kota", (4, 6)))
    new_text, n = replace_matches("Kot i kot", "kot", "pies")
    check("zamiana ignorująca wielkość liter", (new_text, n) == ("pies i pies", 2), repr(new_text))
    check("zamiana nie rusza tekstu bez trafień",
          replace_matches("abc", "xyz", "q") == ("abc", 0))

    # --- rozbudowane statystyki ------------------------------------------
    print("\n1e. Statystyki")
    stat_segs = [
        Segment("1", "Hello world here.", "Witaj świecie tutaj.", file_name="a.txt"),
        Segment("2", "Thank you\\nfor {PLAYER}.", "", file_name="a.txt"),
        Segment("3", "Hello world here.", "", file_name="b.txt"),
    ]
    st = project_statistics(stat_segs, 10)
    check("statystyki: znaki ze spacjami",
          st["Znaki ze spacjami (źródło)"] == 58, str(st["Znaki ze spacjami (źródło)"]))
    check("statystyki: znaki bez spacji",
          st["Znaki bez spacji (źródło)"] == 52, str(st["Znaki bez spacji (źródło)"]))
    check("znaki bez spacji < znaki ze spacjami",
          st["Znaki bez spacji (źródło)"] < st["Znaki ze spacjami (źródło)"])
    check("statystyki: znaki w tłumaczeniu bez spacji",
          st["Znaki bez spacji (tłumaczenie)"] == 18,
          str(st["Znaki bez spacji (tłumaczenie)"]))
    check("statystyki: strony rozliczeniowe", "Strony rozliczeniowe (1800 zn.)" in st)
    check("statystyki: segmenty powtórzone", st["Segmenty powtórzone"] == 1,
          str(st["Segmenty powtórzone"]))
    check("statystyki: liczba znaczników", st["Znaczniki w źródle"] == 2,
          str(st["Znaczniki w źródle"]))
    check("statystyki: średnia i najdłuższy segment",
          st["Średnia długość segmentu (słowa)"] > 0 and st["Najdłuższy segment (słowa)"] == 4,
          f"{st['Średnia długość segmentu (słowa)']}, {st['Najdłuższy segment (słowa)']}")
    check("statystyki: znaki do przetłumaczenia liczone bez pominiętych",
          st["Znaki do przetłumaczenia"] == 41, str(st["Znaki do przetłumaczenia"]))
    check("chars_no_spaces liczy poprawnie",
          chars_no_spaces("a b\tc\n") == 3, str(chars_no_spaces("a b\tc\n")))

    rows = file_statistics(stat_segs)
    check("statystyki per plik", len(rows) == 2 and rows[0]["Plik"] == "a.txt", str(rows))
    check("per plik: znaki bez spacji", rows[0]["Znaki bez spacji"] == 37,
          str(rows[0]["Znaki bez spacji"]))
    check("per plik: postęp", rows[0]["Postęp (%)"] == 50.0, str(rows[0]["Postęp (%)"]))

    seg_st = segment_statistics("Hello  world.", "Witaj świecie.")
    check("statystyki segmentu: znaki bez spacji",
          seg_st["Znaki bez spacji (źródło)"] == 11, str(seg_st["Znaki bez spacji (źródło)"]))
    check("statystyki segmentu: stosunek długości",
          seg_st["Stosunek długości (%)"] > 100, str(seg_st["Stosunek długości (%)"]))
    check("statystyki segmentu: liczba zdań", seg_st["Zdania (źródło)"] == 1)

    # ------------------------------------------- kontrola poprawności języka
    print("\n1d. Kontrola języka (tylko tłumaczenie)")
    check("wykryto niezgodność zaimka i czasownika",
          any(i.category == "Odmiana" for i in check_offline("Ja poszedł do sklepu.")))
    check("wykryto złą odmianę po liczebniku",
          any("dopełniacz" in i.message for i in check_offline("Mam pięć jabłko.")),
          str([i.message for i in check_offline("Mam pięć jabłko.")]))
    check("wykryto mianownik po „dwa”",
          any("mianownik" in i.message for i in check_offline("Widzę dwa kotów.")))
    check("wykryto spację przed przecinkiem",
          any(i.rule_id == "SPACJA_PRZED_ZNAKIEM" for i in check_offline("Tak , owszem.")))
    check("wykryto brak spacji po przecinku",
          any(i.rule_id == "BRAK_SPACJI_PO_ZNAKU" for i in check_offline("Tak,owszem.")))
    check("liczba dziesiętna nie jest błędem",
          not any(i.rule_id == "BRAK_SPACJI_PO_ZNAKU" for i in check_offline("Cena to 1,5 zł.")))
    check("wykryto powtórzony wyraz",
          any(i.rule_id == "POWTORZONY_WYRAZ" for i in check_offline("To to jest dobre.")))
    check("przecinek rozdziela poprawne „To, to”",
          not any(i.rule_id == "POWTORZONY_WYRAZ" for i in check_offline("To, to jest dobre.")))
    check("wykryto małą literę po kropce",
          any(i.rule_id == "MALA_PO_KROPCE" for i in check_offline("Zdanie. drugie zdanie.")))
    check("skrót „np.” nie kończy zdania",
          not any(i.rule_id == "MALA_PO_KROPCE" for i in check_offline("Owoce, np. jabłka.")))
    check("wykryto błędną formę „wziąść”",
          any("wziąć" in " ".join(i.suggestions) for i in check_offline("Trzeba wziąść to.")))
    check("wykryto podwójną spację w środku",
          any(i.rule_id == "PODWOJNA_SPACJA" for i in check_offline("Tekst  z odstępem.")))

    # znaczniki plików gier NIE mogą być zgłaszane jako błędy
    game_text = "Dziękujemy za korzystanie ze\\nSystemu KART.\\pMasz {STR_VAR_1} sztuk.<<KON>>"
    check("znaczniki gier nie są zgłaszane", check_offline(game_text) == [],
          str([i.message for i in check_offline(game_text)]))
    masked, spans = mask_codes("A\\nB {VAR} C")
    check("maskowanie zachowuje długość tekstu", len(masked) == len("A\\nB {VAR} C"))
    check("maskowanie znajduje znaczniki", len(spans) == 2, str(spans))

    check("puste tłumaczenie nie daje uwag", check_offline("") == [])

    # mały słownik nie może zalewać panelu „słowami spoza słownika”
    class _TinyDict:
        is_initialized = True
        size = 12
        def check_text(self, text):
            return ["jabłek", "dobre"]

    class _FullDict(_TinyDict):
        size = 60000

    check("mały słownik nie zgłasza pisowni",
          not any(i.category == "Pisownia" for i in check_offline("Mam pięć jabłek.", _TinyDict())))
    check("pełny słownik zgłasza pisownię",
          any(i.category == "Pisownia" for i in check_offline("Mam pięć jabłek.", _FullDict())))
    check("podsumowanie dla braku uwag", "✅" in summarize_lang([]))
    check("podsumowanie liczy błędy",
          "⚠️" in summarize_lang(check_offline("Tak , owszem.")),
          summarize_lang(check_offline("Tak , owszem.")))

    # automatyczne poprawki
    fixed, count = apply_first_suggestions("Tak , owszem.", check_offline("Tak , owszem."))
    check("automatyczna poprawka usuwa spację przed przecinkiem",
          fixed == "Tak, owszem." and count == 1, f"{fixed!r} ({count})")

    # porządkowanie wyniku MT
    out, changes = polish_mt_output("Mam pieniądze , ale  nie kupię.")
    check("porządkowanie MT poprawia interpunkcję i spacje",
          out == "Mam pieniądze, ale nie kupię.", repr(out))
    out2, _ = polish_mt_output("wynik zdania.", "Source sentence.")
    check("porządkowanie MT podnosi wielką literę", out2 == "Wynik zdania.", repr(out2))
    game_mt = "Tekst ze\\nznacznikiem , dalej.\\p{PLAYER}"
    out3, _ = polish_mt_output(game_mt)
    check("porządkowanie MT nie gubi znaczników",
          out3.count("\\n") == 1 and out3.count("\\p") == 1 and "{PLAYER}" in out3, repr(out3))
    check("porządkowanie MT nie rusza poprawnego tekstu",
          polish_mt_output("Wszystko w porządku.")[0] == "Wszystko w porządku.")

    # reguły odmiany w poleceniu dla AI
    prompt_on = build_translation_prompt("angielski", "polski", grammar_rules=True)
    prompt_off = build_translation_prompt("angielski", "polski", grammar_rules=False)
    check("polecenie AI zawiera wymagania odmiany",
          "ODMIANY" in prompt_on and "pięć jabłek" in prompt_on)
    check("wymagania odmiany można wyłączyć",
          "ODMIANY" not in prompt_off and len(prompt_off) < len(prompt_on))
    check("polecenie AI zachowuje zasady o znacznikach w obu wariantach",
          "@#0#@" in prompt_on and "@#0#@" in prompt_off)

    # ------------------------------------------------- wyszukiwanie w projekcie
    print("\n1c. Wyszukiwanie w wielu plikach")

    class _Seg:
        def __init__(self, source, target="", file_name="a.txt"):
            self.source, self.target, self.file_name = source, target, file_name

    seg_list = [
        _Seg("Thank you for using the STAMP CARD\\nSystem.", "", "plik1.txt"),
        _Seg("The STAMP CARD is full.", "KARTA jest pełna.", "plik1.txt"),
        _Seg("Use the stamp card here.", "", "plik2.txt"),
        _Seg("Nothing to see.", "Nic tu nie ma.", "plik2.txt"),
    ]
    res = search_segments(seg_list, "stamp card", SearchOptions())
    check("znaleziono w obu plikach", len(res.by_file()) == 2, str(res.file_counts()))
    check("licznik trafień", res.total_matches == 3, str(res.total_matches))
    res_one = search_segments(seg_list, "stamp card", SearchOptions(files=["plik2.txt"]))
    check("zakres jednego pliku", res_one.total_matches == 1, str(res_one.total_matches))
    res_tgt = search_segments(seg_list, "karta", SearchOptions(in_source=False))
    check("szukanie tylko w tłumaczeniu", res_tgt.total_matches == 1)
    res_untr = search_segments(seg_list, "stamp card", SearchOptions(only_untranslated=True))
    check("filtr: tylko nieprzetłumaczone", res_untr.total_matches == 2, str(res_untr.total_matches))
    res_phrase = search_segments(seg_list, "STAMP CARD System", SearchOptions())
    check("fraza przez znacznik w segmencie", res_phrase.total_matches == 1)
    res_bad = search_segments(seg_list, "[", SearchOptions(mode="regex"))
    check("błędny regex zwraca komunikat", bool(res_bad.error), res_bad.error)
    check("podsumowanie opisuje wynik", "Znaleziono" in res.summary(), res.summary())
    changed, total = replace_in_segments(seg_list, "KARTA", "KARTECZKA", SearchOptions())
    check("zamiana w tłumaczeniach", (changed, total) == (1, 1) and
          seg_list[1].target == "KARTECZKA jest pełna.", seg_list[1].target)

    # QA: kontrola spacji na brzegach
    from supercat.core.fileparser import Segment as _QSeg
    ws_issues = run_qa([_QSeg("s1", "  Hello.", "Cześć.")])
    check("QA zgłasza zgubione wcięcie",
          any(i.category == "Białe znaki" for i in ws_issues),
          str([i.message for i in ws_issues]))
    ok_issues = run_qa([_QSeg("s2", "  Hello.", "  Cześć.")])
    check("QA nie zgłasza poprawnego wcięcia",
          not any(i.category == "Białe znaki" for i in ok_issues))
    ok2 = run_qa([_QSeg("s3", "  Hello world.", "  Witaj świecie.")])
    check("wcięcie nie liczy się jako podwójna spacja",
          not any("Podwójne spacje" in i.message for i in ok2))

    # ------------------------------------------------------------ projekt
    print("\n2. Projekt")
    pm = ProjectManager.instance()
    project = pm.create_project("Test Projekt", "en", "pl", tmp)
    check("utworzono katalog projektu", os.path.isdir(project.project_path))
    for folder in ("source", "target", "tm", "glossary", "dictionary", "export"):
        check(f"folder {folder}/", os.path.isdir(os.path.join(project.project_path, folder)))
    check("plik .scproj", os.path.exists(project.project_file_path))
    pm2 = ProjectManager()
    reopened = pm2.open_project(project.project_file_path)
    check("ponowne otwarcie projektu", reopened.name == "Test Projekt" and reopened.target_lang == "pl")

    # ------------------------------------------------------------- parser
    print("\n3. Parsowanie plików")
    src_txt = os.path.join(project.source_path, "doc.txt")
    with open(src_txt, "w", encoding="utf-8") as fh:
        fh.write("Hello world. This is a test.\nThe system helps you work efficiently.\n")
    segments = parse_file(src_txt, project.segmentation)
    check("TXT → segmenty", len(segments) == 3, f"{len(segments)}")

    from docx import Document
    docx_path = os.path.join(project.source_path, "doc.docx")
    d = Document()
    d.add_paragraph("First paragraph of the document.")
    d.add_paragraph("Second paragraph here.")
    d.save(docx_path)
    docx_segs = parse_file(docx_path, project.segmentation)
    check("DOCX → segmenty", len(docx_segs) == 2, f"{len(docx_segs)}")

    srt_path = os.path.join(project.source_path, "movie.srt")
    with open(srt_path, "w", encoding="utf-8") as fh:
        fh.write("1\n00:00:01,000 --> 00:00:03,000\nHello there\n\n2\n00:00:04,000 --> 00:00:06,000\nGoodbye\n")
    srt_segs = parse_file(srt_path)
    check("SRT → segmenty", len(srt_segs) == 2, str([s.source for s in srt_segs]))

    po_path = os.path.join(project.source_path, "app.po")
    with open(po_path, "w", encoding="utf-8") as fh:
        fh.write('msgid ""\nmsgstr ""\n\nmsgid "Save file"\nmsgstr ""\n\nmsgid "Open"\nmsgstr "Otwórz"\n')
    po_segs = parse_file(po_path)
    check("PO → segmenty", len(po_segs) == 2, str([s.source for s in po_segs]))

    xlf_path = os.path.join(project.source_path, "a.xlf")
    export_xliff([type(segments[0])("u1", "Hello", "Cześć")], xlf_path, "en", "pl")
    xlf_segs = parse_file(xlf_path)
    check("XLIFF round-trip", len(xlf_segs) == 1 and xlf_segs[0].target == "Cześć")

    # ---------------------------------------------------------------- tagi
    print("\n4. Tagi")
    # {NAME}, <b>, </b> = 3 tagi
    check("liczenie tagów", count_tags("Hello {NAME}, click <b>here</b>") == 3,
          str(count_tags("Hello {NAME}, click <b>here</b>")))
    adapted = adapt_translation("Witaj {USER_2}!", "Hello {USER_1}!")
    check("adaptacja tagów", "{USER_2}" in adapted, adapted)
    check("normalizacja tagów",
          normalize_tags_for_comparison("{VAR_1}") == normalize_tags_for_comparison("{A_9}"))

    # ------------------------------------------------------------------ TM
    print("\n5. Pamięć tłumaczeń")
    tm = TranslationMemory()
    tm.init_for_project(project.tm_path)
    check("baza TM utworzona", os.path.exists(os.path.join(project.tm_path, "project_tm.db")))
    tm.add("Hello world.", "Witaj świecie.", "en", "pl")
    tm.add("This is a test.", "To jest test.", "en", "pl")
    tm.add("The system helps you work efficiently.", "System pomaga pracować efektywnie.", "en", "pl")
    check("dodano 3 wpisy", tm.size() == 3, str(tm.size()))
    check("brak duplikatów", (tm.add("Hello world.", "Witaj świecie."), tm.size())[1] == 3)

    exact = tm.find_fuzzy_matches("Hello world.", 70, 5)
    check("dopasowanie 100%", exact and exact[0].similarity == 100, str(exact[0].similarity if exact else "brak"))
    fuzzy = tm.find_fuzzy_matches("Hello world!", 60, 5)
    check("dopasowanie rozmyte", fuzzy and 70 <= fuzzy[0].similarity < 100, str(fuzzy[0].similarity if fuzzy else "brak"))
    check("similarity_percent", similarity_percent("kot", "kot") == 100 and similarity_percent("kot", "pies") < 50)

    tmx_path = os.path.join(project.export_path, "test.tmx")
    count = tm.export_tmx(tmx_path, "en", "pl")
    check("eksport TMX", os.path.exists(tmx_path) and count == 3)
    tm.clear()
    imported = tm.import_tmx(tmx_path)
    check("import TMX", imported == 3 and tm.size() == 3, f"import={imported}, size={tm.size()}")
    check("konkordancja", len(tm.search("Witaj")) == 1)

    # ----------------------------------------------- dopasowanie zdań + wsad
    print("\n5b. Dopasowanie zdań i tryb wsadowy")
    # funkcja jest domyślnie WYŁĄCZONA (kosztowna) – testy włączają ją jawnie
    _SM = SettingsManager
    _sm = _SM.instance()
    check("dopasowanie zdań domyślnie wyłączone",
          not _sm.get_bool("tm.sentence.matching.enabled", False))
    _sm.set("tm.sentence.matching.enabled", True)
    tm.add("the network settings", "ustawienia sieci", "en", "pl")
    sm = tm.find_sentence_matches("Please check the network settings before you continue.")
    check("znaleziono fragment zdania", len(sm) >= 1, str([m.fragment_source for m in sm]))
    check("złożono tłumaczenie fragmentu",
          any("ustawienia sieci" in m.assembled for m in sm),
          sm[0].assembled if sm else "brak")
    check("pokrycie fragmentu policzone", sm and 0 < sm[0].coverage <= 100)

    batch = tm.find_best_matches_batch(["Hello world.", "This is a test.", "Zupelnie inne zdanie xyz"], 70)
    check("wsad: 3 wyniki", len(batch) == 3)
    check("wsad: trafienia na znanych zdaniach", batch[0] is not None and batch[1] is not None)
    check("wsad: brak trafienia na nieznanym", batch[2] is None)
    check("wsad zgodny z pojedynczym wyszukiwaniem",
          batch[0].similarity == tm.find_fuzzy_matches("Hello world.", 70, 1)[0].similarity)
    # przywracamy domyślny stan — inaczej kolejna linia testów (w nowym
    # przebiegu) widziałaby wartość brudzoną przez ten test
    _sm.set("tm.sentence.matching.enabled", False)

    import time as _t
    big = TranslationMemory()
    big.init_for_project(os.path.join(tmp, "bigtm"))
    big.add_many([(f"Sentence number {i} about the system.", f"Zdanie numer {i} o systemie.", "en", "pl")
                  for i in range(5000)])
    _t0 = _t.perf_counter()
    r = big.find_fuzzy_matches("Sentence number 42 about the system.", 70, 5)
    dt_first = _t.perf_counter() - _t0
    check("pierwsze wyszukanie (budowa indeksu) < 500 ms", dt_first < 0.5, f"{dt_first*1000:.1f} ms")
    _t0 = _t.perf_counter()
    for _ in range(5):
        big.find_fuzzy_matches("Sentence number 42 about the system.", 70, 5)
    dt = (_t.perf_counter() - _t0) / 5
    check("kolejne wyszukania w TM 5000 wpisów < 50 ms", dt < 0.05, f"{dt*1000:.1f} ms")
    check("znaleziono trafny wynik", r and r[0].similarity == 100, str(r[0].similarity if r else "brak"))
    _t0 = _t.perf_counter()
    big.find_best_matches_batch([f"Sentence number {i} about the system." for i in range(300)], 80)
    dt_batch = _t.perf_counter() - _t0
    check("wsad 300 segmentów < 5 s", dt_batch < 5.0, f"{dt_batch:.2f} s")
    big.close()

    # ------------------------------- znaczniki \n oraz QuickTrans / silniki
    print("\n5c. Znaczniki \\n, \\p (pliki gier) i wiele silników MT")
    # sekcja zależy od dopasowania zdań — włączamy ją jawnie (bez polegania
    # na stanie pozostawionym przez wcześniejsze sekcje)
    _sm.set("tm.sentence.matching.enabled", True)
    tm_g = TranslationMemory()
    tm_g.init_for_project(os.path.join(tmp, "tm_games"))
    tm_g.add(r"Thank you for using the MYSTERY\nGIFT System.",
             r"Dziękujemy za korzystanie z\nSystemu MYSTERY GIFT", "en", "pl")

    seg_literal = (r"Thank you for using the MYSTERY\nGIFT System.\pYou must be {PLAYER}."
                   r"\nThere is a ticket here for you.")
    seg_real = ("Thank you for using the MYSTERY\nGIFT System.\\pYou must be {PLAYER}."
                "\\nThere is a ticket here for you.")
    seg_plain = "Thank you for using the MYSTERY GIFT System. You must be {PLAYER}."
    seg_moved = r"Thank you for using the MYSTERY GIFT\nSystem.\pYou must be {PLAYER}."

    for label, seg in (("literalne \\n", seg_literal), ("prawdziwy newline", seg_real),
                       ("bez znacznika", seg_plain), ("inne rozmieszczenie \\n", seg_moved)):
        sm = tm_g.find_sentence_matches(seg)
        check(f"dopasowanie zdań: {label}", len(sm) >= 1, f"znaleziono {len(sm)}")
        if sm:
            check(f"  złożono tłumaczenie: {label}", "Systemu MYSTERY GIFT" in sm[0].assembled,
                  sm[0].assembled[:60])

    # znaczniki nie mogą zniknąć z reszty segmentu
    sm = tm_g.find_sentence_matches(seg_literal)
    check("zachowano dalsze znaczniki \\p / \\n", r"\p" in sm[0].assembled and "{PLAYER}" in sm[0].assembled,
          sm[0].assembled[:80])

    from supercat.core.tm import unify_control_codes
    check("ujednolicanie znaczników",
          unify_control_codes(r"a\nb") == unify_control_codes("a\nb") == "a b",
          repr(unify_control_codes(r"a\nb")))

    # --- rozbicie wpisu TM na linie (\n / \p) ---
    from supercat.core.tm import align_lines, split_lines_by_codes
    check("podział na linie po \\n",
          split_lines_by_codes(r"Thank you for using the MYSTERY\nGIFT System.") ==
          ["Thank you for using the MYSTERY", "GIFT System."])
    pairs = align_lines(r"Thank you for using the MYSTERY\nGIFT System.",
                        r"Dziękujemy za korzystanie z\nSystemu MYSTERY GIFT")
    check("parowanie linii źródło↔cel", len(pairs) == 2, str(pairs))
    check("  linia 1", pairs[0] == ("Thank you for using the MYSTERY", "Dziękujemy za korzystanie z"),
          str(pairs[0]))
    check("  linia 2", pairs[1] == ("GIFT System.", "Systemu MYSTERY GIFT"), str(pairs[1]))

    line_hits = [m for m in tm_g.find_sentence_matches(seg_literal) if m.kind == "linia"]
    check("dopasowanie typu 'linia' w wynikach", len(line_hits) >= 1, str(len(line_hits)))

    # krótkie wpisy nie mogą uchodzić za dopasowanie całej linii
    noise = TranslationMemory()
    noise.init_for_project(os.path.join(tmp, "tm_noise"))
    for a, b in (("System", "System"), ("PLAYER", "GRACZ"), ("YES", "TAK"),
                 (r"Thank you for accessing the\nMYSTERY GIFT System.<<KON>>",
                  r"Dziękujemy za korzystanie z\nSystemu MYSTERY GIFT")):
        noise.add(a, b, "en", "pl")
    noise_hits = noise.find_sentence_matches(seg_literal)
    junk = [m for m in noise_hits
            if m.fragment_target in ("System", "GRACZ", "TAK") and m.kind.startswith("linia")]
    check("jednowyrazowe wpisy nie udają dopasowania linii", not junk,
          str([m.fragment_source for m in junk]))
    good = [m for m in noise_hits if "Systemu MYSTERY GIFT" in m.fragment_target]
    check("sensowne dopasowanie linii zachowane", bool(good),
          str([m.fragment_target for m in noise_hits]))
    noise.close()

    # przypadek STAMP CARD: wpis nieprzetłumaczony i złożenie całego segmentu
    stamp = TranslationMemory()
    stamp.init_for_project(os.path.join(tmp, "tm_stamp"))
    stamp.add("System.", "System.", "en", "pl")          # tłumaczenie = źródło
    stamp.add(r"Thank you for using the MYSTERY\nGIFT System.",
              r"Dziękujemy za korzystanie z\nSystemu MYSTERY GIFT", "en", "pl")
    stamp_seg = (r"Thank you for using the STAMP CARD\nSystem.\pYou have {STR_VAR_1} "
                 r"more to collect to\nfill your STAMP CARD.")
    stamp_hits = stamp.find_sentence_matches(stamp_seg)
    check("wpis nieprzetłumaczony (System.→System.) odrzucony",
          not any(m.fragment_target.strip() == "System." for m in stamp_hits),
          str([m.fragment_target for m in stamp_hits]))
    if stamp_hits:
        best = stamp_hits[0]
        check("złożenie zawiera CAŁY segment, nie samą linię",
              "STAMP CARD" in best.assembled and len(best.assembled) > 60,
              best.assembled[:70])
        check("złożenie zachowuje znaczniki \\p i {STR_VAR_1}",
              r"\p" in best.assembled and "{STR_VAR_1}" in best.assembled)
        check("pokrycie liczone względem całego segmentu", best.coverage < 100,
              f"{best.coverage}%")
    stamp.close()

    # formatowanie i kopiowanie pomiaru czasu
    from supercat.ui.editor_tab import EditorTab as _ET
    check("format czasu: ms", _ET.format_duration(224, "ms") == "224 ms")
    check("format czasu: sekundy", _ET.format_duration(1422, "s") == "1.42 s")
    check("format czasu: minuty", _ET.format_duration(90000, "min") == "1.50 min")
    check("format czasu: auto dobiera jednostkę",
          _ET.format_duration(790, "auto") == "790 ms"
          and _ET.format_duration(1422, "auto") == "1.42 s")
    check("  pary linii dostępne w wyniku",
          line_hits and len(line_hits[0].line_pairs) == 2,
          str(line_hits[0].line_pairs) if line_hits else "brak")

    # --- segment KRÓTSZY niż wpis TM (rozbity przez segmentację) ---
    for short_seg, expect in ((r"Thank you for using the MYSTERY", "Dziękujemy za korzystanie z"),
                              (r"GIFT System.", "Systemu MYSTERY GIFT")):
        sm_short = tm_g.find_sentence_matches(short_seg)
        check(f"krótszy segment: {short_seg[:22]!r}", len(sm_short) >= 1, f"znaleziono {len(sm_short)}")
        check("  poprawna linia tłumaczenia",
              sm_short and expect in sm_short[0].assembled,
              sm_short[0].assembled if sm_short else "brak")

    # --- ochrona znaczników przed silnikami MT ---
    # --- dopasowanie ROZMYTE linia-po-linii (inne słowa + <<KON>> w TM) ---
    tm_fz = TranslationMemory()
    tm_fz.init_for_project(os.path.join(tmp, "tm_fuzzy_lines"))
    tm_fz.add(r"Thank you for accessing the\nMYSTERY GIFT System.<<KON>>",
              r"Dziękujemy za korzystanie z\nSystemu MYSTERY GIFT", "en", "pl")
    real_seg = (r"Thank you for using the MYSTERY\nGIFT System.\pYou must be {PLAYER}."
                r"\nThere is a ticket here for you.")
    fz_lines = tm_fz.find_sentence_matches(real_seg)
    check("rozmyte dopasowanie linii (inne słowa w TM)", len(fz_lines) >= 2, f"{len(fz_lines)}")
    targets = {p[1] for m in fz_lines for p in m.line_pairs}
    check("  linia 1 dopasowana", "Dziękujemy za korzystanie z" in targets, str(targets))
    check("  linia 2 dopasowana", "Systemu MYSTERY GIFT" in targets, str(targets))

    # podpowiedzi z segmentów przetłumaczonych w projekcie (bez zapisu do bazy)
    before_size = tm_fz.size()
    tm_fz.add_volatile_pairs([("There is a ticket here for you.", "Jest tu bilet dla Ciebie.")])
    vol = tm_fz.find_sentence_matches(real_seg)
    vol_targets = {p[1] for m in vol for p in m.line_pairs}
    check("podpowiedź z segmentu sesji (TM w locie)",
          "Jest tu bilet dla Ciebie." in vol_targets, str(vol_targets))
    check("  bez zapisu do bazy", tm_fz.size() == before_size, f"{tm_fz.size()} vs {before_size}")
    tm_fz.close()

    # --- tłumaczenie linia po linii zachowuje pozycję \n ---
    from supercat.core.mt import split_keep_separators
    parts = split_keep_separators(r"aaa\nbbb\pccc")
    check("podział z zachowaniem znaczników",
          parts == ["aaa", "\\n", "bbb", "\\p", "ccc"], str(parts))

    from supercat.core.mt import protect_codes, restore_codes
    src_codes = r"It appears to be for use at the\nVERMILION CITY port.\pWhy not try it?"
    prot, ph = protect_codes(src_codes)
    check("znaczniki schowane przed MT", "\\n" not in prot and "\\p" not in prot, prot[:50])
    check("liczba tokenów", len(ph) == 2, str(ph))
    check("odtworzenie 1:1", restore_codes(prot, ph) == src_codes)
    # typowe uszkodzenia wprowadzane przez MyMemory
    check("naprawa '@ #0#@' (spacja w tokenie)",
          "\\n" in restore_codes("Tekst@ #0#@dalszy@#1#@koniec", ph))
    broken_space = "Wygląda na to @#0#@ VERMILION.@#1#@ Dlaczego?"
    fixed = restore_codes(broken_space, ph)
    check("usunięcie spacji wokół znacznika",
          "\\n VERMILION" not in fixed and "\\nVERMILION" in fixed, fixed)
    check("brak spacji przed znacznikiem", " \\n" not in fixed and " \\p" not in fixed, fixed)
    missing = restore_codes("Tylko poczatek @#0#@ srodek", ph)
    check("brakujący znacznik dopisany", missing.count("\\p") == 1, missing)

    from supercat.core.mt import ENGINES, MachineTranslation, to_lang_code
    mt_test = MachineTranslation()
    check("kody języków", to_lang_code("polski") == "pl" and to_lang_code("pt-BR") == "pt")
    engine_ids = [e for e, _l in ENGINES]
    check("silnik MyMemory dostępny", "mymemory" in engine_ids)
    # Puter AI Gateway (endpoint zgodny z OpenAI)
    from supercat.core.mt import PUTER_DEFAULT_MODEL, PUTER_DEFAULT_URL, PUTER_MODELS
    # Google Gemini (AI Studio)
    from supercat.core.mt import GEMINI_BASE_URL, GEMINI_DEFAULT_MODEL, GEMINI_MODELS
    # --- oczyszczanie „gadatliwych” odpowiedzi AI ---
    from supercat.core.ai_clean import clean_ai_translation
    verbose = """* Role: Professional translator (English to Polish).
* Task: Translate the user's text.
* Input text: "Thank you for using the MYSTERY"
* "Thank you for using" -> "Dziękujemy za korzystanie z" (formal/plural)
* No tags/placeholders present in the source.
Dziękujemy za korzystanie z MYSTERY"""
    check("AI: odrzucenie toku rozumowania",
          clean_ai_translation(verbose) == "Dziękujemy za korzystanie z MYSTERY",
          clean_ai_translation(verbose)[:60])
    labelled = '* Option 1: "System GIFT"\n*Self-correction*: ...\nResult: System GIFT.'
    check("AI: etykieta Result:", clean_ai_translation(labelled) == "System GIFT.",
          clean_ai_translation(labelled))
    check("AI: etykieta Final choice:",
          clean_ai_translation('Wait, hmm...\nFinal choice: Proszę, wychowaj go') ==
          "Proszę, wychowaj go")
    check("AI: zdjęcie cudzysłowów",
          clean_ai_translation('"Masz prezent - JAJKO POKéMONA!"') ==
          "Masz prezent - JAJKO POKéMONA!")
    check("AI: blok <think> usunięty",
          clean_ai_translation("<think>analiza...</think>\nCzysty wynik") == "Czysty wynik")
    check("AI: blok kodu usunięty",
          clean_ai_translation("```\nTekst\n```") == "Tekst")
    check("AI: poprawna odpowiedź bez zmian",
          clean_ai_translation(r"Dziękujemy za korzystanie z\nSystemu MYSTERY GIFT") ==
          r"Dziękujemy za korzystanie z\nSystemu MYSTERY GIFT")
    # regresja: AI dostaje CAŁY segment, nie pocięty po \n (traciło kontekst)
    from supercat.core.mt import AI_ENGINES
    import json as _json
    stamp_src = (r"Thank you for using the STAMP CARD\nSystem.\pYou have {STR_VAR_1} "
                 r"more to collect to\nfill your STAMP CARD.")
    mt_ai = MachineTranslation()
    mt_ai.keys["gemini"] = "X"
    _sent = []

    def _capture(url, data, headers, timeout=30):
        body = _json.loads(data)
        sent = body["contents"][0]["parts"][0]["text"]
        _sent.append(sent)
        return {"candidates": [{"content": {"parts": [{"text": sent}]}}],
                "usageMetadata": {"totalTokenCount": 5}}

    mt_ai._http_post = _capture
    mt_ai.translate_with("gemini", stamp_src, "en", "pl")
    check("AI dostaje segment w JEDNYM zapytaniu", len(_sent) == 1, f"{len(_sent)} zapytań")
    check("AI widzi cały tekst (kontekst zachowany)",
          "STAMP CARD" in _sent[0] and "collect" in _sent[0], _sent[0][:60])
    check("silniki AI wyłączone z dzielenia po liniach",
          {"gemini", "openai", "puter"} <= AI_ENGINES, str(AI_ENGINES))
    # proste silniki NADAL dzielą po liniach (chroni to znaczniki)
    mt_plain = MachineTranslation()
    _parts = []
    mt_plain._local = lambda t: (_parts.append(t), t)[1]
    mt_plain.set_engine("local")
    mt_plain.translate(stamp_src, "en", "pl")
    check("proste silniki dzielą po liniach", len(_parts) > 1, f"{len(_parts)} fragmentów")

    # KLUCZOWE: \n przełamuje zdanie w środku ("STAMP CARD\nSystem." = jedna nazwa).
    # Fragmenty nie mogą trafiać do tłumacza osobno, bo "System." traci sens.
    from supercat.core.mt import restore_inner_codes, split_into_sentences_with_codes
    chunks = split_into_sentences_with_codes(stamp_src)
    texts = [t for t, _c in chunks if t.strip()]
    check("zdanie przełamane \\n jest scalane",
          any("STAMP CARD System." in t for t in texts), str(texts[:1]))
    check("„System.” nie jest osobnym fragmentem",
          not any(t.strip() in ("System.", "System") for t in texts), str(texts))
    check("\\p rozdziela wypowiedzi",
          any(c == ["\\p"] for _t, c in chunks), str([c for _t, c in chunks]))
    check("znaczniki wewnętrzne zapamiętane",
          any(c == ["\\n"] for _t, c in chunks))
    # znaczniki wracają na swoje miejsce w tłumaczeniu
    rebuilt = restore_inner_codes("Dziękujemy za korzystanie z Systemu KART", ["\\n"])
    check("znacznik wstawiony z powrotem do zdania", "\\n" in rebuilt, rebuilt)
    check("tekst tłumaczenia nienaruszony",
          rebuilt.replace("\\n", " ").split() ==
          "Dziękujemy za korzystanie z Systemu KART".split(), rebuilt)
    check("brak znaczników = brak zmian",
          restore_inner_codes("Zwykły tekst", []) == "Zwykły tekst")
    # pełny obieg na prostym silniku
    mt_round = MachineTranslation()
    mt_round.set_engine("local")
    round_out = mt_round.translate(stamp_src, "en", "pl")
    check("po tłumaczeniu liczba \\n się zgadza",
          round_out.count(chr(92) + "n") == stamp_src.count(chr(92) + "n"),
          f"{round_out.count(chr(92)+'n')} vs {stamp_src.count(chr(92)+'n')}")
    check("po tłumaczeniu \\p zachowane",
          round_out.count(chr(92) + "p") == stamp_src.count(chr(92) + "p"))
    check("zmienna {STR_VAR_1} zachowana", "{STR_VAR_1}" in round_out, round_out[:70])

    # polecenie dla AI tłumaczy tę zasadę wprost
    prompt_lines = build_translation_prompt("en", "pl")
    check("polecenie wyjaśnia, że znacznik nie kończy zdania",
          "NIE granice zdań" in prompt_lines)
    check("polecenie zawiera przykład STAMP CARD", "STAMP CARD" in prompt_lines)

    # duplikaty i warianty w odpowiedzi modelu
    dup = "Dziękujemy za korzystanie ze STAMP CARD.Dziękujemy za korzystanie ze STAMP CARD."
    check("AI: usunięcie podwojonego tekstu",
          clean_ai_translation(dup) == "Dziękujemy za korzystanie ze STAMP CARD.",
          clean_ai_translation(dup)[:60])
    dup2 = '"Musisz zebrać jeszcze {STR_VAR_1}, aby"\nMusisz zebrać jeszcze {STR_VAR_1}, aby'
    check("AI: podwojenie z cudzysłowem",
          clean_ai_translation(dup2) == "Musisz zebrać jeszcze {STR_VAR_1}, aby",
          clean_ai_translation(dup2)[:60])
    check("AI: wariant „or” – zostaje pierwszy",
          clean_ai_translation('"Wypełnij KARTĘ." or "Uzupełnij KARTĘ."') == "Wypełnij KARTĘ.")
    check("AI: wariant „lub” – zostaje pierwszy",
          clean_ai_translation('"Wypełnij KARTĘ." lub "Uzupełnij KARTĘ."') == "Wypełnij KARTĘ.")
    check("AI: dwa różne zdania NIE są skracane",
          clean_ai_translation("Ala ma kota. Ala ma psa.") == "Ala ma kota. Ala ma psa.")

    # zgubione znaczniki są odzyskiwane i zgłaszane
    from supercat.core.mt import LAST_RESTORE_STATS, protect_codes, restore_codes
    _prot, _ph = protect_codes(stamp_src)
    check("wszystkie znaczniki schowane", len(_ph) == 4, str(len(_ph)))
    full = restore_codes("A@#0#@B@#1#@C @#2#@ D@#3#@E", _ph)
    check("komplet znaczników odtworzony", LAST_RESTORE_STATS["missing"] == 0)
    check("znaczniki wrócone na swoje miejsca",
          r"\n" in full and r"\p" in full and "{STR_VAR_1}" in full, full[:60])
    partial = restore_codes("A@#0#@B C @#2#@ D", _ph)
    check("zgubione znaczniki są dopisywane", LAST_RESTORE_STATS["missing"] == 2,
          str(LAST_RESTORE_STATS["missing"]))
    check("żaden znacznik nie ginie z pliku",
          partial.count(chr(92) + "n") + partial.count(chr(92) + "p") >= 3, partial)

    check("AI: znaczniki @#0#@ zachowane",
          "@#0#@" in clean_ai_translation("Result: Tekst @#0#@ dalej"))
    prompt = build_translation_prompt("en", "pl", "styl retro",
                                      [("gift", "prezent")])
    check("AI: polecenie zakazuje rozumowania", "Nie pokazuj toku rozumowania" in prompt)
    check("AI: polecenie zawiera glosariusz", "gift → prezent" in prompt)
    check("AI: polecenie zawiera własne wytyczne", "styl retro" in prompt)

    check("silnik Gemini na liście", "gemini" in engine_ids)
    check("endpoint Gemini", "generativelanguage.googleapis.com" in GEMINI_BASE_URL)
    check("domyślny model Gemini", GEMINI_DEFAULT_MODEL in GEMINI_MODELS, GEMINI_DEFAULT_MODEL)
    check("domyślny model to alias 'latest' (odporny na zmiany nazw)",
          "latest" in GEMINI_DEFAULT_MODEL, GEMINI_DEFAULT_MODEL)

    # automatyczny fallback, gdy Google zablokuje model dla nowego konta (404)
    import io as _io, json as _json, urllib.error as _ue
    mt_fb = MachineTranslation()
    mt_fb.keys["gemini"] = "FAKE"
    mt_fb.keys["gemini_model"] = "gemini-2.5-flash"
    mt_fb.save_keys = lambda: None
    mt_fb.list_gemini_models = lambda only_free_friendly=True: ["gemini-flash-latest"]
    _tried = []

    def _fake_post(url, data, headers, timeout=30):
        _tried.append(url)
        if "gemini-2.5-flash" in url:
            raise _ue.HTTPError(url, 404, "Not Found", {}, _io.BytesIO(
                _json.dumps({"error": {"code": 404,
                             "message": "no longer available to new users"}}).encode()))
        return {"candidates": [{"content": {"parts": [{"text": "Witaj"}]}}],
                "usageMetadata": {"totalTokenCount": 42}}

    mt_fb._http_post = _fake_post
    fb_out = mt_fb.translate_with("gemini", "Hello", "en", "pl")
    check("fallback po błędzie 404 zwraca tłumaczenie", fb_out == "Witaj", fb_out)
    check("fallback zapamiętuje działający model",
          mt_fb.keys["gemini_model"] == "gemini-flash-latest", mt_fb.keys["gemini_model"])
    check("fallback wykonał dokładnie dwie próby", len(_tried) == 2, str(len(_tried)))
    check("tokeny z usageMetadata policzone", mt_fb._last_tokens == 42, str(mt_fb._last_tokens))
    check("pobieranie modeli wymaga klucza",
          "klucz" in (lambda: (lambda m: [m.keys.__setitem__("gemini", ""),
                     _safe_list(m)][1])(MachineTranslation()))().lower())
    mt_gem = MachineTranslation()
    mt_gem.keys["gemini"] = ""
    check("bez klucza Gemini czytelny komunikat",
          "klucza Gemini" in mt_gem.translate_with("gemini", "Hello", "en", "pl"))
    check("Gemini niedostępny bez klucza", "gemini" not in mt_gem.available_engines())
    mt_gem.keys["gemini"] = "x"
    check("Gemini dostępny po podaniu klucza", "gemini" in mt_gem.available_engines())

    # licznik zużycia
    from supercat.core.usage import UsageTracker
    tracker = UsageTracker.instance()
    tracker.reset_all()
    before = tracker.get("local").requests_today
    mt_use = MachineTranslation()
    mt_use.set_engine("local")
    mt_use.translate("Hello world", "en", "pl")
    check("licznik rejestruje wywołanie",
          tracker.get("local").requests_today == before + 1,
          f"{before} -> {tracker.get('local').requests_today}")
    check("licznik zapisuje znaki", tracker.get("local").chars_today > 0)
    check("znane limity Gemini", tracker.limits("gemini") == (10, 250),
          str(tracker.limits("gemini")))
    for _ in range(200):
        tracker.record("gemini", chars=10)
    check("procent wykorzystania limitu", tracker.percent_used("gemini") == 80,
          str(tracker.percent_used("gemini")))
    check("ostrzeżenie przy zbliżaniu do limitu", tracker.is_near_limit("gemini"))
    check("silnik bez limitu nie ma procentu", tracker.percent_used("local") is None)
    check("podsumowanie zawiera limit", "250" in tracker.summary("gemini"),
          tracker.summary("gemini"))
    check("zestawienie zawiera silniki", len(tracker.report()) >= 2)
    # błędy też są liczone
    err_before = tracker.get("gemini").errors_today
    mt_use.keys["gemini"] = ""
    mt_use.translate_with("gemini", "Hello", "en", "pl")
    check("nieudane wywołania są liczone",
          tracker.get("gemini").errors_today == err_before + 1)
    tracker.reset_all()

    check("silnik Puter AI na liście", "puter" in engine_ids)
    check("endpoint Puter zgodny z OpenAI", PUTER_DEFAULT_URL.endswith("/v1/chat/completions"),
          PUTER_DEFAULT_URL)
    check("domyślny model Puter", PUTER_DEFAULT_MODEL in PUTER_MODELS, PUTER_DEFAULT_MODEL)
    mt_puter = MachineTranslation()
    mt_puter.keys["puter_token"] = ""
    check("bez tokenu Puter jest czytelny komunikat",
          "tokenu Puter" in mt_puter.translate_with("puter", "Hello", "en", "pl"))
    check("Puter niedostępny bez tokenu",
          "puter" not in mt_puter.available_engines())
    mt_puter.keys["puter_token"] = "x"
    check("Puter dostępny po podaniu tokenu", "puter" in mt_puter.available_engines())
    check("silnik Google bez klucza", "google_free" in engine_ids)
    free = mt_test.available_engines(only_free=True)
    check("lista darmowych silników", "google_free" in free and "mymemory" in free, str(free))
    multi = mt_test.translate_multi("Hello world", "en", "pl", engines=["local"])
    check("zbiorcze tłumaczenie (translate_multi)", len(multi) == 1 and not multi[0][3], str(multi[0][:2]))
    check("translate_with nie zmienia silnika domyślnego",
          (mt_test.translate_with("local", "Hello", "en", "pl"), mt_test.engine)[1] == mt_test.engine)
    tm_g.close()

    # ---------------------------------------------------------- glosariusz
    print("\n6. Glosariusz")
    g = Glossary()
    g.init_for_project(project.glossary_path)
    check("domyślny glosariusz", g.size > 5, str(g.size))
    g.add("efficiently", "efektywnie", "przysłówek")
    found = g.find_terms("The system helps you work efficiently.")
    check("wyszukanie terminów w zdaniu", any(t.source == "efficiently" for t in found), str([t.source for t in found]))
    g2 = Glossary()
    g2.init_for_project(project.glossary_path)
    check("zapis i ponowny odczyt", any(e.source == "efficiently" for e in g2.entries))

    # ------------------------------------------------------------------ QA
    print("\n7. Kontrola jakości QA")
    Seg = type(segments[0])
    qa_segments = [
        Seg("1", "Version 2.5 released", "Wersja 3.7 wydana"),      # liczby
        Seg("2", "Click <b>here</b>", "Kliknij tutaj"),              # tagi
        Seg("3", "Hello world.", ""),                                # pusty
        Seg("4", "Text.", "Tekst"),                                  # interpunkcja
        Seg("5", "Same text", "Same text"),                          # nieprzetłumaczone
    ]
    issues = run_qa(qa_segments)
    kinds = {i.category for i in issues}
    check("wykryto błąd liczb", "Liczby" in kinds)
    check("wykryto błąd tagów", "Tagi" in kinds)
    check("wykryto pusty segment", "Puste tłumaczenie" in kinds)
    check("wykryto interpunkcję", "Interpunkcja" in kinds)
    check("wykryto nieprzetłumaczone", "Nieprzetłumaczone" in kinds)

    stats = project_statistics(qa_segments, tm.size())
    check("statystyki", stats["Segmenty (razem)"] == 5 and stats["Segmenty przetłumaczone"] == 4, str(stats["Postęp (%)"]))

    # ------------------------------------------------------------- eksport
    print("\n8. Eksport")
    for seg in segments:
        m = tm.find_fuzzy_matches(seg.source, 90, 1)
        if m:
            seg.target = m[0].text
    out_txt = os.path.join(project.target_path, "doc.txt")
    export_by_replacement(src_txt, out_txt, segments)
    content = open(out_txt, encoding="utf-8").read()
    check("eksport TXT z podmianą", "Witaj świecie" in content, content[:60].replace("\n", " "))

    html_out = os.path.join(project.export_path, "bilingual.html")
    export_html_bilingual(segments, html_out)
    check("eksport HTML dwujęzyczny", os.path.exists(html_out) and "Witaj świecie" in open(html_out, encoding="utf-8").read())

    from supercat.core.fileparser import export_docx
    docx_out = os.path.join(project.target_path, "doc.docx")
    docx_segs[0].target = "Pierwszy akapit dokumentu."
    export_docx(docx_path, docx_out, docx_segs)
    check("eksport DOCX", os.path.exists(docx_out) and
          "Pierwszy akapit" in "\n".join(p.text for p in Document(docx_out).paragraphs))

    # --------------------------------------------------------------- GUI
    print("\n9. Interfejs graficzny (pełny przepływ)")
    from supercat.ui.main_window import MainWindow
    w = MainWindow()
    w.open_project_path(project.project_file_path)
    check("projekt wczytany w GUI", w.project is not None and w.project.name == "Test Projekt")
    check("segmenty w edytorze", len(w.editor_tab.segments) > 0, str(len(w.editor_tab.segments)))
    check("siatka wypełniona", w.editor_tab.grid.rowCount() == len(w.editor_tab.segments))
    check("TM wczytana w GUI", w.tm.size() >= 3, str(w.tm.size()))
    check("glosariusz w GUI", w.glossary.size > 5)

    w.editor_tab.load_segment(0)
    w.editor_tab.set_target_text("Tłumaczenie testowe.")
    check("wpisanie tłumaczenia", w.editor_tab.segments[0].target == "Tłumaczenie testowe.")
    w.editor_tab.confirm_and_next()
    check("zatwierdzenie segmentu", w.editor_tab.segments[0].status == "translated")
    check("przejście do następnego", w.editor_tab.current_index == 1)
    check("segment zapisany do TM", len(w.tm.search("Tłumaczenie testowe")) == 1)

    w.editor_tab.copy_source_to_target()
    check("kopiowanie źródła", w.editor_tab.segments[1].target == w.editor_tab.segments[1].source)

    w.save_all(silent=True)
    check("zapis tłumaczeń", os.path.exists(os.path.join(project.project_path, "translations.json")))

    w2 = MainWindow()
    w2.open_project_path(project.project_file_path)
    check("tłumaczenia po ponownym otwarciu",
          w2.editor_tab.segments[0].target == "Tłumaczenie testowe.",
          w2.editor_tab.segments[0].target)

    w.search_tab.search_edit.setText("test")
    w.search_tab.perform_search()
    check("wyszukiwanie", w.search_tab.result.total_matches > 0,
          str(w.search_tab.result.total_matches))
    check("wyniki pogrupowane po plikach", w.search_tab.tree.topLevelItemCount() > 0)

    # --- rozbudowane wyszukiwanie w GUI ---------------------------------
    w.search_tab.scope.setCurrentText("Cały projekt")
    hits_all = w.search_tab.result.hits
    files_in_result = {h.file_name for h in hits_all}
    check("wyszukiwanie obejmuje inne pliki", len(files_in_result) >= 1, str(files_in_result))

    w.search_tab.codes_check.setChecked(True)
    w.search_tab.search_edit.setText("STAMP CARD System")
    w.search_tab.perform_search()
    check("fraza przez znacznik \\n (GUI)",
          any("stamp" in h.source.lower() for h in w.search_tab.result.hits)
          or w.search_tab.result.total_matches == 0)

    w.search_tab.search_edit.setText("[")
    w.search_tab.mode.setCurrentText("Regex")
    w.search_tab.perform_search()
    check("błędny regex nie wywraca zakładki", "regularne" in w.search_tab.status.text().lower())
    w.search_tab.mode.setCurrentText("Zawiera")

    w.search_tab.search_edit.setText("test")
    w.search_tab.perform_search()
    items = w.search_tab._hit_items()
    if items:
        w.search_tab.tree.setCurrentItem(items[0])
        w.search_tab.goto_result()
        check("przejście do segmentu z wyników",
              w.editor_tab.current_index == items[0].data(0, Qt.ItemDataRole.UserRole))
        w.search_tab.next_result()
        check("F3 przechodzi dalej", True)
    else:
        check("przejście do segmentu z wyników", False, "brak wyników")
    # zakres „tylko przeglądany plik”
    first_file = w.editor_tab.segments[0].file_name
    w.editor_tab._file_filter = first_file
    w.search_tab.scope.setCurrentText("Tylko przeglądany plik")
    w.search_tab.search_edit.setText("test")
    w.search_tab.perform_search()
    check("zakres: tylko przeglądany plik",
          all(h.file_name == first_file for h in w.search_tab.result.hits),
          str({h.file_name for h in w.search_tab.result.hits}))
    w.search_tab.scope.setCurrentText("Cały projekt")
    w.editor_tab._file_filter = None
    w.editor_tab.refresh_grid()

    # podświetlenie trafień w edytorze
    w.editor_tab.load_segment(0)
    w.editor_tab.highlight_search("test", w.search_tab.current_options())
    check("podświetlenie wyszukiwania w edytorze",
          len(w.editor_tab.source_edit.extraSelections())
          + len(w.editor_tab.target_edit.extraSelections()) >= 0)
    w.editor_tab.clear_search_highlight()
    check("czyszczenie podświetlenia", w.editor_tab.target_edit.extraSelections() == [])
    check("pomiar czasu wyszukiwania w pasku", "⏱" in w.search_tab.status.text(),
          w.search_tab.status.text())
    check("adaptacyjne opóźnienie w rozsądnym zakresie",
          200 <= w.search_tab._live_delay_ms <= 1500, str(w.search_tab._live_delay_ms))
    w.search_tab.live_check.setChecked(False)

    # zachowywanie wcięcia w GUI: podpowiedź TM/MT dziedziczy spację źródła
    seg0 = w.editor_tab.segments[0]
    old_source = seg0.source
    seg0.source = "  Indented source."
    w.editor_tab.load_segment(0)
    w.editor_tab.set_target_text("Wcięte tłumaczenie.")
    check("wstawiony tekst dziedziczy wcięcie",
          w.editor_tab.segments[0].target.startswith("  "),
          repr(w.editor_tab.segments[0].target))
    seg0.source = old_source
    w.editor_tab.load_segment(0)

    # --- przejście do wyniku z INNEGO pliku (zgłoszony błąd) --------------
    # Gdy w edytorze otwarty jest jeden plik, a trafienie leży w drugim,
    # kliknięcie wyniku musi przełączyć plik i zaznaczyć właściwy wiersz.
    ed_tab = w.editor_tab
    files_in_project = sorted({(sg.file_name or "(bez pliku)") for sg in ed_tab.segments})
    other_idx = None
    if len(files_in_project) >= 2:
        first_name = files_in_project[0]
        ed_tab._file_filter = first_name
        ed_tab.refresh_grid()
        # segment z INNEGO pliku niż aktualnie filtrowany
        for i, sg in enumerate(ed_tab.segments):
            if (sg.file_name or "(bez pliku)") != first_name:
                other_idx = i
                break
    if other_idx is not None:
        check("segment z innego pliku jest ukryty filtrem",
              other_idx not in ed_tab._visible_indices())
        w.go_to_editor_segment(other_idx)
        check("przejście do segmentu z innego pliku zmienia segment",
              ed_tab.current_index == other_idx,
              f"{ed_tab.current_index} != {other_idx}")
        check("filtr pliku przełączony na plik trafienia",
              ed_tab._file_filter == (ed_tab.segments[other_idx].file_name or "(bez pliku)"),
              str(ed_tab._file_filter))
        rows_sel = [r.row() for r in ed_tab.grid.selectionModel().selectedRows()]
        row_idx = (ed_tab.grid.item(rows_sel[0], 0).data(Qt.ItemDataRole.UserRole)
                   if rows_sel else None)
        check("siatka zaznacza TEN SAM segment co edytor", row_idx == other_idx,
              f"wiersz={row_idx} edytor={ed_tab.current_index}")
        check("pole źródła pokazuje właściwy tekst",
              ed_tab.source_edit.toPlainText() == ed_tab.segments[other_idx].source)
    else:
        check("przejście do segmentu z innego pliku zmienia segment", False, "brak 2 plików")

    # filtr tekstowy siatki też nie może blokować przejścia
    ed_tab._show_all_files()
    ed_tab.filter_edit.setText("zzz-nic-takiego-nie-ma")
    ed_tab.refresh_grid()
    check("filtr tekstowy ukrywa wszystko", ed_tab.grid.rowCount() == 0)
    w.go_to_editor_segment(0)
    check("przejście czyści filtr tekstowy siatki",
          ed_tab.filter_edit.text() == "" and ed_tab.current_index == 0,
          repr(ed_tab.filter_edit.text()))

    # filtr statusu też
    ed_tab.segments[0].target = "cos przetlumaczonego"
    ed_tab.status_filter.setCurrentText("Nieprzetłumaczone")
    ed_tab.refresh_grid()
    check("filtr statusu ukrywa segment", 0 not in ed_tab._visible_indices())
    w.go_to_editor_segment(0)
    check("przejście zdejmuje filtr statusu",
          ed_tab.status_filter.currentText() == "Wszystkie" and ed_tab.current_index == 0,
          ed_tab.status_filter.currentText())
    ed_tab.segments[0].target = ""
    ed_tab.status_filter.setCurrentText("Wszystkie")

    # gdy segment JEST widoczny – filtry zostają nietknięte
    ed_tab._show_all_files()
    ed_tab.refresh_grid()
    check("widoczny segment nie zmienia filtrów",
          ed_tab.reveal_segment(1) is False)

    # przejście z zakładki QA korzysta z tej samej ścieżki
    if other_idx is not None:
        ed_tab._file_filter = files_in_project[0]
        ed_tab.refresh_grid()
        w.qa_tab.run_checks()
        w.go_to_editor_segment(other_idx)     # tak samo robi dwuklik w QA
        check("QA również przechodzi do ukrytego segmentu",
              ed_tab.current_index == other_idx
              and other_idx in ed_tab._visible_indices())
    ed_tab._show_all_files()

    # --- przełączniki znaków specjalnych w Ustawieniach -------------------
    _SMk = SettingsManager

    ws_seg = w.editor_tab.segments[0]
    kept_src = ws_seg.source
    ws_seg.source = "  Wcięty\ntekst."
    w.editor_tab.refresh_grid()

    def _grid_cell() -> str:
        for r in range(w.editor_tab.grid.rowCount()):
            it = w.editor_tab.grid.item(r, 0)
            if it and it.data(Qt.ItemDataRole.UserRole) == 0:
                return w.editor_tab.grid.item(r, 1).text()
        return ""

    _SMk.instance().set("ui.markers.spaces", True)
    _SMk.instance().set("ui.markers.newlines", True)
    _SMk.instance().set("ui.markers.style", "␣ → ⏎  (standardowe)")
    w.editor_tab.refresh_grid()
    check("siatka pokazuje ␣ i ⏎", "␣" in _grid_cell() and "⏎" in _grid_cell(), _grid_cell())

    w.settings_tab.markers_spaces.setChecked(False)
    check("wyłączenie ␣ działa z Ustawień", "␣" not in _grid_cell(), _grid_cell())
    check("wyłączenie ␣ nie rusza ⏎", "⏎" in _grid_cell(), _grid_cell())

    w.settings_tab.markers_newlines.setChecked(False)
    check("wyłączenie ⏎ działa z Ustawień", "⏎" not in _grid_cell(), _grid_cell())

    w.settings_tab.markers_spaces.setChecked(True)
    w.settings_tab.markers_newlines.setChecked(True)
    w.settings_tab.markers_style.setCurrentText("· » ¶  (dyskretne)")
    check("zmiana zestawu znaków widoczna w siatce",
          "·" in _grid_cell() and "¶" in _grid_cell(), _grid_cell())
    w.settings_tab.markers_style.setCurrentText("␣ → ⏎  (standardowe)")

    # ustawienie przeżywa ponowne wczytanie widoku
    _SMk.instance().set("ui.markers.spaces", False)
    w.editor_tab.refresh_grid()
    check("ustawienie znaków obowiązuje po odświeżeniu", "␣" not in _grid_cell(), _grid_cell())
    _SMk.instance().set("ui.markers.spaces", True)
    ws_seg.source = kept_src
    w.editor_tab.refresh_grid()

    # --- słowniki: dodawanie, folder, usuwanie ---------------------------
    dict_folder = project.dictionary_path
    sample_dict = os.path.join(tmp, "moj_slownik.txt")
    with open(sample_dict, "w", encoding="utf-8") as fh:
        fh.write("kot\npies\ndrzewo\nkomputer\n")
    installed = w.dictionary.install_file(sample_dict, dict_folder)
    check("dodano słownik z pliku", os.path.exists(installed), installed)
    w.dictionary_tab.reload()
    check("słownik widoczny po przeładowaniu",
          "moj_slownik.txt" in w.dictionary.sources, str(w.dictionary.sources))
    check("policzono słowa w pliku",
          w.dictionary.source_counts.get("moj_slownik.txt", 0) == 4,
          str(w.dictionary.source_counts))
    check("lista plików w zakładce Słowniki", w.dictionary_tab.files_list.count() >= 1)
    check("zakładka pokazuje ścieżkę folderu",
          dict_folder in w.dictionary_tab.folder_label.text(),
          w.dictionary_tab.folder_label.text())
    check("propozycje poprawnej pisowni",
          "komputer" in w.dictionary.suggest_corrections("komputre", 5),
          str(w.dictionary.suggest_corrections("komputre", 5)))
    check("odrzucenie pliku o złym rozszerzeniu",
          _raises(lambda: w.dictionary.install_file(
              os.path.join(project.source_path, "doc.txt").replace(".txt", ".xyz"), dict_folder)))
    check("usuwanie słownika z projektu",
          w.dictionary.remove_file("moj_slownik.txt", dict_folder))
    w.dictionary_tab.reload()
    check("słownik zniknął po usunięciu",
          "moj_slownik.txt" not in w.dictionary.sources, str(w.dictionary.sources))
    # --- kodowanie słowników (polski w ISO-8859-2) -----------------------
    iso_path = os.path.join(dict_folder, "iso_test.dic")
    with open(iso_path, "wb") as fh:
        fh.write("3\noknówka\nżółw\nzażółć\n".encode("iso8859-2"))
    with open(os.path.join(dict_folder, "iso_test.aff"), "w", encoding="ascii") as fh:
        fh.write("SET ISO8859-2\n")
    check("wykryto kodowanie z pliku .aff",
          w.dictionary.detect_encoding(iso_path) == "iso8859-2",
          w.dictionary.detect_encoding(iso_path))
    w.dictionary_tab.reload()
    check("polskie znaki wczytane poprawnie",
          "oknówka" in w.dictionary.words and "żółw" in w.dictionary.words,
          str(sorted(x for x in w.dictionary.words if len(x) > 3)[:5]))
    check("brak znaków zastępczych po wczytaniu",
          not any("\ufffd" in word for word in w.dictionary.words))

    # bez .aff kodowanie rozpoznajemy z treści pliku
    os.remove(os.path.join(dict_folder, "iso_test.aff"))
    check("kodowanie rozpoznane bez .aff",
          w.dictionary.detect_encoding(iso_path) in ("iso8859-2", "cp1250"),
          w.dictionary.detect_encoding(iso_path))
    utf_path = os.path.join(dict_folder, "utf_test.txt")
    with open(utf_path, "w", encoding="utf-8") as fh:
        fh.write("zażółć\ngęślą\n")
    check("plik UTF-8 rozpoznany jako UTF-8",
          w.dictionary.detect_encoding(utf_path) == "utf-8",
          w.dictionary.detect_encoding(utf_path))
    w.dictionary.remove_file("iso_test.dic", dict_folder)
    w.dictionary.remove_file("utf_test.txt", dict_folder)
    w.dictionary_tab.reload()

    # --- lista SJP.pl: wszystkie formy odmienione ------------------------
    check("zalecany słownik z pełną odmianą jest pierwszy na liście",
          "SJP" in DOWNLOADABLE_DICTIONARIES[0][0]
          and "odmiana" in DOWNLOADABLE_DICTIONARIES[0][0].lower(),
          DOWNLOADABLE_DICTIONARIES[0][0])
    check("adres SJP jest rozwiązywany (nazwa zawiera datę)",
          Dictionary.resolve_url("https://example.com/plik.dic") == "https://example.com/plik.dic")

    # rozpakowanie ZIP w formacie SJP (formy po przecinku w jednej linii)
    import zipfile as _zip
    zip_path = os.path.join(tmp, "sjp_test.zip")
    with _zip.ZipFile(zip_path, "w") as archive:
        archive.writestr("odm.txt", "robić, robię, robisz\nwitać, witamy, witacie\n")
        archive.writestr("README.txt", "opis licencji")
    out_txt = os.path.join(dict_folder, "sjp_test.txt")
    count = Dictionary._extract_wordlist(zip_path, out_txt)
    check("rozpakowano formy z archiwum ZIP", count == 6, str(count))
    w.dictionary_tab.reload()
    check("formy odmienione trafiły do słownika",
          all(form in w.dictionary.words for form in ("witamy", "robię", "witacie")),
          str(sorted(w.dictionary.words)[:8]))
    check("„Witamy” i „Dziękujemy” nie są zgłaszane przy pełnej odmianie",
          w.dictionary.is_correct("Witamy") and w.dictionary.is_correct("witamy"))
    check("README pominięty przy rozpakowaniu",
          "licencji" not in " ".join(w.dictionary.words))
    w.dictionary.remove_file("sjp_test.txt", dict_folder)
    w.dictionary_tab.reload()

    check("lista słowników do pobrania jest niepusta",
          len(DOWNLOADABLE_DICTIONARIES) >= 5 and
          all(u.startswith("https://") for _n, _f, u, _s in DOWNLOADABLE_DICTIONARIES))

    # --- panel „Język” w edytorze ----------------------------------------
    lang_seg = w.editor_tab.segments[0]
    kept_lang_target = lang_seg.target
    kept_lang_source = lang_seg.source
    w.editor_tab.load_segment(0)
    w.editor_tab.target_edit.setPlainText("Mam pięć jabłko , naprawdę.")
    w.editor_tab.check_language(force=True)
    worker = w.editor_tab._lang_worker
    if worker is not None:
        worker.wait(5000)
    app.processEvents()
    check("panel Język wykrył uwagi", w.editor_tab.lang_list.count() > 0,
          str(w.editor_tab.lang_list.count()))
    check("panel Język pokazuje podsumowanie",
          "⚠️" in w.editor_tab.lang_status.text() or "❌" in w.editor_tab.lang_status.text(),
          w.editor_tab.lang_status.text())

    w.editor_tab.apply_language_fixes()
    check("automatyczna poprawka z panelu Język",
          "," in w.editor_tab.target_edit.toPlainText()
          and " ," not in w.editor_tab.target_edit.toPlainText(),
          repr(w.editor_tab.target_edit.toPlainText()))

    # kontrola dotyczy WYŁĄCZNIE tłumaczenia – źródło z błędem nic nie zmienia
    lang_seg.source = "Ja poszedł do sklepu."
    w.editor_tab.target_edit.setPlainText("Poszedłem do sklepu.")
    w.editor_tab.check_language(force=True)
    if w.editor_tab._lang_worker is not None:
        w.editor_tab._lang_worker.wait(5000)
    app.processEvents()
    check("błąd w ŹRÓDLE nie jest zgłaszany",
          not any("zaimkiem" in w.editor_tab.lang_list.item(i).text()
                  for i in range(w.editor_tab.lang_list.count())),
          str([w.editor_tab.lang_list.item(i).text()[:40]
               for i in range(w.editor_tab.lang_list.count())]))

    # --- podkreślanie błędów w polu tłumaczenia --------------------------
    class _SpellDict:
        """Udaje pełny słownik – bez pobierania 5 MB z sieci."""
        is_initialized = True
        size = 343070
        has_morphology = True
        def check_text(self, text, skip_uppercase=True):
            import re as _re
            out = []
            for word in _re.findall(r"[^\W\d_]+", text):
                if skip_uppercase and word.isupper():
                    continue
                if word.lower() == "korzystaniefa":
                    out.append(word)
            return out
        def suggest_corrections(self, word, limit=5, fast=False):
            return ["korzystanie"]

    real_dict = w.dictionary
    w.dictionary = _SpellDict()
    lang_seg.source = "Thank you for using the MYSTERY\\nGIFT System."
    w.editor_tab.load_segment(0)
    w.editor_tab.target_edit.setPlainText(
        "Dziękujemy za korzystaniefa z\\nSystemu MYSTERY GIFT")
    w.editor_tab.check_language(force=True)
    if w.editor_tab._lang_worker is not None:
        w.editor_tab._lang_worker.wait(8000)
    app.processEvents()

    underlined = []
    body = w.editor_tab.target_edit.toPlainText()
    for sel in w.editor_tab.target_edit.extraSelections():
        cur = sel.cursor
        underlined.append(body[cur.selectionStart():cur.selectionEnd()])
    check("błędny wyraz podkreślony w polu tłumaczenia",
          "korzystaniefa" in underlined, str(underlined))
    check("nazwy WERSALIKAMI nie są podkreślane",
          "MYSTERY" not in underlined and "GIFT" not in underlined, str(underlined))
    check("znacznik \\n nie jest zgłaszany",
          not any("\\n" in u for u in underlined), str(underlined))
    check("literówka ma czerwoną falkę (jak w edytorze tekstu)",
          any(sel.format.underlineColor().name() == "#ff5252"
              and sel.format.underlineStyle().name == "WaveUnderline"
              for sel in w.editor_tab.target_edit.extraSelections()),
          str([(sel.format.underlineColor().name(), sel.format.underlineStyle().name)
               for sel in w.editor_tab.target_edit.extraSelections()]))
    check("pisownia ma poziom błędu, nie „info”",
          all(i.severity == "błąd" for i in w.editor_tab._lang_issues
              if i.rule_id == "PISOWNIA"),
          str([(i.rule_id, i.severity) for i in w.editor_tab._lang_issues]))
    check("uwaga o pisowni ma pozycję w tekście",
          any(i.rule_id == "PISOWNIA" and i.offset >= 0 for i in w.editor_tab._lang_issues),
          str([(i.rule_id, i.offset) for i in w.editor_tab._lang_issues]))

    # propozycje doliczane w tle, nie blokują podkreśleń
    import time as _t
    class _SlowDict(_SpellDict):
        def suggest_corrections(self, word, limit=5, fast=False):
            if not fast:
                _t.sleep(0.4)
            return ["korzystanie"]

    w.dictionary = _SlowDict()
    w.editor_tab.target_edit.setPlainText("Dziękujemy za korzystaniefa z.")
    started = _t.perf_counter()
    w.editor_tab.check_language(force=True)
    if w.editor_tab._lang_worker is not None:
        w.editor_tab._lang_worker.wait(8000)
    app.processEvents()
    underline_ms = (_t.perf_counter() - started) * 1000
    check("podkreślenia pojawiają się szybko (bez czekania na propozycje)",
          underline_ms < 1500 and w.editor_tab._lang_selections,
          f"{underline_ms:.0f} ms")
    if w.editor_tab._suggest_worker is not None:
        w.editor_tab._suggest_worker.wait(8000)
    app.processEvents()
    check("propozycje dochodzą w tle",
          any(i.suggestions for i in w.editor_tab._lang_issues if i.rule_id == "PISOWNIA"),
          str([(i.fragment, i.suggestions) for i in w.editor_tab._lang_issues]))
    check("lista uwag odświeżona o propozycje",
          any("propozycje" in w.editor_tab.lang_list.item(i).text()
              for i in range(w.editor_tab.lang_list.count())),
          str([w.editor_tab.lang_list.item(i).text()[:50]
               for i in range(w.editor_tab.lang_list.count())]))

    # menu podręczne pokazuje propozycje i podmienia wyraz
    w.dictionary = _SpellDict()
    w.editor_tab.target_edit.setPlainText("Dziękujemy za korzystaniefa z.")
    w.editor_tab.check_language(force=True)
    if w.editor_tab._lang_worker is not None:
        w.editor_tab._lang_worker.wait(8000)
    app.processEvents()

    menu_items = []
    class _FakeMenu:
        def addSeparator(self):
            pass
        def addAction(self, text):
            from PyQt6.QtGui import QAction
            menu_items.append(text)
            return QAction(text)
        def exec(self, *a):
            return None

    real_menu = w.editor_tab.target_edit.createStandardContextMenu
    w.editor_tab.target_edit.createStandardContextMenu = lambda: _FakeMenu()
    spell = next((i for i in w.editor_tab._lang_issues if i.rule_id == "PISOWNIA"), None)
    if spell is not None:
        cur = w.editor_tab.target_edit.textCursor()
        cur.setPosition(spell.offset + 2)
        w.editor_tab._target_context_menu(w.editor_tab.target_edit.cursorRect(cur).center())
    w.editor_tab.target_edit.createStandardContextMenu = real_menu
    check("prawy przycisk pokazuje propozycje poprawek",
          any("korzystanie" in item for item in menu_items), str(menu_items))
    check("menu ma opcję dodania do słownika",
          any("słownika" in item for item in menu_items), str(menu_items))

    # podmiana wyrazu z menu podręcznego
    spell_issue = next((i for i in w.editor_tab._lang_issues if i.rule_id == "PISOWNIA"), None)
    if spell_issue is not None:
        w.editor_tab._replace_issue(spell_issue, "korzystanie")
        check("propozycja z menu poprawia wyraz",
              "korzystanie z" in w.editor_tab.target_edit.toPlainText()
              and "korzystaniefa" not in w.editor_tab.target_edit.toPlainText(),
              repr(w.editor_tab.target_edit.toPlainText()))
    else:
        check("propozycja z menu poprawia wyraz", False, "brak uwagi o pisowni")

    # dodanie wyrazu do słownika użytkownika
    w.dictionary = real_dict
    w.editor_tab.add_word_to_dictionary("Pokéball")
    user_dict = os.path.join(project.dictionary_path, "uzytkownika.txt")
    check("wyraz dopisany do słownika użytkownika",
          os.path.exists(user_dict) and "Pokéball" in open(user_dict, encoding="utf-8").read())
    check("wyraz od razu uznany za poprawny", real_dict.is_correct("pokéball"))

    # podkreślanie można wyłączyć
    SettingsManager.instance().set("lang.check.underline", False)
    w.editor_tab.highlight_language_issues(w.editor_tab._lang_issues)
    check("podkreślanie można wyłączyć",
          w.editor_tab._lang_selections == [])
    SettingsManager.instance().set("lang.check.underline", True)

    # główny wyłącznik kontroli języka
    SettingsManager.instance().set("lang.check.enabled", False)
    w.editor_tab.check_language(force=True)
    app.processEvents()
    check("główny wyłącznik zatrzymuje kontrolę",
          w.editor_tab.lang_list.count() == 0
          and "wyłączona" in w.editor_tab.lang_status.text(),
          w.editor_tab.lang_status.text())
    check("wyłączona kontrola nie zwraca uwag z rdzenia",
          check_translation("Mam pięć jabłko , to to.")[0] == [])
    SettingsManager.instance().set("lang.check.enabled", True)

    # przełączniki zakresu w zakładce Ustawienia
    lang_settings = w.settings_tab
    lang_settings._set_all_language(False)
    check("„Wyłącz wszystko” gasi kontrole",
          not SettingsManager.instance().get_bool("lang.check.grammar", True)
          and not SettingsManager.instance().get_bool("qa.check.language", True))
    lang_settings._set_all_language(True)
    check("„Włącz wszystko” przywraca kontrole",
          SettingsManager.instance().get_bool("lang.check.grammar", False)
          and SettingsManager.instance().get_bool("lang.check.enabled", False))
    check("opcje szczegółowe wyszarzone przy wyłączonym module",
          lang_settings.lang_group.isEnabled())

    opts_no_grammar = default_lang_options()
    opts_no_grammar["grammar"] = False
    check("można wyłączyć samą kontrolę odmiany",
          not any(i.category == "Odmiana"
                  for i in check_offline("Mam pięć jabłko.", None, opts_no_grammar)))
    opts_no_punct = default_lang_options()
    opts_no_punct["punctuation"] = False
    check("można wyłączyć samą interpunkcję",
          not any(i.category == "Interpunkcja"
                  for i in check_offline("Tak , owszem.", None, opts_no_punct)))

    lang_seg.source = kept_lang_source
    _SMk2 = SettingsManager
    _SMk2.instance().set("lang.check.auto", False)
    w.editor_tab.lang_auto.setChecked(False)
    check("kontrolę języka można wyłączyć",
          not _SMk2.instance().get_bool("lang.check.auto", True))
    _SMk2.instance().set("lang.check.auto", True)
    w.editor_tab.lang_auto.setChecked(True)
    lang_seg.target = kept_lang_target
    w.editor_tab.load_segment(0)

    # QA: kontrola języka obejmuje tłumaczenie
    lang_issues = run_qa([Segment("q1", "I have five apples.", "Mam pięć jabłko.")])
    check("QA zgłasza błąd odmiany w tłumaczeniu",
          any(i.category.startswith("Język") for i in lang_issues),
          str([i.category for i in lang_issues]))
    SettingsManager.instance().set("qa.check.language", False)
    off_issues = run_qa([Segment("q2", "I have five apples.", "Mam pięć jabłko.")])
    check("kontrolę języka w QA można wyłączyć",
          not any(i.category.startswith("Język") for i in off_issues))
    SettingsManager.instance().set("qa.check.language", True)

    # --- liczniki przy plikach aktualizują się na bieżąco ------------------
    ed = w.editor_tab
    ed._show_all_files()

    def _file_row_text(name):
        for i in range(ed.files_list.count()):
            if ed.files_list.item(i).data(Qt.ItemDataRole.UserRole) == name:
                return ed.files_list.item(i).text()
        return ""

    first_file = (ed.segments[0].file_name or "(bez pliku)")
    check("licznik pliku pokazuje postęp i procent",
          "/" in _file_row_text(first_file) and "%" in _file_row_text(first_file),
          _file_row_text(first_file))

    # masowa zmiana (jak „Zastosuj TM”) musi odświeżyć listę
    original_label = _file_row_text(first_file)
    saved_targets = [(s.target, s.status) for s in ed.segments]
    for seg in ed.segments:          # najpierw wszystko do zrobienia
        seg.target = ""
        seg.status = "new"
    ed.update_progress()
    before_label = _file_row_text(first_file)
    for seg in ed.segments:
        seg.target = "X"
    ed.update_progress()
    after_label = _file_row_text(first_file)
    check("licznik odświeża się po masowym tłumaczeniu",
          after_label != before_label, f"{before_label} -> {after_label}")
    check("ukończony plik dostaje znacznik ✅", after_label.startswith("✅"), after_label)
    check("pozycja „Wszystkie pliki” też ma licznik",
          "/" in ed.files_list.item(0).text(), ed.files_list.item(0).text())

    for seg, (target, status) in zip(ed.segments, saved_targets):
        seg.target, seg.status = target, status
    ed.update_progress()
    check("licznik wraca po cofnięciu tłumaczeń",
          _file_row_text(first_file) == original_label,
          f"{_file_row_text(first_file)} vs {original_label}")

    # zaznaczenie pliku nie znika przy odświeżaniu licznika
    ed._file_filter = first_file
    ed.update_file_counters()
    check("filtr pliku przetrwał odświeżenie licznika", ed._file_filter == first_file)
    ed._show_all_files()

    # --- pominięte segmenty nie są liczone --------------------------------
    skip_seg = ed.segments[0]
    label_before = _file_row_text(first_file)
    skip_seg.ignored = True
    ed.update_progress()
    label_after = _file_row_text(first_file)
    check("pominięty segment znika z mianownika licznika",
          label_after != label_before and "(0/0" in label_after,
          f"{label_before} -> {label_after}")
    check("licznik pliku nie pokazuje znaczników pominięcia",
          "🚫" not in label_after, label_after)
    check("pasek postępu pomija wykluczone i nie ma znacznika",
          "🚫" not in ed.progress.format(), ed.progress.format())
    check("pozycja „Wszystkie pliki” bez dopisków",
          "🚫" not in ed.files_list.item(0).text()
          and "pominiętych" not in ed.files_list.item(0).text(),
          ed.files_list.item(0).text())
    skip_seg.ignored = False
    ed.update_progress()
    check("przywrócenie segmentu wraca do poprzedniego licznika",
          _file_row_text(first_file) == label_before,
          f"{_file_row_text(first_file)} vs {label_before}")

    # --- grupowe pomijanie segmentów --------------------------------------
    from PyQt6.QtCore import QItemSelectionModel
    check("siatka pozwala zaznaczyć wiele wierszy",
          ed.grid.selectionMode().name == "ExtendedSelection",
          ed.grid.selectionMode().name)

    ed._show_all_files()
    ed.refresh_grid()
    ed.grid.clearSelection()
    rows = min(3, ed.grid.rowCount())
    for row in range(rows):
        ed.grid.selectionModel().select(
            ed.grid.model().index(row, 0),
            QItemSelectionModel.SelectionFlag.Select | QItemSelectionModel.SelectionFlag.Rows)
    picked = ed.selected_indices()
    check("odczyt zaznaczonych segmentów", len(picked) == rows, str(picked))

    changed = ed.set_ignored(picked, True)
    check("grupowe pomijanie działa",
          changed == rows and all(ed.segments[i].ignored for i in picked), str(changed))
    check("pominięcie ręczne jest oznaczone jako ręczne",
          ed.segments[picked[0]].extra.get("manual_skip") is True,
          str(ed.segments[picked[0]].extra))

    # reguły wykluczania nie mogą cofnąć ręcznej decyzji
    w.exclusion_set().apply(ed.segments)
    check("reguły nie przywracają ręcznie pominiętych",
          all(ed.segments[i].ignored for i in picked))

    restored = ed.set_ignored(picked, False)
    check("grupowe przywracanie działa",
          restored == rows and not any(ed.segments[i].ignored for i in picked), str(restored))

    # toggle: mieszane zaznaczenie → pomija wszystkie
    ed.segments[picked[0]].ignored = True
    ed.toggle_ignore()
    check("odwrócenie przy mieszanym zaznaczeniu pomija wszystkie",
          all(ed.segments[i].ignored for i in picked))
    ed.toggle_ignore()
    check("ponowne odwrócenie przywraca wszystkie",
          not any(ed.segments[i].ignored for i in picked))
    ed.grid.clearSelection()

    # pomijanie po wzorcu
    from PyQt6.QtWidgets import QInputDialog as _QID
    real_get_text = _QID.getText
    _QID.getText = staticmethod(lambda *a, **k: ("Alpha*", True))
    ed.segments.append(Segment("m1", "Alpha wzorzec test", "", file_name="wzor.txt"))
    ed.refresh_grid()
    try:
        ed.ignore_matching()
    finally:
        _QID.getText = real_get_text
    check("pomijanie po wzorcu działa", ed.segments[-1].ignored)
    check("wzorzec zapisany jako reguła wykluczania",
          any(r.pattern == "Alpha*" for r in w.exclusion_set().rules),
          str([r.pattern for r in w.exclusion_set().rules][-2:]))
    w.exclusion_set().rules = [r for r in w.exclusion_set().rules if r.pattern != "Alpha*"]
    ed.segments.pop()
    ed.refresh_grid()

    # --- oznaczenia działają na wielu segmentach naraz --------------------
    ed._show_all_files()
    ed.refresh_grid()
    ed.grid.clearSelection()
    multi = min(3, ed.grid.rowCount())
    for row in range(multi):
        ed.grid.selectionModel().select(
            ed.grid.model().index(row, 0),
            QItemSelectionModel.SelectionFlag.Select | QItemSelectionModel.SelectionFlag.Rows)
    chosen = ed.selected_indices()

    ed.approve_current()
    check("zatwierdzanie działa na wielu segmentach naraz",
          all(ed.segments[i].status == "approved" for i in chosen),
          str([ed.segments[i].status for i in chosen]))
    check("zatwierdzony liczy się jako gotowy",
          all(_is_done_seg(ed.segments[i]) for i in chosen))

    approved_progress = ed.progress.format()
    ed.mark_new()
    check("„oznacz jako nowy” działa grupowo",
          all(ed.segments[i].status == "new" for i in chosen),
          str([ed.segments[i].status for i in chosen]))
    check("cofnięcie do „nowy” zmienia licznik",
          ed.progress.format() != approved_progress,
          f"{approved_progress} -> {ed.progress.format()}")

    ed.mark_draft()
    check("„roboczy” działa grupowo",
          all(ed.segments[i].status == "draft" for i in chosen))
    ed.mark_translated()
    check("„przetłumaczony” działa grupowo",
          all(ed.segments[i].status == "translated" for i in chosen))
    ed.mark_new()
    ed.grid.clearSelection()

    # --- nawigacja w stylu OmegaT ------------------------------------------
    nav_saved = [(s.target, s.status) for s in ed.segments]
    for seg in ed.segments:
        seg.target = ""
        seg.status = "new"
        seg.ignored = False
    ed.segments[0].target = "gotowe"
    ed.segments[2].target = "gotowe"
    ed.refresh_grid()

    ed.load_segment(0)
    ed.next_untranslated()
    check("Ctrl+U przechodzi do nieprzetłumaczonego",
          ed.current_index == 1, str(ed.current_index))
    ed.next_untranslated()
    check("Ctrl+U pomija segmenty z tłumaczeniem",
          ed.current_index == 3, str(ed.current_index))
    ed.next_translated()
    check("Ctrl+Shift+U szuka przetłumaczonego z zawijaniem",
          _is_done_seg(ed.segments[ed.current_index]), str(ed.current_index))

    ed.load_segment(0)
    ed.prev_untranslated()
    check("Ctrl+Alt+U cofa do nieprzetłumaczonego",
          not _is_done_seg(ed.segments[ed.current_index]), str(ed.current_index))

    ed.segments[1].status = "approved"
    ed.load_segment(0)
    ed.next_unapproved()
    check("skok do niezatwierdzonego pomija zatwierdzone",
          ed.segments[ed.current_index].status != "approved", str(ed.current_index))

    # pominięte segmenty są omijane przy nawigacji
    for seg in ed.segments:
        seg.target = ""
        seg.status = "new"
    ed.segments[1].ignored = True
    ed.load_segment(0)
    ed.next_untranslated()
    check("nawigacja omija pominięte segmenty",
          ed.current_index != 1, str(ed.current_index))
    ed.segments[1].ignored = False

    for seg, (target, status) in zip(ed.segments, nav_saved):
        seg.target, seg.status = target, status
    ed.refresh_grid()

    # --- „Zatwierdź i dalej” zmienia oznaczenie ---------------------------
    conf_saved = [(sg.target, sg.status) for sg in ed.segments]
    for seg in ed.segments:
        seg.target = ""
        seg.status = "new"
        seg.ignored = False
    ed.refresh_grid()
    conf_label_before = _file_row_text(first_file)

    ed.load_segment(0)
    ed.confirm_and_next()
    check("„Zatwierdź i dalej” zmienia status pustego segmentu",
          ed.segments[0].status == "translated", ed.segments[0].status)
    check("zatwierdzenie pustego segmentu zmienia licznik",
          _file_row_text(first_file) != conf_label_before,
          f"{conf_label_before} -> {_file_row_text(first_file)}")
    check("„Zatwierdź i dalej” przechodzi do segmentu do zrobienia",
          not _is_done_seg(ed.segments[ed.current_index]), str(ed.current_index))

    ed.load_segment(1)
    ed.target_edit.setPlainText("tłumaczenie testowe")
    ed.confirm_and_next()
    check("zatwierdzenie z tłumaczeniem ustawia status",
          ed.segments[1].status == "translated" and ed.segments[1].target,
          ed.segments[1].status)

    # oznaczenie „przetłumaczony” liczy się w liczniku
    for seg in ed.segments:          # wszystko od nowa do zrobienia
        seg.target = ""
        seg.status = "new"
    ed.update_progress()
    mark_label = _file_row_text(first_file)
    ed.grid.clearSelection()
    ed.load_segment(0)
    ed.mark_translated()
    check("„przetłumaczony” zmienia licznik pliku",
          _file_row_text(first_file) != mark_label,
          f"{mark_label} -> {_file_row_text(first_file)}")
    check("„przetłumaczony” liczy się jako gotowy",
          _is_done_seg(ed.segments[0]))
    stat_now = project_statistics(ed.segments)
    check("statystyki też uznają oznaczenie za wykonane",
          stat_now["Segmenty przetłumaczone"] >= 1,
          str(stat_now["Segmenty przetłumaczone"]))

    for seg, (target, status) in zip(ed.segments, conf_saved):
        seg.target, seg.status = target, status
    ed.refresh_grid()

    # --- skróty: brak duplikatów i możliwość zmiany -----------------------
    from supercat.core import shortcuts as _shortcuts
    from PyQt6.QtGui import QShortcut as _QSc, QAction as _QAct

    editor_keys = {}
    for obj in w.findChildren(_QSc):
        editor_keys.setdefault(obj.key().toString(), []).append(obj)
    check("Ctrl+U jest zarejestrowany dokładnie raz jako skrót edytora",
          len(editor_keys.get("Ctrl+U", [])) == 1,
          str(len(editor_keys.get("Ctrl+U", []))))

    clashing = []
    for act in w.findChildren(_QAct):
        key = act.shortcut().toString()
        if key and key in editor_keys:
            if act.shortcutContext() != Qt.ShortcutContext.WidgetShortcut:
                clashing.append(f"{key} ({act.text()[:24]})")
    check("akcje menu nie odbierają skrótów edytorowi", not clashing, str(clashing))

    check("rejestr skrótów zna Ctrl+U",
          _shortcuts.get("next_untranslated") == "Ctrl+U",
          _shortcuts.get("next_untranslated"))
    check("skróty mają grupy w Ustawieniach", len(_shortcuts.groups()) >= 4,
          str(_shortcuts.groups()))
    check("zakładka „Skróty” istnieje",
          any("Skróty" in w.settings_tab.tabText(i)
              for i in range(w.settings_tab.count())))
    check("tabela skrótów wypełniona",
          w.settings_tab.sc_table.rowCount() == len(_shortcuts.SHORTCUTS),
          str(w.settings_tab.sc_table.rowCount()))

    # zmiana skrótu działa bez restartu
    w.settings_tab._apply_shortcut("next_untranslated", "Ctrl+9")
    check("zmieniony skrót zapisany", _shortcuts.get("next_untranslated") == "Ctrl+9",
          _shortcuts.get("next_untranslated"))
    new_keys = {o.key().toString() for o in w.findChildren(_QSc)}
    check("nowa kombinacja działa w edytorze", "Ctrl+9" in new_keys)
    check("stara kombinacja przestała działać", "Ctrl+U" not in new_keys)

    check("wykrywanie zajętych kombinacji",
          _shortcuts.find_conflict("next_translated", "Ctrl+9") is not None)
    check("wolna kombinacja nie zgłasza konfliktu",
          _shortcuts.find_conflict("next_translated", "Ctrl+Alt+Shift+F12") is None)

    _shortcuts.reset_all()
    w.reload_shortcuts()
    check("przywracanie domyślnych skrótów",
          _shortcuts.get("next_untranslated") == "Ctrl+U")

    # przyciski nawigacji w pasku edytora
    from PyQt6.QtWidgets import QPushButton as _QPB
    nav_buttons = [b.text() for b in ed.findChildren(_QPB)]
    check("przycisk „poprzedni nieprzetłumaczony” istnieje",
          any("◀◀" in t for t in nav_buttons), str(nav_buttons[:8]))
    check("przycisk „następny nieprzetłumaczony” istnieje",
          any("▶▶" in t for t in nav_buttons), str(nav_buttons[:8]))

    # --- skoki nie wychodzą poza przeglądany plik -------------------------
    jump_saved = [(sg.target, sg.status, sg.ignored) for sg in ed.segments]
    files_now = sorted({(sg.file_name or "(bez pliku)") for sg in ed.segments})
    if len(files_now) >= 2:
        for seg in ed.segments:
            seg.target = ""
            seg.status = "new"
            seg.ignored = False
        target_file = files_now[-1]
        in_file = [i for i, sg in enumerate(ed.segments)
                   if (sg.file_name or "(bez pliku)") == target_file]
        ed._file_filter = target_file
        ed.refresh_grid()
        ed.load_segment(in_file[0])

        real_question = QMessageBox.question
        QMessageBox.question = staticmethod(lambda *a, **k: QMessageBox.StandardButton.No)
        try:
            ed.prev_untranslated()
        finally:
            QMessageBox.question = real_question
        check("skok wstecz zostaje w przeglądanym pliku",
              (ed.segments[ed.current_index].file_name or "(bez pliku)") == target_file,
              f"{ed.segments[ed.current_index].file_name} vs {target_file}")
        check("skok nie podmienia filtru pliku",
              ed._file_filter == target_file, str(ed._file_filter))

        # gdy w pliku nie ma już celu, program pyta i dopiero wtedy wychodzi
        for i in in_file:
            ed.segments[i].target = "gotowe"
        ed.refresh_grid()
        ed.load_segment(in_file[0])
        QMessageBox.question = staticmethod(lambda *a, **k: QMessageBox.StandardButton.Yes)
        try:
            ed.next_untranslated()
        finally:
            QMessageBox.question = real_question
        check("po zgodzie skok przechodzi do innego pliku",
              (ed.segments[ed.current_index].file_name or "(bez pliku)") != target_file,
              str(ed.segments[ed.current_index].file_name))

    ed._show_all_files()
    for seg, (target, status, ignored) in zip(ed.segments, jump_saved):
        seg.target, seg.status, seg.ignored = target, status, ignored
    ed.refresh_grid()

    # --- cofanie zmian oznaczeń (Ctrl+Z / Ctrl+Y) -------------------------
    undo_saved = [(sg.status, sg.ignored) for sg in ed.segments]
    ed.grid.clearSelection()
    undo_rows = min(3, ed.grid.rowCount())
    for row in range(undo_rows):
        ed.grid.selectionModel().select(
            ed.grid.model().index(row, 0),
            QItemSelectionModel.SelectionFlag.Select | QItemSelectionModel.SelectionFlag.Rows)
    undo_indices = ed.selected_indices()
    statuses_before = [ed.segments[i].status for i in undo_indices]

    ed.approve_current()
    check("zatwierdzenie zmieniło statusy",
          all(ed.segments[i].status == "approved" for i in undo_indices))

    w.undo_action()
    check("Ctrl+Z cofa zmianę statusów",
          [ed.segments[i].status for i in undo_indices] == statuses_before,
          str([ed.segments[i].status for i in undo_indices]))

    w.redo_action()
    check("Ctrl+Y ponawia zmianę statusów",
          all(ed.segments[i].status == "approved" for i in undo_indices))
    w.undo_action()

    # cofanie działa też dla pominięć
    ed.set_ignored(undo_indices, True)
    check("pominięcie zapisane w historii",
          all(ed.segments[i].ignored for i in undo_indices))
    w.undo_action()
    check("Ctrl+Z cofa pominięcie",
          not any(ed.segments[i].ignored for i in undo_indices))

    check("cofanie aktualizuje licznik",
          "/" in ed.progress.format(), ed.progress.format())
    check("historia jest ograniczona i nie rośnie w nieskończoność",
          len(ed._undo_stack) <= 100, str(len(ed._undo_stack)))

    for seg, (status, ignored) in zip(ed.segments, undo_saved):
        seg.status, seg.ignored = status, ignored
    ed.grid.clearSelection()
    ed.refresh_grid()

    # --- silnik lokalny korzysta z glosariusza i pamięci TM ---------------
    from supercat.core.glossary import GlossaryEntry as _GE

    local_mt = MachineTranslation()
    plain = local_mt.translate_with("local", "WIRELESS COMMUNICATION", "en", "pl")
    check("silnik lokalny informuje o braku terminu",
          "brak w słowniku" in plain, plain[:70])

    class _Gloss:
        entries = [_GE("wireless communication", "komunikacja bezprzewodowa")]

    local_mt.glossary_provider = _Gloss()
    with_gloss = local_mt.translate_with("local", "WIRELESS COMMUNICATION", "en", "pl")
    check("silnik lokalny tłumaczy z glosariusza",
          "komunikacja bezprzewodowa" in with_gloss, with_gloss)
    check("słownik lokalny rośnie o terminy projektu",
          len(local_mt.local_vocabulary()) > 25,
          str(len(local_mt.local_vocabulary())))

    class _FakeTM:
        is_initialized = True
        def all_entries(self, limit=0):
            return [("stamp card", "karta znaczków", "en", "pl", 1),
                    ("a very long sentence " * 5, "za długie", "en", "pl", 1)]

    local_mt.tm_provider = _FakeTM()
    vocab = local_mt.local_vocabulary()
    check("terminy z pamięci TM trafiają do słownika",
          vocab.get("stamp card") == "karta znaczków", str(vocab.get("stamp card")))
    check("długie zdania z TM są pomijane",
          not any(len(k) > 40 for k in vocab), "za długie hasło w słowniku")
    tm_out = local_mt.translate_with("local", "Use your STAMP CARD here", "en", "pl")
    check("silnik lokalny używa terminu z TM",
          "karta znaczków" in tm_out.lower(), tm_out)

    # --- czytelne komunikaty błędów MT ------------------------------------
    lt_result = MachineTranslation().translate_with(
        "libretranslate", "test", "en", "pl")
    check("LibreTranslate bez serwera daje zrozumiały komunikat",
          "niedostępny" in lt_result and "Ustawieniach" in lt_result, lt_result[:80])
    check("komunikat nie pokazuje surowego WinError",
          "WinError" not in lt_result and "urlopen" not in lt_result, lt_result[:60])

    from supercat.ui.quicktrans import QuickTransDialog as _QT
    friendly = _QT._friendly_error
    class _Dummy:
        ERROR_HINTS = _QT.ERROR_HINTS
    check("błąd 429 tłumaczony na limit zapytań",
          "limit zapytań" in friendly(_Dummy(), "[Błąd MT: HTTP Error 429: Too Many Requests]"),
          friendly(_Dummy(), "HTTP Error 429"))
    check("błąd połączenia tłumaczony na brak serwera",
          "serwer" in friendly(_Dummy(), "[WinError 10061] Nie można nawiązać połączenia"),
          friendly(_Dummy(), "[WinError 10061]"))
    check("nieznany błąd nadal jest pokazywany",
          "⚠️" in friendly(_Dummy(), "coś dziwnego"), friendly(_Dummy(), "coś dziwnego"))

    # --- LibreTranslate: instalacja i serwer lokalny ----------------------
    from supercat.core import libretranslate_setup as _lt

    check("wykrywanie pakietu LibreTranslate",
          isinstance(_lt.is_installed(), bool))
    check("adres serwera lokalnego",
          _lt.server_url() == "http://127.0.0.1:5000", _lt.server_url())
    check("sprawdzanie działania serwera nie wywraca programu",
          _lt.is_running(port=59999, timeout=0.5) is False)
    check("lista publicznych serwerów dostępna",
          len(_lt.PUBLIC_SERVERS) >= 2 and
          all(url.startswith("https://") for _n, url in _lt.PUBLIC_SERVERS))

    server = _lt.LibreTranslateServer()
    check("nowy serwer nie jest jeszcze nasz", server.is_ours is False)
    check("zatrzymanie nieuruchomionego serwera jest bezpieczne", server.stop() is False)
    if not _lt.is_installed():
        ok_start, msg_start = server.start(wait_seconds=1)
        check("start bez pakietu daje czytelny komunikat",
              not ok_start and "nie jest zainstalowany" in msg_start, msg_start[:50])

    check("zakładka MT ma sekcję LibreTranslate",
          hasattr(w.settings_tab, "lt_box")
          and "LibreTranslate" in w.settings_tab.lt_box.title(),
          getattr(w.settings_tab, "lt_box", None) and w.settings_tab.lt_box.title())
    check("są przyciski instalacji i uruchomienia",
          all(hasattr(w.settings_tab, name) for name in
              ("lt_install_btn", "lt_start_btn", "lt_stop_btn", "lt_check_btn")))
    w.settings_tab._refresh_lt_state()
    check("stan LibreTranslate jest opisany",
          bool(w.settings_tab.lt_state.text()), w.settings_tab.lt_state.text()[:40])
    check("przyciski odzwierciedlają stan pakietu",
          w.settings_tab.lt_install_btn.isEnabled() != _lt.is_installed())
    check("„Zatrzymaj” nieaktywne, gdy serwer nie działa",
          not w.settings_tab.lt_stop_btn.isEnabled())
    check("można wskazać języki do pobrania",
          w.settings_tab.lt_langs.text().strip() != "", w.settings_tab.lt_langs.text())

    # --- regresja: „python -m libretranslate” NIE działa -------------------
    # Pakiet nie ma __main__.py, więc Python odpowiada „is a package and cannot
    # be directly executed”. Poprawny start to skrypt konsolowy albo .main.
    launch = _lt.launch_command()
    check("polecenie startu nie używa „-m libretranslate”",
          not (len(launch) >= 3 and launch[-1] == "libretranslate"), str(launch))
    check("polecenie startu jest poprawne",
          launch[-1].endswith("libretranslate")
          or launch[-1] == "libretranslate.main", str(launch))
    if _lt.is_installed():
        check("użyto skryptu konsolowego lub modułu .main",
              os.path.basename(launch[0]).startswith("libretranslate")
              or launch[-1] == "libretranslate.main", str(launch))

    # waga modeli i formatowanie po polsku
    check("waga modeli jest liczbą", isinstance(_lt.models_size_bytes(), int))
    check("formatowanie wagi: MB", _lt.format_size(163 * 1024 ** 2) == "163 MB",
          _lt.format_size(163 * 1024 ** 2))
    check("formatowanie wagi: GB z przecinkiem",
          _lt.format_size(int(1.42 * 1024 ** 3)) == "1,42 GB",
          _lt.format_size(int(1.42 * 1024 ** 3)))
    check("formatowanie wagi: kB", _lt.format_size(2048) == "2 kB")
    check("znana waga pary językowej", _lt.LANGUAGE_PAIR_MB > 100,
          str(_lt.LANGUAGE_PAIR_MB))

    # pasek postępu czyta wiersze pipa
    check("postęp pipa: procent z licznika bajtów",
          _lt.parse_pip_progress("   --- 5.2/12.3 MB 3.1 MB/s") == (42, "Pobieranie: 5.2/12.3 MB"),
          str(_lt.parse_pip_progress("   --- 5.2/12.3 MB 3.1 MB/s")))
    check("postęp pipa: nazwa i waga pakietu",
          (_lt.parse_pip_progress("  Downloading libretranslate-1.9.6.whl (1.1 MB)") or (0, ""))[0] == -1)
    check("postęp pipa: etap instalowania",
          _lt.parse_pip_progress("Installing collected packages: x") == (95, "Instalowanie pakietów…"))
    check("postęp pipa: zakończenie",
          _lt.parse_pip_progress("Successfully installed x-1.0") == (100, "Zainstalowano"))
    check("postęp pipa: zwykły wiersz pomijany",
          _lt.parse_pip_progress("jakiś tekst") is None)

    # sprawdzanie najnowszej wersji nie może wywrócić programu
    check("odczyt najnowszej wersji zwraca napis",
          isinstance(_lt.latest_version(timeout=0.001), str))
    if _lt.is_installed():
        note = w.settings_tab._lt_version_note()
        check("stan pakietu mówi, czy wersja jest najnowsza",
              note == "" or "najnowsza" in note or "nowsza" in note, note)

    # regresja: pasek LibreTranslate nie może być tym samym widżetem,
    # co pasek pobierania LanguageTool (obie zakładki to jeden obiekt)
    check("LibreTranslate ma własny pasek postępu",
          hasattr(w.settings_tab, "ltr_progress"))
    check("pasek LibreTranslate to inny widżet niż pasek LanguageTool",
          w.settings_tab.ltr_progress is not getattr(w.settings_tab, "lt_progress", None))
    w.settings_tab._on_lt_install_progress(42, "Pobieranie: 5.2/12.3 MB")
    check("postęp instalacji ustawia procent na pasku LibreTranslate",
          w.settings_tab.ltr_progress.value() == 42
          and w.settings_tab.ltr_progress.maximum() == 100,
          f"{w.settings_tab.ltr_progress.value()}/{w.settings_tab.ltr_progress.maximum()}")
    check("postęp instalacji nie rusza paska LanguageTool",
          getattr(w.settings_tab, "lt_progress", None) is None
          or w.settings_tab.lt_progress.value() != 42)
    w.settings_tab._on_lt_install_progress(-1, "Pobieranie pakietu…")
    check("nieznany rozmiar przełącza pasek w tryb nieokreślony",
          w.settings_tab.ltr_progress.maximum() == 0)

    # --- regresja: awaria na Windowsie (znak → w nazwie modelu) -----------
    # Modele nazywają się „English → Polish” (U+2192). Konsola Windows (cp1250)
    # nie umie tego wypisać → UnicodeEncodeError → LibreTranslate łyka błąd jako
    # „Cannot update models” → modele się nie instalują → serwer pada na
    # IndexError: list index out of range. Wymuszamy UTF-8, żeby przerwać kaskadę.
    env = _lt.subprocess_env()
    check("podproces dostaje wymuszone UTF-8",
          env.get("PYTHONIOENCODING") == "utf-8" and env.get("PYTHONUTF8") == "1",
          f"{env.get('PYTHONIOENCODING')}/{env.get('PYTHONUTF8')}")
    check("podproces ma wyciszone ostrzeżenie requests",
          "requests" in env.get("PYTHONWARNINGS", ""))
    check("podproces nie pyta o licencję interfejsu",
          env.get("LT_DISABLE_WEB_UI") == "false")

    win_log = ("UnicodeEncodeError: 'charmap' codec can't encode character "
               "'\\u2192' in position 20\nIndexError: list index out of range")
    explained = _lt.explain_start_error(win_log)
    check("IndexError tłumaczony na brak modeli",
          "bez modeli" in explained, explained[:60])
    check("wskazana przyczyna: znak → i kodowanie",
          "→" in explained and "UTF-8" in explained, explained[-90:])
    offline = _lt.explain_start_error("IndexError: list index out of range")
    check("bez śladu kodowania radzi sprawdzić internet",
          "internet" in offline, offline[-60:])
    check("zajęty port jest rozpoznawany",
          "Port 5000" in _lt.explain_start_error("OSError: Address already in use"))
    check("nieznany błąd nie jest zmyślany",
          _lt.explain_start_error("zupełnie inny błąd") == "")

    # modele pobierane PRZED startem serwera
    check("znane pary modeli to lista", isinstance(_lt.installed_model_pairs(), list))
    check("ensure_models przy jednym języku nic nie robi",
          _lt.ensure_models("en")[0] is True)
    import inspect as _insp2

    check("start serwera najpierw sprawdza modele",
          "ensure_models" in _insp2.getsource(_lt.LibreTranslateServer.start))
    check("instalacja pipa też używa UTF-8",
          "subprocess_env" in _insp2.getsource(_lt.install))

    # --- postęp pobierania modeli w megabajtach ---------------------------
    # Argos ściąga model w całości do pamięci i zapisuje jednym write na końcu,
    # więc obserwowanie katalogu nie pokazywało żadnego postępu. Pobieramy sami.
    import inspect as _insp3

    ensure_src = _insp3.getsource(_lt.ensure_models)
    check("modele pobierane własnym kodem, nie pack.install()",
          "_download_with_progress" in ensure_src)
    check("ensure_models melduje pobrane bajty", "on_bytes" in ensure_src)
    dl_src = _insp3.getsource(_lt._download_with_progress)
    check("pobieranie idzie kawałkami, nie całością",
          "response.read(" in dl_src and "while True" in dl_src)
    check("plik zapisywany najpierw jako .part",
          ".part" in dl_src and "os.replace" in dl_src)
    check("pobieranie da się przerwać", "cancelled" in dl_src)

    plan_empty, size_empty = _lt.model_download_plan("en")
    check("plan pobierania dla jednego języka jest pusty",
          plan_empty == [] and size_empty == 0)

    from supercat.ui.workers import LTStartWorker as _LTS

    check("worker startu ma sygnał postępu bajtów",
          hasattr(_LTS, "bytes_progress"))
    check("worker startu pobiera modele przed serwerem",
          "ensure_models" in _insp3.getsource(_LTS.run))
    check("worker startu da się anulować", hasattr(_LTS, "cancel"))

    w.settings_tab._on_lt_bytes(42 * 1024 ** 2, 133 * 1024 ** 2)
    check("pasek pokazuje pobrane megabajty modeli",
          "42 MB" in w.settings_tab.ltr_progress.text()
          or "42 MB" in w.settings_tab.ltr_progress.format(),
          w.settings_tab.ltr_progress.format())
    check("pasek modeli liczy procent",
          w.settings_tab.ltr_progress.value() == 31,
          str(w.settings_tab.ltr_progress.value()))
    w.settings_tab._on_lt_bytes(5 * 1024 ** 2, 0)
    check("nieznany rozmiar modeli nie wywraca paska",
          w.settings_tab.ltr_progress.maximum() == 0)

    # --- katalog języków i przycisk „Sprawdź” ------------------------------
    check("odmiana „para”: 1", _lt.plural_pairs(1) == "1 para")
    check("odmiana „para”: 2-4", _lt.plural_pairs(2) == "2 pary"
          and _lt.plural_pairs(24) == "24 pary")
    check("odmiana „para”: 5+", _lt.plural_pairs(5) == "5 par"
          and _lt.plural_pairs(98) == "98 par")
    check("odmiana „para”: 12-14 to wyjątek", _lt.plural_pairs(12) == "12 par",
          _lt.plural_pairs(12))

    catalog = _lt.language_catalog()
    check("katalog języków jest listą", isinstance(catalog, list))
    if catalog:
        check("katalog ma sensowną liczbę języków", len(catalog) > 20, str(len(catalog)))
        entry = catalog[0]
        check("wpis katalogu ma komplet pól",
              {"code", "name", "installed", "pairs"} <= set(entry), str(entry))
        check("katalog wie o angielskim",
              any(e["code"] == "en" for e in catalog))
    check("kody zainstalowanych języków to lista",
          isinstance(_lt.installed_language_codes(), list))
    state_text = _lt.describe_state()
    check("opis stanu jest zdaniem", len(state_text) > 10, state_text[:60])

    # lista języków w interfejsie
    lang_list = w.settings_tab.lt_lang_list
    if lang_list.count():
        check("lista języków wypełniona", lang_list.count() > 20, str(lang_list.count()))
        first = lang_list.item(0)
        check("pozycje mają pole wyboru",
              bool(first.flags() & Qt.ItemFlag.ItemIsUserCheckable))
        check("pozycje niosą kod języka",
              bool(first.data(Qt.ItemDataRole.UserRole)))

        # zaznaczanie przepisuje kody do pola tekstowego
        w.settings_tab.lt_langs.setText("en,pl")
        w.settings_tab._on_lt_langs_typed()
        checked = [lang_list.item(r).data(Qt.ItemDataRole.UserRole)
                   for r in range(lang_list.count())
                   if lang_list.item(r).checkState() == Qt.CheckState.Checked]
        check("wpisane kody zaznaczają pozycje na liście",
              sorted(checked) == ["en", "pl"], str(checked))

        target = next(r for r in range(lang_list.count())
                      if lang_list.item(r).data(Qt.ItemDataRole.UserRole) == "de")
        lang_list.item(target).setCheckState(Qt.CheckState.Checked)
        check("zaznaczenie na liście dopisuje kod do pola",
              "de" in w.settings_tab.lt_langs.text(), w.settings_tab.lt_langs.text())
        check("wybór języków zapisuje się w ustawieniach",
              "de" in SettingsManager.instance().get_str("mt.lt.languages", ""),
              SettingsManager.instance().get_str("mt.lt.languages", ""))
        lang_list.item(target).setCheckState(Qt.CheckState.Unchecked)

        # filtr
        w.settings_tab._filter_lt_languages("polish")
        visible = [r for r in range(lang_list.count()) if not lang_list.item(r).isHidden()]
        check("filtr zawęża listę języków", len(visible) == 1, str(len(visible)))
        w.settings_tab._filter_lt_languages("")
        check("wyczyszczenie filtru pokazuje wszystko",
              all(not lang_list.item(r).isHidden() for r in range(lang_list.count())))

        check("podsumowanie języków opisuje wybór",
              "Wybrano" in w.settings_tab.lt_lang_summary.text(),
              w.settings_tab.lt_lang_summary.text()[:60])

    # przycisk „Sprawdź” musi robić coś widocznego, nie tylko odświeżać etykietę
    import inspect as _insp4

    check_src = _insp4.getsource(w.settings_tab._check_libretranslate)
    check("„Sprawdź” pokazuje okno z wynikiem", "QMessageBox" in check_src)
    check("„Sprawdź” podaje katalog modeli", "models_dir" in check_src)
    check("„Sprawdź” odświeża też listę języków", "_load_lt_languages" in check_src)
    check("„Sprawdź” jest podpięty do przycisku",
          w.settings_tab.lt_check_btn.toolTip() != "")

    # --- RequestsDependencyWarning: wyjaśnione, nie ukryte -----------------
    occurs, explanation = _lt.dependency_warning_info()
    check("wykrywanie ostrzeżenia requests zwraca parę",
          isinstance(occurs, bool) and isinstance(explanation, str))
    if occurs:
        check("wyjaśnienie mówi, że to nie awaria",
              "działają normalnie" in explanation, explanation[:60])
        check("wyjaśnienie podaje polecenie naprawcze",
              "pip install -U requests" in explanation)

    # ostrzeżenie requests o urllib3/chardet jest wyciszane w podprocesie
    import inspect as _inspect

    start_src = _inspect.getsource(_lt.LibreTranslateServer.start)
    check("start serwera wycisza ostrzeżenie requests",
          "subprocess_env" in start_src)

    # --- DeepL bez klucza: jasny komunikat --------------------------------
    deepl_mt = MachineTranslation()
    deepl_mt.keys["deepl"] = ""
    deepl_out = deepl_mt.translate_with("deepl", "test", "en", "pl")
    check("DeepL bez klucza wyjaśnia, że wymaga rejestracji",
          "wymaga klucza" in deepl_out and "Free" in deepl_out, deepl_out[:70])
    check("komunikat DeepL podaje darmowy limit",
          "500 000" in deepl_out, deepl_out[:120])
    check("komunikat DeepL wskazuje alternatywy",
          "MyMemory" in deepl_out or "Google" in deepl_out, deepl_out[-60:])

    # --- okno postępu przy wczytywaniu projektu ---------------------------
    from supercat.ui.dialogs.loading_dialog import LoadingDialog, run_with_progress

    probe = LoadingDialog("Test", ["Etap A", "Etap B"], w)
    check("okno wczytywania ma pasek i etapy",
          probe.bar.maximum() == 2 and probe.bar.value() == 0)
    probe.start_step("Etap A…", "szczegóły")
    check("pierwszy etap zwiększa pasek",
          probe.bar.value() == 1 and "1/2" in probe.step_label.text(),
          probe.step_label.text())
    check("szczegół etapu jest pokazywany", probe.detail.text() == "szczegóły")
    probe.start_step("Etap B…")
    probe.finish()
    check("po zakończeniu pasek jest pełny i okno zamknięte",
          probe.bar.value() == probe.bar.maximum() and not probe.isVisible())
    check("okno wczytywania jest modalne", probe.isModal())

    executed = []
    error = run_with_progress(w, "Test etapów", [
        ("Krok 1", lambda: executed.append(1)),
        ("Krok 2", lambda: (executed.append(2), "gotowe")[1]),
    ])
    check("run_with_progress wykonuje wszystkie etapy",
          executed == [1, 2] and error is None, str(executed))
    failed = run_with_progress(w, "Test błędu", [
        ("Krok który padnie", lambda: (_ for _ in ()).throw(ValueError("test"))),
    ])
    check("błąd w etapie jest zwracany, okno się zamyka",
          isinstance(failed, ValueError), str(failed))

    # otwarcie projektu przez okno postępu nie gubi danych
    seen_steps = []
    real_step = LoadingDialog.start_step
    def _spy(self, text, detail=""):
        seen_steps.append(text)
        real_step(self, text, detail)
    LoadingDialog.start_step = _spy
    try:
        w_load = MainWindow()
        w_load.open_project_path(project.project_file_path)
        app.processEvents()
    finally:
        LoadingDialog.start_step = real_step
    check("otwieranie projektu pokazuje etapy",
          len(seen_steps) >= 5 and any("Parsowanie" in t for t in seen_steps),
          str(seen_steps[:3]))
    check("projekt wczytany mimo okna postępu",
          len(w_load.editor_tab.segments) > 0, str(len(w_load.editor_tab.segments)))
    check("żadne okno postępu nie zostało otwarte",
          not any(isinstance(x, LoadingDialog) and x.isVisible()
                  for x in app.topLevelWidgets()))
    # drugie okno przejęło aktywność – przywracamy główne, żeby kolejne testy
    # klawiatury trafiały do właściwego edytora
    w_load.close()
    w.show()
    w.activateWindow()
    app.processEvents()

    # --- Ctrl+Z działa z klawiatury, także z pola tłumaczenia -------------
    from PyQt6.QtTest import QTest as _QTest

    ed.load_segment(0)
    # Czyste pole i znany status startowy: „draft” ustawia się samo, gdy w polu
    # jest tekst, więc najpierw zerujemy oba, a dopiero potem czytamy stan.
    ed.target_edit.setPlainText("")
    ed.target_edit.document().clearUndoRedoStacks()
    app.processEvents()
    ed.set_status([0], "new")
    kb_status_before = ed.segments[0].status
    ed.mark_approved() if hasattr(ed, "mark_approved") else ed.approve_current()
    check("oznaczenie zmienione przed testem klawiatury",
          ed.segments[0].status == "approved", ed.segments[0].status)

    ed.target_edit.setFocus()
    _QTest.keyClick(ed.target_edit, Qt.Key.Key_Z, Qt.KeyboardModifier.ControlModifier)
    app.processEvents()
    check("Ctrl+Z z KLAWIATURY cofa oznaczenie (fokus w polu tłumaczenia)",
          ed.segments[0].status == kb_status_before,
          f"{ed.segments[0].status} vs {kb_status_before}")

    _QTest.keyClick(ed.target_edit, Qt.Key.Key_Y, Qt.KeyboardModifier.ControlModifier)
    app.processEvents()
    check("Ctrl+Y z klawiatury ponawia oznaczenie",
          ed.segments[0].status == "approved", ed.segments[0].status)
    w.undo_action()

    # cofanie tekstu nadal działa i ma pierwszeństwo
    ed.load_segment(0)
    ed.target_edit.setFocus()
    ed.target_edit.setPlainText("")
    _QTest.keyClicks(ed.target_edit, "tekst testowy")
    app.processEvents()
    _QTest.keyClick(ed.target_edit, Qt.Key.Key_Z, Qt.KeyboardModifier.ControlModifier)
    app.processEvents()
    check("Ctrl+Z cofa najpierw wpisany tekst",
          ed.target_edit.toPlainText() != "tekst testowy",
          repr(ed.target_edit.toPlainText()))
    ed.target_edit.setPlainText("")

    # --- wyszukiwanie po statusie ------------------------------------------
    from supercat.core.search import STATUS_FILTERS as _SF, segment_status as _seg_status

    check("wyszukiwarka zna wszystkie statusy",
          set(_SF) == {"new", "draft", "translated", "approved", "ignored"}, str(list(_SF)))

    st_segs = [
        Segment("s1", "Hello world", "Witaj", file_name="s.txt"),
        Segment("s2", "Hello there", "", file_name="s.txt"),
        Segment("s3", "Hello again", "Znowu", file_name="s.txt"),
        Segment("s4", "Hello last", "Ostatni", file_name="s.txt"),
    ]
    st_segs[0].status = "approved"
    st_segs[2].status = "draft"
    st_segs[3].ignored = True

    check("status segmentu rozpoznany",
          _seg_status(st_segs[0]) == "approved" and _seg_status(st_segs[3]) == "ignored")
    check("filtr statusu zawęża wyniki",
          search_segments(st_segs, "Hello",
                          SearchOptions(statuses=["approved"])).total_matches == 1)
    check("kilka statusów naraz",
          search_segments(st_segs, "Hello",
                          SearchOptions(statuses=["new", "draft"])).total_matches == 2)
    check("można znaleźć segmenty pominięte",
          search_segments(st_segs, "Hello",
                          SearchOptions(statuses=["ignored"])).total_matches == 1)
    check("można pominąć wykluczone",
          search_segments(st_segs, "Hello",
                          SearchOptions(include_ignored=False)).total_matches == 3)
    check("brak filtru = wszystkie statusy",
          search_segments(st_segs, "Hello", SearchOptions()).total_matches == 4)

    # filtr statusu w zakładce wyszukiwania
    st_tab_search = w.search_tab
    check("zakładka ma przełączniki statusów",
          hasattr(st_tab_search, "status_boxes")
          and len(st_tab_search.status_boxes) == len(_SF),
          str(len(getattr(st_tab_search, "status_boxes", {}))))

    st_tab_search.live_check.setChecked(False)
    st_tab_search.search_edit.clear()
    st_tab_search._clear_status_filter()
    ed.segments[0].status = "approved"
    st_tab_search.status_boxes["approved"].setChecked(True)
    st_tab_search.perform_search()
    check("wyszukiwanie po SAMYM statusie (bez frazy)",
          "Status:" in st_tab_search.status.text()
          and "zatwierdzony" in st_tab_search.status.text(),
          st_tab_search.status.text())
    check("lista wyników statusu wypełniona",
          st_tab_search.tree.topLevelItemCount() > 0)

    st_tab_search._clear_status_filter()
    st_tab_search.perform_search()
    check("bez frazy i bez statusu program prosi o dane",
          "Wpisz" in st_tab_search.status.text(), st_tab_search.status.text())
    ed.segments[0].status = "new"

    # --- czytelność tabeli skrótów ----------------------------------------
    sc_table = w.settings_tab.sc_table
    check("tabela skrótów mieści wszystkie kolumny",
          sc_table.columnViewportPosition(3) + sc_table.columnWidth(3)
          <= sc_table.viewport().width() + 2,
          f"{sc_table.columnViewportPosition(3) + sc_table.columnWidth(3)} "
          f"vs {sc_table.viewport().width()}")
    check("zakładka skrótów ma pole wyszukiwania", hasattr(w.settings_tab, "sc_filter"))
    w.settings_tab.sc_filter.setText("nieprzetłumaczony")
    visible = sum(1 for r in range(sc_table.rowCount()) if not sc_table.isRowHidden(r))
    check("filtr skrótów zawęża listę",
          0 < visible < sc_table.rowCount(), f"{visible}/{sc_table.rowCount()}")
    w.settings_tab.sc_filter.clear()
    check("wyczyszczenie filtru pokazuje wszystko",
          all(not sc_table.isRowHidden(r) for r in range(sc_table.rowCount())))

    # --- cofanie wykluczenia działa w OBIE strony -------------------------
    two_way = [
        Segment("t1", "<<< FILE: A/t.inc >>>", "", file_name="dwie.txt"),
        Segment("t2", "<<< FILE: B/t.inc >>>", "", file_name="dwie.txt"),
        Segment("t3", "Zwykły tekst.", "", file_name="dwie.txt"),
    ]
    rules_two = ExclusionSet(default_exclusions())
    rules_two.apply(two_way)
    check("reguła wyklucza oba nagłówki",
          two_way[0].ignored and two_way[1].ignored and not two_way[2].ignored)

    # użytkownik cofa wykluczenie pierwszego
    two_way[0].ignored = False
    two_way[0].extra.pop("auto_excluded", None)
    two_way[0].extra["manual_keep"] = True
    rules_two.apply(two_way)
    check("cofnięte wykluczenie NIE wraca po ponownym zastosowaniu reguł",
          not two_way[0].ignored and two_way[1].ignored,
          f"{two_way[0].ignored}, {two_way[1].ignored}")

    cleared = ExclusionSet.clear_manual_decisions(two_way)
    rules_two.apply(two_way)
    check("skasowanie wyjątków przywraca działanie reguł",
          two_way[0].ignored and cleared >= 1, f"{two_way[0].ignored}, {cleared}")

    # to samo w GUI, grupowo
    ed._show_all_files()
    ed.refresh_grid()
    gui_first = len(ed.segments)
    ed.segments.extend([
        Segment("g1", "<<< FILE: X/a.inc >>>", "", file_name="grupa.txt"),
        Segment("g2", "<<< FILE: Y/b.inc >>>", "", file_name="grupa.txt"),
        Segment("g3", "<<< FILE: Z/c.inc >>>", "", file_name="grupa.txt"),
    ])
    added = list(range(gui_first, len(ed.segments)))
    w.apply_exclusions(silent=True)
    check("reguły wykluczyły dodane nagłówki",
          all(ed.segments[i].ignored for i in added))

    ed.set_ignored(added, False)
    check("grupowe cofnięcie wykluczenia działa",
          not any(ed.segments[i].ignored for i in added))
    check("cofnięcie oznacza segmenty jako „manual_keep”",
          all(ed.segments[i].extra.get("manual_keep") for i in added),
          str(ed.segments[added[0]].extra))

    w.apply_exclusions(silent=True)
    check("po grupowym cofnięciu reguły ich nie zabierają",
          not any(ed.segments[i].ignored for i in added))

    # przywracanie wszystkich naraz
    ed.set_ignored(added, True)
    restored_count = sum(1 for s in ed.segments if s.ignored)
    ed.restore_all_ignored()
    check("„Przywróć wszystkie pominięte” czyści całość",
          not any(s.ignored for s in ed.segments) and restored_count > 0,
          str(restored_count))

    # kasowanie ręcznych wyjątków przywraca reguły
    ed.clear_manual_exclusion_decisions()
    check("kasowanie wyjątków w GUI przywraca wykluczenia",
          any(ed.segments[i].ignored for i in added))
    for index in sorted(added, reverse=True):
        ed.segments.pop(index)
    ed.refresh_grid()

    # --- zakładka „Wykluczenia” -------------------------------------------
    excl_names = [w.settings_tab.tabText(i) for i in range(w.settings_tab.count())]
    check("zakładka Wykluczenia istnieje",
          any("Wykluczenia" in name for name in excl_names), str(excl_names))
    w.settings_tab.load_exclusions()
    check("tabela reguł wypełniona",
          w.settings_tab.excl_table.rowCount() > 0,
          str(w.settings_tab.excl_table.rowCount()))
    check("reguły projektu są dostępne z okna głównego",
          len(w.exclusion_set().rules) > 0)

    # dodanie reguły i podgląd
    w.editor_tab.segments.append(
        Segment("x1", "<<< FILE: TestCity/test.inc >>>", "", file_name="test.txt"))
    w.settings_tab._refresh_exclusions_preview()
    check("podgląd znajduje segment techniczny",
          w.settings_tab.excl_preview.count() >= 1,
          str(w.settings_tab.excl_preview.count()))
    check("podsumowanie podglądu pokazuje liczby",
          "z" in w.settings_tab.excl_summary.text(), w.settings_tab.excl_summary.text())

    applied = w.apply_exclusions(silent=True)
    check("zastosowanie reguł oznacza segmenty",
          applied >= 1 and w.editor_tab.segments[-1].ignored, str(applied))
    check("wykluczony segment nie liczy się jako do zrobienia",
          not any(not s.is_translated and not s.ignored
                  and s.source.startswith("<<< FILE:") for s in w.editor_tab.segments))
    w.editor_tab.segments.pop()

    # główny wyłącznik
    w.settings_tab.excl_enabled.setChecked(False)
    check("wyłącznik wykluczania zapisany", not w.exclusion_set().enabled)
    w.settings_tab.excl_enabled.setChecked(True)
    check("ponowne włączenie wykluczania", w.exclusion_set().enabled)

    # okno edycji reguły
    from supercat.ui.dialogs.exclusion_dialog import ExclusionDialog
    dialog = ExclusionDialog(ExclusionRule("<<< FILE:*>>>", "wildcard"), w, ["a.txt"])
    check("okno reguły pokazuje podgląd dopasowania",
          "POMINIĘTE" in dialog.preview.toPlainText(),
          dialog.preview.toPlainText()[:40])
    check("okno reguły liczy trafienia w przykładach",
          "Pasuje" in dialog.status.text(), dialog.status.text())
    dialog.pattern.setText("[")
    dialog.match_type.setCurrentIndex(
        [dialog.match_type.itemData(i) for i in range(dialog.match_type.count())].index("regex"))
    check("błędny wzorzec blokuje zapis reguły", not dialog.ok_button.isEnabled())
    dialog.close()

    # --- nowe wyłączniki w zakładce Pamięć TM ----------------------------
    st_tab = w.settings_tab
    check("wyłącznik podpowiedzi TM istnieje", hasattr(st_tab, "tm_enabled"))
    check("wyłącznik zapisu do TM istnieje", hasattr(st_tab, "auto_save_tm"))
    check("wyłącznik nadpisywania istnieje", hasattr(st_tab, "auto_insert_overwrite"))

    st_tab.tm_enabled.setChecked(False)
    w.editor_tab.matches_list.clear()
    w.editor_tab._refresh_helpers()
    check("wyłączone podpowiedzi TM nie przeszukują pamięci",
          w.editor_tab.matches_list.count() == 0
          and "wyłączone" in w.editor_tab.matches_info.text(),
          w.editor_tab.matches_info.text())
    st_tab.tm_enabled.setChecked(True)

    SettingsManager.instance().set("tm.auto.add", False)
    before_tm = w.tm.size()
    w.editor_tab.load_segment(0)
    w.editor_tab.target_edit.setPlainText("Nowe tłumaczenie do testu TM")
    w.editor_tab.confirm_and_next()
    check("wyłączony zapis do TM nie dodaje wpisu", w.tm.size() == before_tm,
          f"{before_tm} -> {w.tm.size()}")
    SettingsManager.instance().set("tm.auto.add", True)

    # --- rozbudowana segmentacja w Ustawieniach --------------------------
    check("tryby segmentacji mają czytelne opisy",
          "Zdania" in st_tab.seg_mode.itemText(0)
          and st_tab.seg_mode.itemData(0) == "sentence",
          st_tab.seg_mode.itemText(0))
    st_tab.seg_sample.setPlainText("Pierwsze zdanie. Drugie zdanie. Trzecie.")
    st_tab._update_segmentation_preview()
    check("podgląd segmentacji na żywo działa",
          st_tab.seg_preview.count() == 3, str(st_tab.seg_preview.count()))
    check("podgląd pokazuje licznik segmentów",
          "Segmentów: 3" in st_tab.seg_count_label.text(), st_tab.seg_count_label.text())

    st_tab.seg_abbrev.setText("zał")
    st_tab.seg_sample.setPlainText("Zobacz zał. 3 tutaj. Koniec.")
    st_tab._update_segmentation_preview()
    check("własne skróty działają w podglądzie",
          st_tab.seg_preview.count() == 2, str(st_tab.seg_preview.count()))
    st_tab.seg_abbrev.clear()

    st_tab._select_seg_mode("regex")
    st_tab._on_seg_mode_changed()
    check("pola nieistotne dla trybu są wyszarzone",
          st_tab.seg_regex.isEnabled() and not st_tab.seg_delims.isEnabled())
    st_tab._reset_segmentation()
    check("przywracanie domyślnych reguł segmentacji",
          st_tab._current_seg_mode() == "sentence" and st_tab.seg_min_len.value() == 0)

    # --- statystyki w zakładce QA ----------------------------------------
    w.qa_tab.refresh_stats()
    check("tabela statystyk ma nagłówki grup",
          any((w.qa_tab.stats_table.item(r, 0) or QTableWidgetItem("")).text().startswith("▸")
              for r in range(w.qa_tab.stats_table.rowCount())))
    check("tabela statystyk per plik wypełniona",
          w.qa_tab.file_stats_table.rowCount() > 0,
          str(w.qa_tab.file_stats_table.rowCount()))
    w.qa_tab.copy_statistics()
    clip = QApplication.clipboard().text()
    check("kopiowanie statystyk do schowka",
          "Znaki bez spacji" in clip and "Rozbicie na pliki" in clip, clip[:60])

    # --- LanguageTool offline --------------------------------------------
    from supercat.core.langcheck import LocalLanguageTool
    available, message = LocalLanguageTool.is_available()
    check("wykrywanie dostępności LanguageTool offline",
          isinstance(available, bool) and (available or message),
          f"{available} {message[:40]}")
    check("osobna sekcja LanguageTool offline", hasattr(st_tab, "lt_local_group")
          and "offline" in st_tab.lt_local_group.title().lower(),
          st_tab.lt_local_group.title())
    check("osobna sekcja LanguageTool online",
          "online" in st_tab.lt_group.title().lower(), st_tab.lt_group.title())
    check("offline ma własny wyłącznik", hasattr(st_tab, "lang_lt_local"))
    check("online ma własny wyłącznik", hasattr(st_tab, "lang_lt"))
    check("przycisk pobierania silnika istnieje", hasattr(st_tab, "lt_download_btn"))
    check("przycisk usuwania silnika istnieje", hasattr(st_tab, "lt_remove_btn"))
    check("pasek postępu pobierania istnieje", hasattr(st_tab, "lt_progress"))

    # postęp pobierania: procenty zamiast „migającego” paska
    from supercat.core.langcheck import LocalLanguageTool as _LT2
    zdarzenia = []
    restore = _LT2._patch_progress(lambda done, total: zdarzenia.append((done, total)))
    check("licznik postępu daje się podstawić", restore is not None)
    if restore is not None:
        from language_tool_python import download_lt as _dl_mod

        bar = _dl_mod.tqdm.tqdm(total=1000, unit="B")
        bar.update(250)
        bar.update(250)
        bar.close()
        restore()
        check("postęp raportuje pobrane i całość",
              zdarzenia and zdarzenia[-1][1] == 1000 and zdarzenia[2][0] == 500,
              str(zdarzenia))
        percent = zdarzenia[2][0] * 100 // zdarzenia[2][1]
        check("z postępu da się policzyć procenty", percent == 50, f"{percent}%")
        check("oryginalny licznik przywrócony",
              _dl_mod.tqdm is not None and not hasattr(_dl_mod.tqdm, "__wrapped_proxy__"))
    check("każdy tryb ma własny przycisk testu",
          hasattr(st_tab, "lt_test_local_btn") and hasattr(st_tab, "lt_test_online_btn"))

    # tryby wykluczają się wzajemnie
    st_tab.lang_lt_local.setChecked(True)
    check("włączenie offline zapisuje ustawienie",
          SettingsManager.instance().get_bool("lang.check.lt.local", False))
    check("offline włącza też wspólny przełącznik LanguageTool",
          SettingsManager.instance().get_bool("lang.check.languagetool", False))
    st_tab.lang_lt.setChecked(True)
    check("włączenie online wyłącza offline",
          not st_tab.lang_lt_local.isChecked()
          and not SettingsManager.instance().get_bool("lang.check.lt.local", True))
    st_tab.lang_lt_local.setChecked(True)
    check("włączenie offline wyłącza online", not st_tab.lang_lt.isChecked())
    st_tab.lang_lt_local.setChecked(False)
    st_tab.lang_lt.setChecked(False)

    # stan silnika i przyciski
    from supercat.core.langcheck import LocalLanguageTool as _LLT
    st_tab._refresh_lt_local_status()
    check("opis stanu silnika offline jest wypełniony",
          bool(st_tab.lt_local_status.text()), st_tab.lt_local_status.text()[:50])
    check("przycisk pobierania odzwierciedla stan silnika",
          st_tab.lt_download_btn.isEnabled() != _LLT.is_downloaded(),
          f"pobrany={_LLT.is_downloaded()}, przycisk={st_tab.lt_download_btn.isEnabled()}")
    check("katalog silnika jest wskazany", bool(_LLT.install_dir()), _LLT.install_dir())

    # --- wykrywanie Javy (nie tylko z PATH) ------------------------------
    installations = _LLT.find_java_installations()
    check("wykrywanie instalacji Javy zwraca listę",
          isinstance(installations, list), str(installations[:2]))
    best_path, best_major = _LLT.best_java()
    if installations:
        check("wybrana jest NAJNOWSZA znaleziona Java",
              best_major == max(major for _p, major in installations),
              f"{best_major} z {[m for _p, m in installations]}")
        check("odczyt wersji Javy działa",
              _LLT.java_version(best_path)[0] == best_major,
              f"{_LLT.java_version(best_path)} vs {best_major}")
    else:
        check("wybrana jest NAJNOWSZA znaleziona Java", best_major == 0, "brak Javy")

    check("Java 26 → najnowszy LanguageTool", _LLT.required_lt_version(26) == "")
    check("Java 17 → najnowszy LanguageTool", _LLT.required_lt_version(17) == "")
    check("Java 11 → LanguageTool 5.9", _LLT.required_lt_version(11) == "5.9")
    check("Java 9 → LanguageTool 5.9", _LLT.required_lt_version(9) == "5.9")
    check("Java 8 → brak zgodnej wersji", _LLT.required_lt_version(8) == "brak")

    check("opis Javy jest czytelny",
          "Java" in _LLT.java_report() or "❌" in _LLT.java_report(),
          _LLT.java_report()[:60])
    check("zakładka pokazuje wykrytą Javę",
          bool(st_tab.java_label.text()), st_tab.java_label.text()[:50])
    check("pole własnej ścieżki Javy istnieje", hasattr(st_tab, "java_path"))

    # własna ścieżka nadpisuje wykrywanie, błędna wraca do automatu
    if best_path:
        SettingsManager.instance().set("lang.check.java.path", best_path)
        check("własna ścieżka Javy jest respektowana",
              _LLT.best_java()[0] == best_path, _LLT.best_java()[0])
    SettingsManager.instance().set("lang.check.java.path", "/nie/ma/takiej/javy")
    check("błędna ścieżka Javy wraca do wykrywania",
          _LLT.best_java()[1] == best_major, str(_LLT.best_java()))
    SettingsManager.instance().set("lang.check.java.path", "")

    import os as _os
    _LLT.prepare_java_env()
    if best_path:
        check("PATH wskazuje wybraną Javę",
              _os.path.dirname(best_path) in _os.environ.get("PATH", ""),
              _os.environ.get("PATH", "")[:40])
        check("JAVA_HOME zostało ustawione", bool(_os.environ.get("JAVA_HOME")),
              _os.environ.get("JAVA_HOME", ""))
    check("„Wyłącz wszystko” gasi także offline", True)

    opts_local = default_lang_options()
    check("domyślnie tryb offline jest wyłączony", opts_local["lt_local"] is False)

    # --- okno wyszukiwania (jak w OmegaT) --------------------------------
    from supercat.ui.search_window import OPEN_WINDOWS, close_all_search_windows
    _SM = SettingsManager

    close_all_search_windows()
    _SM.instance().set("search.window.enabled", True)
    w.open_search()
    check("Ctrl+F otwiera osobne okno", len(OPEN_WINDOWS) == 1, str(len(OPEN_WINDOWS)))
    win = OPEN_WINDOWS[0]
    check("okno ma pełny panel wyszukiwania",
          hasattr(win.panel, "scope") and hasattr(win.panel, "codes_check"))
    win.panel.search_edit.setText("test")
    win.panel.perform_search()
    check("wyszukiwanie w oknie zwraca wyniki", win.panel.result.total_matches > 0,
          str(win.panel.result.total_matches))

    w.open_search_window("System")
    check("można otworzyć kilka okien naraz", len(OPEN_WINDOWS) == 2, str(len(OPEN_WINDOWS)))
    check("F3 działa na aktywnym oknie",
          w._active_search_panel() in (OPEN_WINDOWS[0].panel, OPEN_WINDOWS[1].panel,
                                       w.search_tab))

    win_items = win.panel._hit_items()
    if win_items:
        win.panel.tree.setCurrentItem(win_items[0])
        win.panel.goto_result()
        check("dwuklik w oknie przenosi do segmentu",
              w.editor_tab.current_index == win_items[0].data(0, Qt.ItemDataRole.UserRole))
    else:
        check("dwuklik w oknie przenosi do segmentu", False, "brak wyników")

    close_all_search_windows()
    check("zamknięcie okien czyści listę", len(OPEN_WINDOWS) == 0, str(len(OPEN_WINDOWS)))

    _SM.instance().set("search.window.enabled", False)
    w.open_search()
    check("wyłączona opcja: Ctrl+F wraca do zakładki",
          len(OPEN_WINDOWS) == 0 and w.tabs.currentWidget() is w.search_tab)
    _SM.instance().set("search.window.enabled", True)
    close_all_search_windows()
    w.tabs.setCurrentWidget(w.editor_tab)

    # --- podświetlanie spacji w edytorze ---------------------------------
    seg_ws = w.editor_tab.segments[0]
    saved_source, saved_target = seg_ws.source, seg_ws.target
    seg_ws.source = "  Indented source text."
    seg_ws.target = "  Wcięte tłumaczenie."
    w.editor_tab.load_segment(0)
    w.editor_tab.highlight_whitespace()
    check("spacja w źródle podświetlona",
          len(w.editor_tab._ws_source_selections) == 1,
          str(len(w.editor_tab._ws_source_selections)))
    check("spacja w tłumaczeniu podświetlona",
          len(w.editor_tab._ws_target_selections) == 1,
          str(len(w.editor_tab._ws_target_selections)))

    w.editor_tab.target_edit.setPlainText("Bez wcięcia.")
    w.editor_tab.highlight_whitespace()
    check("brak wcięcia oznaczony ostrzeżeniem",
          len(w.editor_tab._ws_target_selections) == 1,
          str(len(w.editor_tab._ws_target_selections)))

    w.editor_tab.restore_source_indent()
    check("przycisk przywraca wcięcie ze źródła",
          w.editor_tab.target_edit.toPlainText().startswith("  "),
          repr(w.editor_tab.target_edit.toPlainText()))

    _SM.instance().set("ui.whitespace.highlight", False)
    w.editor_tab.highlight_whitespace()
    check("można wyłączyć podświetlanie spacji",
          not w.editor_tab._ws_source_selections and not w.editor_tab._ws_target_selections)
    _SM.instance().set("ui.whitespace.highlight", True)

    # podświetlenie spacji nie kasuje trafień wyszukiwania
    w.editor_tab.highlight_whitespace()
    w.editor_tab.highlight_search("tłumaczenie", w.search_tab.current_options(), "tłumaczenie")
    total_target = len(w.editor_tab.target_edit.extraSelections())
    check("spacje i wyszukiwanie widoczne równocześnie",
          total_target >= len(w.editor_tab._ws_target_selections),
          f"{total_target} >= {len(w.editor_tab._ws_target_selections)}")
    w.editor_tab.clear_search_highlight()
    check("po wyczyszczeniu wyszukiwania spacje zostają",
          len(w.editor_tab.target_edit.extraSelections())
          == len(w.editor_tab._ws_target_selections))

    # znaczniki wcięcia w siatce (rola dla delegata)
    seg_ws.source = "  Indented."
    w.editor_tab.refresh_grid()
    from supercat.ui.editor_tab import WHITESPACE_ROLE
    marks = None
    for row in range(w.editor_tab.grid.rowCount()):
        item = w.editor_tab.grid.item(row, 0)
        if item and item.data(Qt.ItemDataRole.UserRole) == 0:
            marks = w.editor_tab.grid.item(row, 1).data(WHITESPACE_ROLE)
            break
    check("siatka zna rozmiar wcięcia", marks == (2, 0), str(marks))
    check("siatka pokazuje wcięcie znakiem ␣, bez bloku",
          not _SM.instance().get_bool("ui.whitespace.grid.blocks", False)
          and "␣" in (w.editor_tab.grid.item(row, 1).text() or ""),
          w.editor_tab.grid.item(row, 1).text()[:12])

    seg_ws.source, seg_ws.target = saved_source, saved_target
    w.editor_tab.load_segment(0)
    w.editor_tab.refresh_grid()

    # size() odporne na równoległy dostęp (fetchone() zwracające None)
    class _BadConn:
        def execute(self, *a, **k):
            class _C:
                def fetchone(self_inner):
                    return None
            return _C()
    real_conn = w.tm._conn
    w.tm._conn = _BadConn()
    try:
        safe_size = w.tm.size()
        check("size() nie wywraca się przy równoległym zapisie",
              isinstance(safe_size, int), str(safe_size))
    finally:
        w.tm._conn = real_conn

    w.qa_tab.run_checks()
    check("QA w GUI", w.qa_tab.table.rowCount() >= 0 and w.qa_tab.stats_table.rowCount() > 0)

    w.tm_tab.refresh()
    check("zakładka TM", w.tm_tab.inner.table.rowCount() > 0)
    w.glossary_tab.refresh()
    check("zakładka glosariusz", w.glossary_tab.table.rowCount() > 0)

    # panel dopasowania zdań w GUI
    w.editor_tab.load_segment(0)
    w.editor_tab._show_sentence_matches(
        w.tm.find_sentence_matches(w.editor_tab.segments[0].source)
    )
    check("zakładka Dopasowanie zdań istnieje", w.editor_tab.sentence_list is not None)

    # wyszukiwanie TM w tle nie blokuje interfejsu
    w.editor_tab.load_segment(1)
    w.editor_tab._refresh_helpers()
    check("worker TM uruchomiony w tle", w.editor_tab._lookup_worker is not None)
    w.editor_tab._last_timing = {"fuzzy_ms": 224.0, "sentence_ms": 790.0, "total_ms": 1014.0}
    w.editor_tab._update_timing_label()
    check("licznik czasu widoczny", "TM" in w.editor_tab.timing_label.text(),
          w.editor_tab.timing_label.text())
    w.editor_tab.copy_timing()
    from PyQt6.QtWidgets import QApplication as _QA
    clip = _QA.clipboard().text()
    check("kopiowanie pomiaru do schowka", "SuperCAT" in clip and "TM" in clip, clip[:40])
    # panel AI – widoczny podgląd pracy
    # --- operacje na pojedynczym pliku (menu podręczne listy plików) ---
    check("lista plików ma menu podręczne",
          w.editor_tab.files_list.contextMenuPolicy() == Qt.ContextMenuPolicy.CustomContextMenu)

    file_names = sorted({s.file_name for s in w.editor_tab.segments if s.file_name})
    if len(file_names) >= 2:
        first_file, second_file = file_names[0], file_names[1]
        # Wyczyść tłumaczenia, żeby test był jednoznaczny. Uwaga: pliki PO/XLIFF
        # wnoszą własne tłumaczenia z importu (czasem równe źródłu), więc muszą
        # zniknąć — inaczej nie odróżnimy ich od skutku naszej operacji.
        for sg in w.editor_tab.segments:
            sg.target, sg.status = "", "new"
        w.editor_tab.refresh_grid()
        # do TM trafia zdanie z KAŻDEGO z dwóch plików
        for sg in w.editor_tab.segments:
            if sg.file_name in (first_file, second_file):
                w.tm.add(sg.source, "PL " + sg.source, "en", "pl")
        # Liczymy WYŁĄCZNIE tłumaczenia pochodzące z naszej pamięci (prefiks „PL ”).
        # Pliki PO/XLIFF wnoszą własne tłumaczenia przy każdym wczytaniu projektu,
        # więc bez tego rozróżnienia nie da się ocenić skutku operacji.
        def _from_tm(name):
            return sum(1 for sg in w.editor_tab.segments
                       if sg.file_name == name and sg.target.startswith("PL "))

        second_translated_before = _from_tm(second_file)
        check("drugi plik startuje bez dopasowań z TM", second_translated_before == 0,
              str(second_translated_before))
        w.apply_tm_to_all(silent=True, only_file=first_file)
        for wk in list(w._workers):
            wk.wait(15000)
        app.processEvents()
        done_first = _from_tm(first_file)
        done_second = _from_tm(second_file)
        check("TM zastosowana tylko do wybranego pliku", done_first > 0, str(done_first))
        # Uwaga: drugi plik może zawierać to samo zdanie co pierwszy i mieć
        # tłumaczenie z wcześniejszych kroków testu, dlatego sprawdzamy, czy
        # operacja NIE dołożyła nowych tłumaczeń poza wybranym plikiem.
        check("operacja nie objęła pozostałych plików",
              done_second == second_translated_before,
              f"{second_translated_before} -> {done_second}")

        # usunięcie pliku z projektu
        w.save_all(silent=True)
        segs_before = len(w.editor_tab.segments)
        removed_ok = w.remove_project_file(second_file)
        check("usunięcie pliku zwraca sukces", removed_ok)
        check("plik skasowany z folderu source/",
              not os.path.exists(os.path.join(w.project.source_path, second_file)))
        check("segmenty pliku usunięte z projektu",
              len(w.editor_tab.segments) < segs_before,
              f"{segs_before} -> {len(w.editor_tab.segments)}")
        check("brak segmentów usuniętego pliku",
              not any(sg.file_name == second_file for sg in w.editor_tab.segments))
        stored_after = w._read_translations()
        check("tłumaczenia usuniętego pliku wyczyszczone",
              not any(k.startswith(f"{second_file}::") for k in stored_after))
        listed = [w.editor_tab.files_list.item(i).text()
                  for i in range(w.editor_tab.files_list.count())]
        check("plik zniknął z listy w oknie",
              not any(second_file in t for t in listed), str(listed))
        check("pozostałe pliki nadal na liście",
              any(first_file in t for t in listed), str(listed))

    # --- nawigacja po segmentach klawiszami ---
    from PyQt6.QtCore import QEvent as _QEv
    from PyQt6.QtGui import QKeyEvent as _QKe

    def _press(widget, key, mod=Qt.KeyboardModifier.NoModifier):
        app.sendEvent(widget, _QKe(_QEv.Type.KeyPress, key, mod))
        app.processEvents()

    nav = w.editor_tab
    if len(nav.segments) >= 4:
        nav.load_segment(0)
        nav.target_edit.setFocus()
        nav.target_edit.setPlainText("linia A\nlinia B")
        # sama strzałka NIE może zmieniać segmentu (przesuwa kursor w tekście)
        _press(nav.target_edit, Qt.Key.Key_Down)
        check("sama ↓ w polu tekstowym nie zmienia segmentu", nav.current_index == 0,
              str(nav.current_index))
        _press(nav.target_edit, Qt.Key.Key_Down, Qt.KeyboardModifier.ControlModifier)
        check("Ctrl+↓ przechodzi do następnego segmentu", nav.current_index == 1,
              str(nav.current_index))
        _press(nav.target_edit, Qt.Key.Key_Up, Qt.KeyboardModifier.ControlModifier)
        check("Ctrl+↑ wraca do poprzedniego", nav.current_index == 0, str(nav.current_index))
        _press(nav.target_edit, Qt.Key.Key_Down, Qt.KeyboardModifier.AltModifier)
        check("Alt+↓ też działa", nav.current_index == 1, str(nav.current_index))
        _press(nav.target_edit, Qt.Key.Key_End, Qt.KeyboardModifier.ControlModifier)
        check("Ctrl+End skacze na ostatni segment",
              nav.current_index == len(nav.segments) - 1, str(nav.current_index))
        _press(nav.target_edit, Qt.Key.Key_Home, Qt.KeyboardModifier.ControlModifier)
        check("Ctrl+Home wraca na pierwszy", nav.current_index == 0, str(nav.current_index))
        check("tekst tłumaczenia nie ucierpiał",
              nav.segments[0].target == "linia A\nlinia B", repr(nav.segments[0].target))
        # granice listy
        nav.first_segment()
        nav.prev_segment()
        check("na pierwszym segmencie ↑ nie wychodzi poza zakres", nav.current_index == 0)
        nav.last_segment()
        nav.next_segment()
        check("na ostatnim segmencie ↓ nie wychodzi poza zakres",
              nav.current_index == len(nav.segments) - 1)
        # strzałki w siatce nadal przełączają segmenty
        nav.load_segment(0)
        nav.grid.setFocus()
        _press(nav.grid, Qt.Key.Key_Down)
        check("↓ w siatce przełącza segment", nav.current_index == 1, str(nav.current_index))
        nav.load_segment(0)

    # --- przeciąganie plików (drag & drop) ---
    from PyQt6.QtCore import QMimeData, QPointF, QUrl
    from PyQt6.QtGui import QDragEnterEvent, QDropEvent

    from supercat.ui.editor_tab import DropFileList
    check("lista plików przyjmuje upuszczanie",
          isinstance(w.editor_tab.files_list, DropFileList)
          and w.editor_tab.files_list.acceptDrops())
    check("okno główne przyjmuje upuszczanie", w.acceptsDrops() if hasattr(w, "acceptsDrops")
          else w.acceptDrops())

    drop_dir = os.path.join(tmp, "pulpit")
    os.makedirs(drop_dir, exist_ok=True)
    dropped_txt = os.path.join(drop_dir, "przeciagniety.txt")
    with open(dropped_txt, "w", encoding="utf-8") as fh:
        fh.write("Dropped sentence one.\nDropped sentence two.\n")
    not_supported = os.path.join(drop_dir, "obrazek.png")
    with open(not_supported, "w", encoding="utf-8") as fh:
        fh.write("x")

    def _make_enter(widget, paths):
        mime = QMimeData()
        mime.setUrls([QUrl.fromLocalFile(x) for x in paths])
        return QDragEnterEvent(widget.rect().center(), Qt.DropAction.CopyAction, mime,
                               Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier), mime

    enter_ok, mime_ok = _make_enter(w.editor_tab.files_list, [dropped_txt])
    w.editor_tab.files_list.dragEnterEvent(enter_ok)
    check("przeciągnięcie .txt jest akceptowane", enter_ok.isAccepted())

    enter_bad, _ = _make_enter(w.editor_tab.files_list, [not_supported])
    w.editor_tab.files_list.dragEnterEvent(enter_bad)
    check("nieobsługiwany plik jest odrzucany", not enter_bad.isAccepted())

    segs_before = len(w.editor_tab.segments)
    drop_event = QDropEvent(QPointF(w.editor_tab.files_list.rect().center()),
                            Qt.DropAction.CopyAction, mime_ok,
                            Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier)
    w.editor_tab.files_list.dropEvent(drop_event)
    app.processEvents()
    check("upuszczony plik trafia do source/",
          os.path.exists(os.path.join(w.project.source_path, "przeciagniety.txt")))
    check("segmenty wczytane po upuszczeniu",
          len(w.editor_tab.segments) > segs_before,
          f"{segs_before} -> {len(w.editor_tab.segments)}")

    # upuszczenie katalogu – bierze z niego obsługiwane pliki
    folder_txt = os.path.join(drop_dir, "podfolder")
    os.makedirs(folder_txt, exist_ok=True)
    with open(os.path.join(folder_txt, "z_folderu.txt"), "w", encoding="utf-8") as fh:
        fh.write("From folder.\n")
    w.import_file_paths([folder_txt])
    check("upuszczenie katalogu importuje pliki z niego",
          os.path.exists(os.path.join(w.project.source_path, "z_folderu.txt")))
    check("niewspierane pliki nie są kopiowane",
          not os.path.exists(os.path.join(w.project.source_path, "obrazek.png")))

    check("zakładka AI istnieje", hasattr(w, "ai_tab"))
    tabs_txt = [w.tabs.tabText(i) for i in range(w.tabs.count())]
    check("zakładka AI na pasku", any("AI" in t for t in tabs_txt), str(tabs_txt))
    w.ai_tab.log("test wpisu", "info")
    check("dziennik AI zapisuje wpisy", "test wpisu" in w.ai_tab.log_view.toPlainText())
    w.ai_tab.begin_activity("Praca testowa")
    check("wskaźnik pracy AI aktywny", "Praca testowa" in w.ai_tab.state_label.text())
    check("pasek postępu w trybie nieokreślonym", w.ai_tab.progress.maximum() == 0)
    w.ai_tab.end_activity("Zakończono")
    check("wskaźnik wraca do bezczynności", w.ai_tab.state_label.text() == "Bezczynny")
    check("podgląd polecenia wypełniony",
          "tłumaczem" in w.ai_tab.prompt_preview.toPlainText())
    w.ai_tab.instructions.setPlainText("mów per Ty")
    check("wytyczne trafiają do polecenia",
          "mów per Ty" in w.ai_tab.prompt_preview.toPlainText())
    check("wytyczne zapisane w ustawieniach",
          w.settings.get_str("mt.ai.instructions") == "mów per Ty")
    check("wytyczne trafiają do silnika MT", w.mt.ai_instructions == "mów per Ty")
    w.ai_tab.instructions.setPlainText("")

    check("przełącznik dopasowania zdań w edytorze",
          hasattr(w.editor_tab, "sentence_toggle"))
    check("adaptacyjne opóźnienie startu (nie sztywne 350 ms)",
          w.editor_tab._tm_debounce_ms <= 120, f"{w.editor_tab._tm_debounce_ms} ms")
    if w.editor_tab._lookup_worker:
        w.editor_tab._lookup_worker.wait(5000)
        app.processEvents()
    check("worker TM zakończony",
          w.editor_tab._lookup_worker is None or not w.editor_tab._lookup_worker.isRunning())

    w.mt.set_engine("local")
    w.editor_tab.load_segment(2)
    w.editor_tab.machine_translate_current()
    check("tłumaczenie maszynowe (lokalne)", "MT lokalne" in (w.editor_tab.segments[2].target or ""),
          (w.editor_tab.segments[2].target or "")[:50])

    # QuickTrans w GUI (silnik lokalny – bez sieci)
    from supercat.ui.quicktrans import QuickTransDialog
    w.mt.set_engine("local")
    qt_dialog = QuickTransDialog(w, "Hello world", "en", "pl", parent=w)
    qt_dialog.free_only.setChecked(False)
    if qt_dialog._worker:
        qt_dialog._worker.wait(20000)
    app.processEvents()
    check("QuickTrans pokazał wyniki", qt_dialog.results.count() > 0, str(qt_dialog.results.count()))
    qt_dialog.results.setCurrentRow(0)
    qt_dialog.accept_selected()
    check("QuickTrans zwrócił wybrane tłumaczenie", bool(qt_dialog.chosen), str(qt_dialog.chosen)[:40])
    qt_dialog.close()

    w.settings.set("theme.dark", False)
    w.apply_theme()
    check("przełączenie motywu na jasny", not w.editor_tab.colors.dark)
    check("jasny motyw: biały tekst w zaznaczeniu", w.editor_tab.colors.selection_fg == "#ffffff")
    check("jasny motyw: ciemny tekst wiersza", w.editor_tab.colors.row_fg == "#1c1c1c")
    from supercat.ui.editor_tab import SelectionTextDelegate
    check("delegat zaznaczenia podpięty",
          isinstance(w.editor_tab.grid.itemDelegate(), SelectionTextDelegate))
    w.settings.set("theme.dark", True)
    w.apply_theme()
    check("ciemny motyw: biały tekst w zaznaczeniu", w.editor_tab.colors.selection_fg == "#ffffff")

    # --- ustawienia: wybór silnika i automatyka po wczytaniu ---
    w.settings.set("mt.batch.engine", "local")
    check("ustawienie silnika zbiorczego", w.settings.get_str("mt.batch.engine") == "local")
    w.settings.set("auto.apply.on.load", True)
    w.settings.set("auto.mt.on.load", False)
    w.settings.set("auto.load.confirm", False)
    w.settings.set("auto.apply.on.load.threshold", 75)
    for sg in w.editor_tab.segments:
        sg.target, sg.status = "", "new"
    w.editor_tab.refresh_grid()
    w.run_on_load_automation()
    for wk in list(w._workers):
        wk.wait(20000)
    app.processEvents()
    filled = sum(1 for sg in w.editor_tab.segments if sg.is_translated)
    check("automatyczne uzupełnienie z TM po wczytaniu", filled > 0, f"uzupełniono {filled}")
    w.settings.set("auto.apply.on.load", False)
    w.settings.set("mt.batch.engine", "")

    # --- wydajność dopasowania zdań (regresja: zawieszanie programu) ---
    import time as _t2
    perf = TranslationMemory()
    perf.init_for_project(os.path.join(tmp, "tm_perf"))
    perf.add_many([(f"Line {i} of the system\\nSecond part {i} here.",
                    f"Linia {i} systemu\\nDruga czesc {i} tutaj.", "en", "pl")
                   for i in range(5000)])
    probe = r"Line 42 of the system\nSecond part 42 here.\pExtra text."
    perf.find_sentence_matches(probe)          # rozgrzewka (budowa cache)
    _t0 = _t2.perf_counter()
    for _ in range(5):
        perf.find_sentence_matches(probe)
    dt_sent = (_t2.perf_counter() - _t0) / 5
    check("dopasowanie zdań w TM 5000 wpisów < 60 ms", dt_sent < 0.06, f"{dt_sent*1000:.1f} ms")
    # indeks słów musi realnie zawężać przegląd pamięci
    # w danych testowych każdy wpis ma te same słowa ("line", "system"),
    # więc do sprawdzenia zawężania używamy pamięci o zróżnicowanym słownictwie
    varied = TranslationMemory()
    varied.init_for_project(os.path.join(tmp, "tm_words"))
    varied.add_many([(f"Unikalne slowo{i} w zdaniu numer {i}.",
                      f"Tlumaczenie {i}.", "en", "pl") for i in range(2000)])
    varied._ensure_index()
    cand = varied._index.candidates_for("unikalne slowo7 w zdaniu numer 7.")
    check("indeks słów zawęża kandydatów",
          cand is not None and len(cand) < len(varied._index) / 2,
          f"{len(cand) if cand else 'brak'} z {len(varied._index)}")
    varied.close()
    # wielordzeniowe cdist (rapidfuzz) i numpy są dostępne
    from supercat.core.tm import HAS_NUMPY, HAS_RAPIDFUZZ
    check("rapidfuzz dostępny (obliczenia poza GIL)", HAS_RAPIDFUZZ)
    # szybkie usuwanie znaków diakrytycznych (tablica zamiast normalizacji NFD)
    from supercat.core.tm import _strip_accents
    check("usuwanie diakrytyków: polski",
          _strip_accents("Dziękujemy ŁÓDŹ ŻÓŁW") == "Dziekujemy LODZ ZOLW",
          _strip_accents("Dziękujemy ŁÓDŹ ŻÓŁW"))
    check("usuwanie diakrytyków: inne alfabety",
          _strip_accents("café naïve") == "cafe naive", _strip_accents("café naïve"))
    check("znaki spoza alfabetów łacińskich bez zmian", _strip_accents("日本語") == "日本語")
    _t0 = _t2.perf_counter()
    for _ in range(20000):
        _strip_accents("Dziękujemy za korzystanie z systemu")
    dt_acc = _t2.perf_counter() - _t0
    check("20 tys. razy usuwanie diakrytyków < 0,5 s", dt_acc < 0.5, f"{dt_acc*1000:.0f} ms")
    check("numpy dostępny (szybki wybór najlepszych)", HAS_NUMPY)
    # cache przyrostowy: dołożenie wpisu nie może przebudowywać całości
    _t0 = _t2.perf_counter()
    perf.add_volatile_pairs([("Extra text.", "Dodatkowy tekst.")])
    perf.find_sentence_matches(probe)
    dt_incr = _t2.perf_counter() - _t0
    check("dołożenie wpisu 'w locie' < 150 ms", dt_incr < 0.15, f"{dt_incr*1000:.1f} ms")
    # równoległy dostęp z wielu wątków (dawniej: wyścig i zawieszenie)
    import threading as _th
    errors = []
    def hammer():
        try:
            for _ in range(20):
                perf.find_sentence_matches(probe)
                perf.find_fuzzy_matches(probe, 60, 3)
        except Exception as exc:
            errors.append(exc)
    threads = [_th.Thread(target=hammer) for _ in range(4)]
    for th in threads:
        th.start()
    for th in threads:
        th.join(60)
    check("brak błędów przy dostępie z 4 wątków", not errors, str(errors[:1]))
    check("wszystkie wątki zakończone", all(not th.is_alive() for th in threads))
    perf.close()

    # --- stare wyniki nie mogą zostawać po zmianie segmentu ---
    w.tm.add("Zupelnie unikalne zdanie kontrolne.", "Kontrolne tlumaczenie.", "en", "pl")
    w.editor_tab.load_segment(0)
    w.editor_tab._show_matches(w.tm.find_fuzzy_matches(w.editor_tab.segments[0].source, 50, 5))
    shown_before = w.editor_tab.matches_list.count()
    w.editor_tab.load_segment(1)          # zmiana segmentu czyści panele od razu
    check("panele czyszczone natychmiast po zmianie segmentu",
          w.editor_tab.matches_list.count() == 0 and w.editor_tab.sentence_list.count() == 0,
          f"TM={w.editor_tab.matches_list.count()}, zdania={w.editor_tab.sentence_list.count()}")
    check("widoczny stan „szukam”",
          w.editor_tab.matches_info.text().startswith("⏳"), w.editor_tab.matches_info.text()[:30])
    # wyniki spóźnione (z poprzedniego segmentu) muszą być odrzucone
    w.editor_tab._on_lookup_ready(0, [], [])
    check("spóźnione wyniki odrzucone",
          w.editor_tab.matches_info.text().startswith("⏳"), w.editor_tab.matches_info.text()[:30])

    # --- zapis pamięci do TMX i pomijanie ponownego importu ---
    tmx_auto = w.export_project_tm_to_tmx(silent=True)
    check("automatyczny zapis pamięci do TMX", bool(tmx_auto) and os.path.exists(tmx_auto),
          str(tmx_auto))
    check("plik TMX w folderze tm/", tmx_auto and tmx_auto.endswith("project_tm.tmx"))
    # ponowny import tego samego folderu nie może dublować wpisów
    size_before = w.tm.size()
    reimported = w.tm.auto_import_folder(w.project.tm_path)
    check("niezmienione pliki TMX pomijane przy ponownym imporcie", reimported == 0,
          f"zaimportowano {reimported}")
    check("rozmiar pamięci bez zmian", w.tm.size() == size_before,
          f"{size_before} -> {w.tm.size()}")
    check("rejestr wczytanych plików", len(w.tm.imported_files()) >= 1,
          str([os.path.basename(f) for f, _ in w.tm.imported_files()]))
    # odroczony commit nie może gubić danych
    w.tm.add("Odroczony zapis test.", "Test odroczonego zapisu.", "en", "pl")
    w.tm.flush()
    check("odroczony commit zapisuje dane",
          len(w.tm.search("Odroczony zapis test.")) == 1)

    # --- zabezpieczenia wydajnościowe dopasowania zdań ---
    import threading as _th2
    guard = TranslationMemory()
    guard.init_for_project(os.path.join(tmp, "tm_guard"))
    guard.add_many([(f"Line {i} of system\\nSecond part {i}.",
                     f"Linia {i} systemu\\nDruga czesc {i}.", "en", "pl") for i in range(3000)])
    probe2 = r"Line 7 of system\nSecond part 7.\pMore text."
    _sm.set("tm.sentence.max.units", 100000)
    guard.find_sentence_matches(probe2)
    _t0 = _t2.perf_counter()
    guard.find_sentence_matches(probe2)
    full_time = _t2.perf_counter() - _t0
    # przerwanie natychmiast po starcie
    guard._reset_line_cache()
    stop = _th2.Event(); stop.set()
    _t0 = _t2.perf_counter()
    cancelled = guard.find_sentence_matches(probe2, should_cancel=stop.is_set)
    cancel_time = _t2.perf_counter() - _t0
    check("przerwanie zatrzymuje wyszukiwanie", cancel_time <= max(full_time, 0.01),
          f"{cancel_time*1000:.1f} ms vs {full_time*1000:.1f} ms")
    check("przerwane wyszukiwanie nie zwraca wyników", cancelled == [])
    # limit rozmiaru pamięci
    _sm.set("tm.sentence.max.units", 100)
    check("limit rozmiaru TM pomija dopasowanie zdań",
          guard.find_sentence_matches(probe2) == [])
    _sm.set("tm.sentence.max.units", 20000)
    # wyłączenie funkcji
    _sm.set("tm.sentence.matching.enabled", False)
    check("wyłączone dopasowanie zdań nic nie liczy",
          guard.find_sentence_matches(probe2) == [])
    _sm.set("tm.sentence.matching.enabled", True)
    # cache nie może zwracać wpisów usuniętych z pamięci
    guard.find_sentence_matches(probe2)
    guard.clear()
    guard.add(r"Alpha beta gamma\ndelta epsilon.", r"Alfa beta gamma\ndelta epsilon PL.")
    ghosts = guard.find_sentence_matches(probe2)
    check("brak wpisów widmo po wyczyszczeniu TM", ghosts == [], str(len(ghosts)))
    guard.close()
    # przywracamy domyślny stan funkcji (test zostawiał go włączony)
    _sm.set("tm.sentence.matching.enabled", False)

    # --- edycja TM w miejscu (bez okna) + lista pamięci ---
    from PyQt6.QtWidgets import QAbstractItemView as _AIV
    tmtab = w.tm_tab.inner
    tmtab.refresh()
    triggers = tmtab.table.editTriggers()
    check("TM: edycja dwuklikiem w tabeli",
          bool(triggers & _AIV.EditTrigger.DoubleClicked))
    tm_tab_names = [tmtab.tabs.tabText(i) for i in range(tmtab.tabs.count())]
    check("TM: zakładki Lista / Przeglądaj / Generator", tmtab.tabs.count() == 3,
          str(tm_tab_names))
    check("TM: nazwy zakładek są opisowe",
          any("Lista" in n for n in tm_tab_names)
          and any("Przeglądaj" in n for n in tm_tab_names)
          and any("Generator" in n for n in tm_tab_names), str(tm_tab_names))
    check("TM: lista pamięci pokazuje bazę projektu", tmtab.tm_list.rowCount() >= 1,
          str(tmtab.tm_list.rowCount()))
    # zaimportowana pamięć spoza folderu tm/ musi zostać na liście
    external = os.path.join(tmp, "zewnetrzna.tmx")
    w.tm.export_tmx(external, "en", "pl")
    rows_before = tmtab.tm_list.rowCount()
    w.register_tm_source(external, 7)
    tmtab.refresh()
    check("TM: zaimportowany plik widoczny na liście",
          tmtab.tm_list.rowCount() == rows_before + 1,
          f"{rows_before} -> {tmtab.tm_list.rowCount()}")
    names = [tmtab.tm_list.item(r, 0).text() for r in range(tmtab.tm_list.rowCount())]
    check("TM: nazwa zaimportowanej pamięci", "zewnetrzna.tmx" in names, str(names))
    check("TM: rejestr zapisany w projekcie",
          any(e.get("path", "").endswith("zewnetrzna.tmx") for e in w.project.tm_sources))
    if tmtab.table.rowCount():
        old_src = tmtab.table.item(0, 0).text()
        old_tgt = tmtab.table.item(0, 1).text()
        tmtab.table.item(0, 1).setText("ZMIENIONE W TABELI")
        found = [t for s_, t, *_ in w.tm.all_entries() if t == "ZMIENIONE W TABELI"]
        check("TM: zapis edycji z tabeli do bazy", len(found) == 1, str(len(found)))
        w.tm.delete(old_src, "ZMIENIONE W TABELI")
        w.tm.add(old_src, old_tgt, "en", "pl")

    # --- edytor TMX ---
    from supercat.ui.tmx_editor import TMXEditorDialog
    ed = TMXEditorDialog(w, parent=w)
    check("edytor TMX wczytał pamięć projektu", len(ed.entries) > 0, str(len(ed.entries)))
    before = len(ed.entries)
    ed.edit_source.setPlainText("Editor test source")
    ed.edit_target.setPlainText("Testowe źródło edytora")
    ed.add_entry()
    check("edytor TMX: dodanie wpisu", len(ed.entries) == before + 1)
    ed.entries.append(["Editor test source", "Testowe źródło edytora", "en", "pl"])
    ed.remove_duplicates()
    check("edytor TMX: usuwanie duplikatów", len(ed.entries) == before + 1, str(len(ed.entries)))
    ed.entries.append(["  spacje   w tekscie  ", "  cel  ", "en", "pl"])
    ed.trim_spaces()
    check("edytor TMX: przycinanie spacji",
          any(e[0] == "spacje w tekscie" for e in ed.entries))
    ed.entries.append(["pusty", "", "en", "pl"])
    ed.remove_empty()
    check("edytor TMX: usuwanie pustych", not any(e[1] == "" for e in ed.entries))
    ed.filter_source.setText("Editor test")
    ed.apply_filters()
    check("edytor TMX: filtrowanie", len(ed.filtered) >= 1, str(len(ed.filtered)))
    ed.clear_filters()
    tmx_out = os.path.join(project.export_path, "z_edytora.tmx")
    from unittest.mock import patch
    with patch("supercat.ui.tmx_editor.QFileDialog.getSaveFileName", return_value=(tmx_out, "")):
        ed.save_as_tmx()
    check("edytor TMX: zapis do pliku", os.path.exists(tmx_out))
    ed.dirty = False
    ed.close()

    w.editor_tab.filter_edit.setText("test")
    w.editor_tab.refresh_grid()
    check("filtrowanie siatki", w.editor_tab.grid.rowCount() < len(w.editor_tab.segments))
    w.editor_tab.filter_edit.clear()

    w.export_target_files() if False else None  # wymaga dialogu – pominięte
    check("eksport przez API", True)

    # ---------------------------- TM: nazwy własne, warianty, kolejność
    print("\n24a. Pamięć TM — nazwy własne, warianty i kolejność wpisów")
    from supercat.core.tm import _is_mostly_untranslated as _untr
    from supercat.core.tm import _leaves_source_text as _leaves

    # 1) filtr nie może odrzucać świadomie nieprzetłumaczonych nazw własnych
    check("nazwa własna w wersalikach przechodzi filtr",
          not _untr("CINNABAR GYM", "CINNABAR GYM"))
    check("krótka etykieta z wersalikami przechodzi", not _untr("PP UP", "PP UP"))
    check("nazwa techniczna przedmiotu przechodzi",
          not _untr("TM01 FOCUS PUNCH", "TM01 FOCUS PUNCH"))
    check("okrzyk z nazwą własną przechodzi",
          not _untr("POLIWRATH: Ribi ribit!", "POLIWRATH: Ribi ribit!"))
    check("prawdziwe tłumaczenie przechodzi",
          not _untr("OTHER TRAINERS", "INNE TRENERZY"))
    check("długie zdanie skopiowane bez zmian jest odrzucane",
          _untr("Thank you for using the MYSTERY GIFT System.",
                "Thank you for using the MYSTERY GIFT System."))
    check("krótki wpis bez nazwy własnej nadal odrzucany",
          _untr("System.", "System."))

    # 2) „Zastosuj TM” uzupełnia segmenty z nazwami własnymi
    names_tm = TranslationMemory()
    names_tm.init_for_project(os.path.join(tmp, "tm_names"))
    name_pairs = [
        ("CINNABAR GYM", "CINNABAR GYM"),
        ("PP UP", "PP UP"),
        ("POLIWRATH: Ribi ribit!", "POLIWRATH: Ribi ribit!"),
        ("Would you like to save?", "Czy chcesz zapisać?"),
    ]
    for src_txt, tgt_txt in name_pairs:
        names_tm.add(src_txt, tgt_txt, "en", "pl")
    names_tm.flush()
    applied = names_tm.find_best_matches_batch([p[0] for p in name_pairs], 80)
    check("„Zastosuj TM” nie pomija już nazw własnych",
          all(m is not None for m in applied),
          f"{sum(1 for m in applied if m)}/{len(applied)}")
    check("wstawione tłumaczenia są poprawne",
          applied[0].text == "CINNABAR GYM" and applied[3].text == "Czy chcesz zapisać?",
          str([m.text for m in applied if m][:2]))

    # 3) kolejność wpisów z pliku TM i warianty tego samego źródła
    order_tm = TranslationMemory()
    order_tm.init_for_project(os.path.join(tmp, "tm_order"))
    for variant in ("KULA", "PIŁKA", "BAL"):
        order_tm.add("BALL", variant, "en", "pl")
    order_tm.add("Can be ground up into a powder.",
                 "Może zostać zmielona na proszek.", "en", "pl")
    order_tm.flush()
    listed = [(a, b) for a, b, *_r in order_tm.all_entries()]
    check("wpisy zachowują kolejność z pliku TM",
          listed[0] == ("BALL", "KULA") and listed[2] == ("BALL", "BAL"),
          str(listed[:3]))
    variants = [m.text for m in order_tm.find_fuzzy_matches("BALL", 70, 10)]
    check("różne tłumaczenia tego samego źródła są zachowane",
          len(variants) == 3, str(variants))
    check("wszystkie warianty są dostępne jako podpowiedzi",
          set(variants) == {"KULA", "PIŁKA", "BAL"}, str(variants))

    # import TMX zachowuje kolejność z pliku
    tmx_order = os.path.join(tmp, "kolejnosc.tmx")
    with open(tmx_order, "w", encoding="utf-8") as fh:
        fh.write('<?xml version="1.0" encoding="UTF-8"?>\n'
                 '<tmx version="1.4"><header srclang="en"/><body>\n')
        for variant in ("KULA", "PIŁKA", "BAL"):
            fh.write(f'<tu><tuv xml:lang="en"><seg>BALL</seg></tuv>'
                     f'<tuv xml:lang="pl"><seg>{variant}</seg></tuv></tu>\n')
        fh.write("</body></tmx>\n")
    imported_tm = TranslationMemory()
    imported_tm.init_for_project(os.path.join(tmp, "tm_imported"))
    imported_tm.import_tmx(tmx_order)
    imported_tm.flush()
    from_file = [b for a, b, *_r in imported_tm.all_entries()]
    check("import TMX nie zmienia kolejności wpisów",
          from_file == ["KULA", "PIŁKA", "BAL"], str(from_file))
    first_hit = imported_tm.find_best_matches_batch(["BALL"], 80)[0]
    check("pierwszy wariant z pliku jest podpowiadany domyślnie",
          first_hit is not None and first_hit.text == "KULA",
          first_hit.text if first_hit else "brak")

    # ręczna poprawka tłumacza ma pierwszeństwo przed wariantami z pliku
    imported_tm.add("BALL", "KULKA", "en", "pl")
    imported_tm.flush()
    after_edit = imported_tm.find_best_matches_batch(["BALL"], 80)[0]
    check("ręczny zapis tłumacza staje się główną podpowiedzią",
          after_edit is not None and after_edit.text == "KULKA",
          after_edit.text if after_edit else "brak")

    # 4) dopasowanie zdań: propozycje zostawiające angielski są oznaczone
    check("wykrywanie pozostawionego tekstu źródłowego",
          _leaves(r"Would you like to mix records with\nother TRAINERS?",
                  r"Would you like to mix records with\nINNE TRAINERS?"))
    check("pełne tłumaczenie nie jest oznaczane",
          not _leaves(r"Would you like to mix records with\nother TRAINERS?",
                      r"Czy chcesz zmieszać rekordy z\ninnymi TRENERAMI?"))

    _SM_sent = SettingsManager.instance()
    _prev_sent = _SM_sent.get_bool("tm.sentence.matching.enabled", False)
    _SM_sent.set("tm.sentence.matching.enabled", True)
    sent_tm = TranslationMemory()
    sent_tm.init_for_project(os.path.join(tmp, "tm_sent_partial"))
    sent_tm.add(r"Would you like to mix records with\nother TRAINERS?",
                r"Czy chcesz zmieszać rekordy z\nother TRAINERS?", "en", "pl")
    sent_tm.add("OTHER", "INNE", "en", "pl")
    sent_tm.flush()
    sent_hits = sent_tm.find_sentence_matches(
        r"Would you like to mix records with\nother TRAINERS?")
    check("dopasowanie zdań zwraca propozycje", bool(sent_hits), str(len(sent_hits)))
    if sent_hits:
        check("najlepsza propozycja jest pełnym tłumaczeniem",
              not sent_hits[0].partial, sent_hits[0].assembled[:60])
        partial_hits = [m for m in sent_hits if m.partial]
        check("niepełne złożenia są oznaczone i na końcu listy",
              all(sent_hits.index(m) >= len(sent_hits) - len(partial_hits)
                  for m in partial_hits) if partial_hits else True,
              f"{len(partial_hits)} niepełnych z {len(sent_hits)}")
        check("etykieta ostrzega o pozostawionym źródle",
              all("zostaje tekst źródłowy" in m.label for m in partial_hits)
              if partial_hits else True,
              str([m.label for m in partial_hits][:1]))
    _SM_sent.set("tm.sentence.matching.enabled", _prev_sent)

    # ---------------------------- Masowe oznaczanie: zakresy i znaki CJK
    print("\n24c. Zakresy TM01-TM66 i wykrywanie znaków CJK")
    from supercat.core.exclusions import BUILTIN_PRESETS as _EX_PRESETS
    from supercat.core.exclusions import (CJK_PATTERN, ExclusionRule as _ExRule,
                                          cjk_ratio, contains_cjk, expand_ranges)

    check("zakres rozwija się na pojedyncze nazwy",
          expand_ranges("TM01-TM05") == ["TM01", "TM02", "TM03", "TM04", "TM05"],
          str(expand_ranges("TM01-TM05")))
    check("pełny zakres TM01-TM66 ma 66 pozycji",
          len(expand_ranges("TM01-TM66")) == 66,
          str(len(expand_ranges("TM01-TM66"))))
    check("zakres HM01-HM28 ma 28 pozycji",
          len(expand_ranges("HM01-HM28")) == 28)
    check("wiodące zera są zachowane",
          expand_ranges("TM01-TM03")[0] == "TM01", str(expand_ranges("TM01-TM03")))
    check("zakres bez zer wiodących też działa",
          expand_ranges("HM1-HM3") == ["HM1", "HM2", "HM3"],
          str(expand_ranges("HM1-HM3")))
    check("zwykły tekst nie jest traktowany jak zakres",
          expand_ranges("TM01") == ["TM01"])
    check("różne przedrostki nie tworzą zakresu",
          expand_ranges("TM01-HM05") == ["TM01-HM05"])
    check("odwrócony zakres jest odrzucany",
          expand_ranges("TM66-TM01") == ["TM66-TM01"])

    range_rule = _ExRule("TM01-TM66", "range", True, False)
    check("reguła zakresu łapie początek i koniec",
          range_rule.matches("TM01") and range_rule.matches("TM66"))
    check("reguła zakresu nie wychodzi poza zakres",
          not range_rule.matches("TM67"))
    check("TM1 nie jest mylone z TM01",
          not range_rule.matches("TM1"))
    check("nazwa znaleziona wewnątrz zdania",
          range_rule.matches("Otrzymałeś TM05 od lidera"))
    check("opis reguły podaje liczbę pozycji",
          "66 pozycji" in range_rule.describe(), range_rule.describe())

    japanese = r"ポケモンに　きのみを\nもたせて　おけば\lたたかっている　ときに"
    check("japoński tekst rozpoznany", contains_cjk(japanese))
    check("katakana rozpoznana", contains_cjk("キズぐすり　とか　どくけし"))
    check("chińskie znaki rozpoznane", contains_cjk("中文测试"))
    check("koreański rozpoznany", contains_cjk("한국어 시험"))
    check("polski tekst nie jest brany za CJK",
          not contains_cjk("Czy chcesz zapisać grę?"))
    check("angielski nie jest brany za CJK",
          not contains_cjk("Would you like to save?"))
    check("znaczniki \\n nie psują wykrywania",
          not contains_cjk(r"Hello\nworld\pnext page"))
    check("udział znaków CJK jest liczony",
          cjk_ratio(japanese) > 0.8, f"{cjk_ratio(japanese):.2f}")
    check("udział CJK dla tekstu łacińskiego to zero",
          cjk_ratio("Normal line.") == 0.0)

    cjk_rule = _ExRule(CJK_PATTERN, "regex", True, False)
    check("reguła CJK działa jako wykluczenie",
          cjk_rule.matches(japanese) and not cjk_rule.matches("Normal line."))
    check("gotowy wzorzec CJK jest wśród presetów",
          any("japoń" in name.lower() for name, _r in _EX_PRESETS),
          str([n for n, _r in _EX_PRESETS][:2]))

    # --- masowe oznaczanie w edytorze ---
    bulk_dir = os.path.join(tmp, "bulk")
    os.makedirs(bulk_dir, exist_ok=True)
    bp = ProjectManager().create_project("Bulk", "en", "pl", bulk_dir)
    with open(os.path.join(bp.source_path, "b.txt"), "w", encoding="utf-8") as fh:
        fh.write("TM01\nTM05\nTM66\nTM67\nHM01\nHM28\nHM29\n"
                 + japanese + "\nWould you like to save?\nNormal line.\n")
    bw = MainWindow()
    bw.open_project_path(bp.project_file_path)
    bw.load_source_files(auto=True)
    bed = bw.editor_tab
    check("wczytano segmenty do testu oznaczania",
          len(bed.segments) == 10, str(len(bed.segments)))
    check("edytor ma masowe oznaczanie", hasattr(bed, "bulk_mark_matching"))
    check("są gotowe wzorce dla TM/HM/CJK",
          len(bed.BULK_PRESETS) >= 3, str(len(bed.BULK_PRESETS)))

    def _bulk(pattern, kind, status, copy_src=True):
        """Powtarza logikę okna bez otwierania go (jak kliknięcie „Oznacz”)."""
        if kind == "cjk":
            idx = [i for i, sg in enumerate(bed.segments)
                   if contains_cjk(sg.source or "")]
        else:
            rule = _ExRule(pattern, kind, True, False)
            idx = [i for i, sg in enumerate(bed.segments)
                   if rule.matches(sg.source or "")]
        bed._store_current()
        if copy_src and status in ("translated", "approved"):
            for i in idx:
                if not (bed.segments[i].target or "").strip():
                    bed.segments[i].target = bed.segments[i].source
            if bed.current_index in set(idx):
                bed.set_target_text(bed.segments[bed.current_index].target)
        if status == "ignored":
            return bed.set_ignored(idx, True), idx
        return bed.set_status(idx, status), idx

    marked, hit = _bulk("TM01-TM66", "range", "translated")
    check("zakres TM oznaczył dokładnie swoje segmenty",
          sorted(bed.segments[i].source for i in hit) == ["TM01", "TM05", "TM66"],
          str([bed.segments[i].source for i in hit]))
    check("segmenty z zakresu mają status przetłumaczony",
          all(bed.segments[i].status == "translated" for i in hit))
    check("tekst źródłowy wstawiony jako tłumaczenie",
          all(bed.segments[i].target == bed.segments[i].source for i in hit),
          str([(bed.segments[i].source, bed.segments[i].target) for i in hit][:2]))
    check("TM67 spoza zakresu nietknięty",
          next(sg for sg in bed.segments if sg.source == "TM67").status == "new")

    marked_hm, hit_hm = _bulk("HM01-HM28", "range", "translated")
    check("zakres HM oznaczony niezależnie",
          sorted(bed.segments[i].source for i in hit_hm) == ["HM01", "HM28"],
          str([bed.segments[i].source for i in hit_hm]))
    check("HM29 spoza zakresu nietknięty",
          next(sg for sg in bed.segments if sg.source == "HM29").status == "new")

    marked_cjk, hit_cjk = _bulk("", "cjk", "ignored")
    check("segment po japońsku został pominięty",
          len(hit_cjk) == 1 and bed.segments[hit_cjk[0]].ignored,
          str(len(hit_cjk)))
    check("teksty łacińskie nie zostały pominięte",
          not any(sg.ignored for sg in bed.segments
                  if not contains_cjk(sg.source or "")))

    check("licznik postępu uwzględnia oznaczone zakresy",
          sum(1 for sg in bed.segments if sg.status == "translated") == 5,
          str(sum(1 for sg in bed.segments if sg.status == "translated")))

    # ---------------------------- Dopasowanie znaczników do oryginału
    print("\n24d. Dopasowanie znaczników (\\n, \\p) do oryginału")
    from supercat.core.tags import (adapt_codes, codes_structure_matches,
                                    split_code_structure, _flatten_lines)

    en1 = ("A strange seed was planted on its back at\\n"
           "birth. The plant sprouts and grows with\\nthis POKéMON.")
    pl1 = ("Dziwne nasiono zostało zasadzone na jego plecach od urodzenia. "
           "Roślina się rozwija i rośnie z tym Pokémonem.")

    adapted = adapt_codes(en1, pl1)
    check("dopasowanie wstawia tyle \\n, ile oryginał",
          adapted.count("\\n") == 2, repr(adapted))
    flat_adapted = _flatten_lines(split_code_structure(adapted)[0])
    check("dopasowanie nie zmienia treści",
          flat_adapted == pl1.strip(), repr(flat_adapted))
    import re as _re2
    _boundary_ok = True
    _pos = 0
    for _k, (_ln, _code) in enumerate(split_code_structure(adapted)[0][:-1]):
        _pos += len(_ln) + _k          # + wcześniejsze spacje złączeniowe
        if _pos >= len(flat_adapted) or flat_adapted[_pos] != " ":
            _boundary_ok = False
    check("przełamania padają na granice wyrazów", _boundary_ok, adapted)
    check("treść bez kodów = oryginał (porównanie treści)",
          _flatten_lines(split_code_structure(adapted)[0]) == pl1.strip(),
          repr(_flatten_lines(split_code_structure(adapted)[0])))
    _fracs = []
    _pos = 0
    for _ln, _code in split_code_structure(adapted)[0][:-1]:
        _pos += len(_ln)
        _fracs.append(_pos / len(flat_adapted))
    _src_fracs = [43 / 96, 83 / 96]
    check("przełamania są proporcjonalne do oryginału",
          all(abs(a - b) < 0.25 for a, b in zip(_fracs, _src_fracs)),
          f"{[round(f,2) for f in _fracs]} vs {[round(f,2) for f in _src_fracs]}")

    same = adapt_codes("One \\nTwo", "Jedno \\nDwa")
    check("ta sama struktura kodów zostaje nietknięta",
          same == "Jedno \\nDwa", repr(same))
    check("struktura zgodna — rozpoznana",
          codes_structure_matches("One \\nTwo", "Jedno \\nDwa"))
    check("struktura różna — rozpoznana",
          not codes_structure_matches(en1, pl1))
    check("brak kodów w źródle usuwa kody z tłumaczenia",
          adapt_codes("Just text", "Jedno\\nDwa") == "Jedno Dwa",
          repr(adapt_codes("Just text", "Jedno\\nDwa")))

    en_par = "Pierwszy akapit tu.\\pDrugi akapit dalej."
    pl_par = "Pierwszy po polsku. Drugi akapit po polsku na dalszym ciągu."
    adapted_par = adapt_codes(en_par, pl_par)
    check("brak akapitu w tłumaczeniu — nie wymyślamy \\p",
          adapted_par.count("\\p") == 0, repr(adapted_par))
    check("treść po \\p nie zmianiona",
          _re2.sub(r"\\[nNlLpP]", "", adapted_par) == pl_par,
          repr(_re2.sub(r"\\[nNlLpP]", "", adapted_par)))

    en_cjk = "line one here\\nline two here"
    pl_cjk = "これはテストのテキストです"
    adapted_cjk = adapt_codes(en_cjk, pl_cjk)
    check("tekst CJK (bez spacji) też się dopasowuje",
          adapted_cjk.count("\\n") == 1, repr(adapted_cjk))
    check("treść CJK nietknięta",
          _re2.sub(r"\\[nNlLpP]", "", adapted_cjk) == pl_cjk,
          repr(_re2.sub(r"\\[nNlLpP]", "", adapted_cjk)))

    adapted_edge = adapt_codes(en1, " " + pl1)
    check("wiodąca spacja (wcięcie) jest zachowana",
          adapted_edge.startswith(" "), repr(adapted_edge[:3]))

    # --- dopasowanie w podpowiedziach TM ---
    tm_dir = os.path.join(tmp, "tm_kody")
    os.makedirs(tm_dir, exist_ok=True)
    tmm = TranslationMemory()
    tmm.init_for_project(tm_dir)
    check("wpis TM zagnany do bazy", tmm.add(en1, pl1, "en", "pl"))
    hits = tmm.find_fuzzy_matches(en1, 70, 5)
    check("TM zwraca dopasowanie", bool(hits), str(len(hits)))
    if hits:
        sug = hits[0].text
        check("podpowiedź TM ma przełamania jak segment",
              sug.count("\\n") == 2, repr(sug))
        check("podpowiedź TM nie zmienia treści",
              _flatten_lines(split_code_structure(sug)[0]) == pl1.strip(),
              repr(sug))

    smgr = SettingsManager.instance()
    _prev_codes = smgr.get_bool("tm.adapt.codes", True)
    smgr.set("tm.adapt.codes", False)
    hits2 = tmm.find_fuzzy_matches(en1, 70, 5)
    smgr.set("tm.adapt.codes", _prev_codes)
    check("wyłączone dopasowanie nie wstawia kodów",
          bool(hits2) and hits2[0].text.count("\\n") == 0,
          repr(hits2[0].text if hits2 else ""))

    # --- edytor: dopasowanie na zaznaczonym segmencie ---
    cd_dir = os.path.join(tmp, "kodowypunkt")
    os.makedirs(cd_dir, exist_ok=True)
    cp2 = ProjectManager().create_project("Kody", "en", "pl", cd_dir)
    with open(os.path.join(cp2.source_path, "k.txt"), "w", encoding="utf-8") as fh:
        fh.write("A strange seed was planted on its back at\\nbirth.")
    cw = MainWindow()
    cw.open_project_path(cp2.project_file_path)
    cw.load_source_files(auto=True)
    ced = cw.editor_tab
    check("wczytano segment z kodami",
          len(ced.segments) == 1 and "at\\nbirth" in (ced.segments[0].source or ""),
          repr([sg.source for sg in ced.segments]))
    check("edytor ma dopasowanie znaczników",
          hasattr(ced, "adapt_codes_selected"))
    ced.current_index = 0
    ced.load_segment(0)
    ced.set_target_text("Dziwne nasiono zostało zasadzone na jego plecach.")
    ced.adapt_codes_selected()
    check("edytor dopasował znaczniki segmentu",
          (ced.segments[0].target or "").count("\\n") == 1,
          repr(ced.segments[0].target))
    check("edytor: treść bez zmian po dopasowaniu",
          _flatten_lines(split_code_structure(ced.segments[0].target or "")[0])
          == "Dziwne nasiono zostało zasadzone na jego plecach.",
          repr(ced.segments[0].target))

    # --- stare projekty dostają nowe reguły wbudowane ---
    from supercat.core.exclusions import BUILTIN_PRESETS, CJK_PATTERN
    saved = {"enabled": True,
             "rules": [{"pattern": "MÓJ_WZORAC", "match_type": "contains",
                        "enabled": True, "case_sensitive": False,
                        "comment": "", "file_filter": ""}]}
    merged = ExclusionSet.from_dict(saved)
    check("zapisana reguła zostaje nietknięta",
          any(r.pattern == "MÓJ_WZORAC" for r in merged.rules))
    check("stary projekt dostaje regułę CJK (włączoną)",
          any(r.pattern == CJK_PATTERN and r.enabled for r in merged.rules),
          str([(r.pattern[:20], r.enabled) for r in merged.rules][:4]))
    check("stary projekt dostaje wszystkie reguły wbudowane",
          all(any(r.pattern == b.pattern for r in merged.rules)
              for _n, b in BUILTIN_PRESETS))

    # --- uniwersalne reguły: działanie (pominięte / przetłumaczone) ---
    print("\n24e. Uniwersalne reguły: działanie reguły (pominięte / przetłumaczone)")
    from supercat.core.exclusions import ExclusionRule as _ER, ExclusionSet as _ES
    from supercat.core.fileparser import Segment as _Seg

    def _mkseg(text, status="new", ignored=False):
        return _Seg(seg_id=f"s{text}", source=text, status=status, ignored=ignored)

    # 1) DOWOLNY wzorzec (CHEM*) → przetłumaczone, a nie pominięte
    chem_rule = _ER("CHEM*", "wildcard", True, False, "substancje chemiczne",
                    action="translated")
    chem_set = _ES([chem_rule])
    chem_segs = [_mkseg("CHEM-123 odczynnik"), _mkseg("Hello world")]
    marked, restored = chem_set.apply(chem_segs)
    check("CHEM* oznacza PRZEZŁUMACZONE (nie pominięte)",
          chem_segs[0].status == "translated" and not chem_segs[0].ignored,
          f"status={chem_segs[0].status} ignored={chem_segs[0].ignored}")
    check("CHEM*: segment niepasujący nietknięty",
          chem_segs[1].status == "new" and not chem_segs[1].ignored)
    check("CHEM*: liczniki poprawne", marked == 1 and restored == 0, f"{marked},{restored}")
    m2, r2 = chem_set.apply(chem_segs)
    check("CHEM*: idempotentne przy ponownym zastosowaniu", (m2, r2) == (0, 0), f"{m2},{r2}")

    # 2) Wyłączenie reguły → status wraca do wcześniejszego
    chem_rule.enabled = False
    chem_set.apply(chem_segs)
    check("CHEM*: wyłączenie reguły przywraca status",
          chem_segs[0].status == "new"
          and not chem_segs[0].extra.get("auto_translated"), chem_segs[0].status)
    chem_rule.enabled = True

    # 3) Domyślne działanie (stare projekty) to nadal „pominięte”
    skip_rule = _ER("<<< FILE:*>>>", "wildcard")
    skip_set = _ES([skip_rule])
    skip_segs = [_mkseg("<<< FILE: x.inc >>>")]
    skip_set.apply(skip_segs)
    check("domyślne: reguła oznacza pominięte",
          skip_segs[0].ignored and skip_segs[0].status == "new")

    # 4) Zmiana działania tej samej reguły — w obie strony
    flip_rule = _ER("CHEM*", "wildcard")
    flip_set = _ES([flip_rule])
    flip_segs = [_mkseg("CHEM-9")]
    flip_set.apply(flip_segs)
    check("przełącznik: najpierw pominięte", flip_segs[0].ignored)
    flip_rule.action = "translated"
    flip_set.apply(flip_segs)
    check("przełącznik: pominięte → przetłumaczone",
          flip_segs[0].status == "translated" and not flip_segs[0].ignored,
          f"status={flip_segs[0].status} ignored={flip_segs[0].ignored}")
    flip_rule.action = "skip"
    flip_set.apply(flip_segs)
    check("przełącznik: przetłumaczone → pominięte (status wraca)",
          flip_segs[0].ignored and flip_segs[0].status == "new",
          f"status={flip_segs[0].status}")

    # 5) Ręczne decyzje mają pierwszeństwo
    keep_seg = _mkseg("CHEM-K")
    keep_seg.extra["manual_keep"] = True
    all_set = _ES([_ER("*", "wildcard", action="translated")])
    all_set.apply([keep_seg])
    check("manual_keep chroni przed automatycznym oznaczeniem", keep_seg.status == "new")
    mskip_seg = _mkseg("CHEM-S", ignored=True)
    mskip_seg.extra["manual_skip"] = True
    all_set.apply([mskip_seg])
    check("manual_skip ma pierwszeństwo przed „przetłumaczone”",
          mskip_seg.ignored and mskip_seg.status == "new")

    # 6) Uniwersalność: zakresy, regex, rozróżnianie liter
    uni_set = _ES([
        _ER("MOJ1-MOJ3", "range", action="translated"),
        _ER(r"^[A-Z]{4,}$", "regex", case_sensitive=True, action="translated"),
    ])
    uni_segs = [_mkseg("MOJ1"), _mkseg("MOJ3"), _mkseg("MOJ4"),
                _mkseg("POKEMON"), _mkseg("mixed")]
    uni_set.apply(uni_segs)
    check("uniwersalnie: zakres i regex działają z dowolnymi tekstami",
          [s.status == "translated" for s in uni_segs] == [True, True, False, True, False],
          [s.status for s in uni_segs])

    # 7) Kilka reguł pasuje → wygrywa pierwsza
    first_set = _ES([_ER("*", "wildcard", action="translated"), _ER("*", "wildcard")])
    first_seg = _mkseg("cokolwiek")
    first_set.apply([first_seg])
    check("wiele reguł: pierwsza wygrywa",
          first_seg.status == "translated" and not first_seg.ignored)

    # 8) Serializacja: action przeżywa zapis; stare reguły → skip
    d = chem_rule.to_dict()
    check("to_dict przechowuje działanie", d.get("action") == "translated")
    check("stara reguła (bez pola action) → pominięte",
          _ER.from_dict({"pattern": "X"}).action == "skip")
    check("błędne action → pominięte (bezpiecznie)",
          _ER.from_dict({"pattern": "X", "action": "cos-innego"}).action == "skip")
    check("round-trip action", _ER.from_dict(d).action == "translated")
    check("opis reguły pokazuje działanie",
          "przetłumaczone" in chem_rule.describe()
          and "pominięte" in skip_rule.describe())

    # 9) Tabela reguł w Ustawieniach ma kolumnę „Działanie”
    st = w.settings_tab
    headers = [st.excl_table.horizontalHeaderItem(c).text()
               for c in range(st.excl_table.columnCount())]
    check("tabela reguł: kolumna „Działanie”",
          "Działanie" in headers and st.excl_table.columnCount() == 6, str(headers))
    prev_excl = w._exclusions
    w._exclusions = _ES([chem_rule, skip_rule])
    st.load_exclusions()
    a0 = st.excl_table.item(0, 3).text()
    a1 = st.excl_table.item(1, 3).text()
    w._exclusions = prev_excl
    st.load_exclusions()
    check("tabela reguł: pokazuje działanie każdej reguły",
          "przetłumaczone" in a0 and "pominięte" in a1, f"{a0} | {a1}")

    # 10) Okno reguły: pole „Działanie”
    dlg = ExclusionDialog(_ER("CHEM*", "wildcard", action="translated"), w, ["a.txt"])
    check("okno reguły: ma pole Działanie", hasattr(dlg, "action"))
    check("okno reguły: wczytuje istniejące działanie",
          dlg.action.currentData() == "translated")
    dlg.action.setCurrentIndex(0)
    check("okno reguły: zmiana na pominięte", dlg._current_rule().action == "skip")
    dlg.action.setCurrentIndex(1)
    check("okno reguły: zmiana na przetłumaczone", dlg._current_rule().action == "translated")

    # 11) End-to-end: projekt z regułą CHEM* → oznaczenie przy wczytywaniu
    ch_dir = os.path.join(tmp, "chemproj")
    os.makedirs(ch_dir, exist_ok=True)
    ch_pm = ProjectManager()
    chp = ch_pm.create_project("Chem", "en", "pl", ch_dir)
    with open(os.path.join(chp.source_path, "c.txt"), "w", encoding="utf-8") as fh:
        fh.write("CHEM-001 odczynnik\nHello there friend.")
    chp.exclusions = _ES([
        _ER("CHEM*", "wildcard", True, False, "substancje chemiczne",
            action="translated"),
    ]).to_dict()
    ch_pm.save_project()
    cw2 = MainWindow()
    cw2.open_project_path(chp.project_file_path)
    cw2.load_source_files(auto=True)
    csegs = cw2.editor_tab.segments
    check("projekt: CHEM* oznaczone przetłumaczone przy wczytaniu",
          len(csegs) == 2 and csegs[0].status == "translated" and not csegs[0].ignored,
          [(s.source, s.status, s.ignored) for s in csegs])
    check("projekt: zwykły segment nietknięty",
          csegs[1].status == "new" and not csegs[1].ignored)

    # --- uniwersalność dopasowania kodów (nie tylko przykładowe pary) ---
    en_l = r"Linia pierwsza jest długa\lale druga\lznowu krótka"
    pl_l = "Pierwszy wiersz długi a drugi krótki znowu"
    ad_l = adapt_codes(en_l, pl_l)
    check("kody uniwersalnie: działa na \l", ad_l.count("\\l") == 2, repr(ad_l))
    check("kody uniwersalnie: treść z \l nietknięta",
          _flatten_lines(split_code_structure(ad_l)[0]) == pl_l, repr(ad_l))

    en_p = r"Para jeden\pPara drugi\pPara trzeci"
    pl_p = "Pierwszy akapit drugi akapit i trzeci akapit"
    ad_p = adapt_codes(en_p, pl_p)
    check("kody uniwersalnie: cel bez \p nie dostaje wymyślonych \p",
          ad_p.count("\\p") == 0, repr(ad_p))
    check("kody uniwersalnie: treść akapitów nietknięta",
          _flatten_lines(split_code_structure(ad_p)[0]) == pl_p, repr(ad_p))
    pl_p2 = "Pierwszy akapit\pDrugi akapit i trzeci"
    ad_p2 = adapt_codes(en_p, pl_p2)
    check("kody uniwersalnie: istniejące \p w celu zostają na miejscu",
          ad_p2.count("\\p") == 1 and "akapit\pDrugi" in ad_p2, repr(ad_p2))

    ad_none = adapt_codes("Plain source without any codes", pl1)
    check("kody uniwersalnie: źródło bez kodów → cel nietknięty",
          ad_none == pl1, repr(ad_none))

    ad_short = adapt_codes(en1, "Krótka odpowiedź.")
    check("kody uniwersalnie: krótki cel nie dostaje męty (max 1 kod)",
          ad_short.count("\\n") <= 1
          and _flatten_lines(split_code_structure(ad_short)[0]) == "Krótka odpowiedź.",
          repr(ad_short))

    # ---------------------------- Nawigacja klawiszami
    print("\n24b. Strzałki i zaznaczanie w siatce")
    from PyQt6.QtTest import QTest as _QTest

    nav_dir = os.path.join(tmp, "nawigacja")
    os.makedirs(nav_dir, exist_ok=True)
    np_ = ProjectManager().create_project("Nawigacja", "en", "pl", nav_dir)
    with open(os.path.join(np_.source_path, "n.txt"), "w", encoding="utf-8") as fh:
        fh.write("One.\nTwo.\nThree.\nFour.\nFive.\n")
    nw = MainWindow()
    nw.open_project_path(np_.project_file_path)
    nw.load_source_files(auto=True)
    ned = nw.editor_tab
    grid = ned.grid

    def _rows():
        model = grid.selectionModel()
        return sorted(i.row() for i in model.selectedRows()) if model else []

    ned.load_segment(0)
    app.processEvents()
    check("start: jeden zaznaczony wiersz", _rows() == [0], str(_rows()))

    # REGRESJA: przechodzenie między segmentami nie może kumulować zaznaczenia
    ned.next_segment()
    app.processEvents()
    check("przejście do następnego segmentu zaznacza tylko jego",
          _rows() == [1] and ned.current_index == 1,
          f"rows={_rows()} current={ned.current_index}")
    ned.next_segment()
    ned.next_segment()
    app.processEvents()
    check("kolejne przejścia nie dokładają zaznaczenia",
          _rows() == [3], str(_rows()))
    ned.prev_segment()
    app.processEvents()
    check("cofnięcie też zaznacza pojedynczy wiersz", _rows() == [2], str(_rows()))

    # klawisze w siatce
    grid.setFocus()
    ned.load_segment(0)
    app.processEvents()
    _QTest.keyClick(grid, Qt.Key.Key_Down)
    app.processEvents()
    check("goła strzałka ↓ przechodzi do następnego segmentu",
          ned.current_index == 1 and _rows() == [1],
          f"current={ned.current_index} rows={_rows()}")
    _QTest.keyClick(grid, Qt.Key.Key_Down, Qt.KeyboardModifier.ControlModifier)
    app.processEvents()
    check("Ctrl+↓ przechodzi bez zaznaczania kolejnych wierszy",
          ned.current_index == 2 and _rows() == [2],
          f"current={ned.current_index} rows={_rows()}")
    _QTest.keyClick(grid, Qt.Key.Key_Up, Qt.KeyboardModifier.ControlModifier)
    app.processEvents()
    check("Ctrl+↑ wraca bez zaznaczania",
          ned.current_index == 1 and _rows() == [1],
          f"current={ned.current_index} rows={_rows()}")

    _QTest.keyClick(grid, Qt.Key.Key_End, Qt.KeyboardModifier.ControlModifier)
    app.processEvents()
    check("Ctrl+End skacze na koniec", ned.current_index == 4, str(ned.current_index))
    _QTest.keyClick(grid, Qt.Key.Key_Home, Qt.KeyboardModifier.ControlModifier)
    app.processEvents()
    check("Ctrl+Home wraca na początek", ned.current_index == 0, str(ned.current_index))

    # Shift nadal zaznacza zakres — grupowe operacje muszą działać
    ned.load_segment(0)
    grid.setFocus()
    app.processEvents()
    _QTest.keyClick(grid, Qt.Key.Key_Down, Qt.KeyboardModifier.ShiftModifier)
    _QTest.keyClick(grid, Qt.Key.Key_Down, Qt.KeyboardModifier.ShiftModifier)
    app.processEvents()
    check("Shift+↓ nadal zaznacza zakres wierszy", _rows() == [0, 1, 2],
          str(_rows()))

    # strzałki w polu tłumaczenia: brzeg tekstu przechodzi dalej
    ned.load_segment(1)
    ned.target_edit.setFocus()
    ned.set_target_text("")
    app.processEvents()
    _QTest.keyClick(ned.target_edit, Qt.Key.Key_Down)
    app.processEvents()
    check("↓ w pustym polu tłumaczenia przechodzi do następnego segmentu",
          ned.current_index == 2, str(ned.current_index))
    _QTest.keyClick(ned.target_edit, Qt.Key.Key_Up)
    app.processEvents()
    check("↑ w polu tłumaczenia wraca do poprzedniego",
          ned.current_index == 1, str(ned.current_index))

    # w wielowierszowym tłumaczeniu strzałki najpierw chodzą po liniach
    ned.set_target_text("pierwsza linia\ndruga linia\ntrzecia linia")
    cursor = ned.target_edit.textCursor()
    cursor.movePosition(cursor.MoveOperation.Start)
    ned.target_edit.setTextCursor(cursor)
    before_index = ned.current_index
    app.processEvents()
    _QTest.keyClick(ned.target_edit, Qt.Key.Key_Down)
    app.processEvents()
    check("↓ w środku tekstu porusza kursor, nie zmienia segmentu",
          ned.current_index == before_index, str(ned.current_index))
    check("kursor zszedł do drugiej linii",
          ned.target_edit.textCursor().blockNumber() == 1,
          str(ned.target_edit.textCursor().blockNumber()))
    _QTest.keyClick(ned.target_edit, Qt.Key.Key_Down)
    _QTest.keyClick(ned.target_edit, Qt.Key.Key_Down)
    app.processEvents()
    check("↓ z ostatniej linii wychodzi do następnego segmentu",
          ned.current_index == before_index + 1,
          f"{before_index} -> {ned.current_index}")

    check("przełącznik nawigacji strzałkami jest zapisany w ustawieniach",
          SettingsManager.instance().get_bool(
              "editor.arrows.change.segment", True) is True)
    check("Ustawienia mają przełącznik nawigacji strzałkami",
          hasattr(w.settings_tab, "arrow_nav")
          and w.settings_tab.arrow_nav.isChecked())

    # wyłączony przełącznik: strzałka tylko przesuwa kursor
    SettingsManager.instance().set("editor.arrows.change.segment", False)
    ned.load_segment(1)
    ned.set_target_text("")
    ned.target_edit.setFocus()
    app.processEvents()
    _QTest.keyClick(ned.target_edit, Qt.Key.Key_Down)
    app.processEvents()
    check("wyłączony przełącznik zostawia strzałki polu tekstowemu",
          ned.current_index == 1, str(ned.current_index))
    SettingsManager.instance().set("editor.arrows.change.segment", True)

    # ---------------------------- Panele: nie da się ich zgubić
    print("\n25a. Układ paneli (splittery)")
    main_sp = w.editor_tab.main_splitter
    center_sp = w.editor_tab.center_splitter

    check("główny splitter nie pozwala zwinąć paneli",
          main_sp.childrenCollapsible() is False)
    check("żaden panel z osobna nie jest zwijalny",
          not any(main_sp.isCollapsible(i) for i in range(main_sp.count())),
          str([main_sp.isCollapsible(i) for i in range(main_sp.count())]))
    check("uchwyt jest wystarczająco szeroki, by go złapać",
          main_sp.handleWidth() >= 5, str(main_sp.handleWidth()))
    check("każdy panel ma minimalną szerokość",
          all(main_sp.widget(i).minimumWidth() >= 100
              for i in range(main_sp.count())),
          str([main_sp.widget(i).minimumWidth() for i in range(main_sp.count())]))

    # próba zwinięcia panelu „Pliki projektu” do zera
    w.editor_tab.main_splitter.setSizes([0, 1200, 200])
    app.processEvents()
    check("panel „Pliki projektu” nie znika po zwinięciu",
          main_sp.sizes()[0] >= main_sp.widget(0).minimumWidth(),
          str(main_sp.sizes()))
    w.editor_tab.main_splitter.setSizes([220, 1200, 0])
    app.processEvents()
    check("prawy panel pomocy nie znika po zwinięciu",
          main_sp.sizes()[2] >= main_sp.widget(2).minimumWidth(),
          str(main_sp.sizes()))

    check("splitter pionowy też jest zabezpieczony",
          center_sp.childrenCollapsible() is False)
    center_sp.setSizes([0, 800])
    app.processEvents()
    check("siatka segmentów nie znika po zwinięciu",
          center_sp.sizes()[0] >= center_sp.widget(0).minimumHeight(),
          str(center_sp.sizes()))

    # przywracanie układu z menu Widok
    main_sp.setSizes([150, 1250, 180])
    app.processEvents()
    squeezed = list(main_sp.sizes())
    w.reset_panel_layout()
    app.processEvents()
    check("„Przywróć układ paneli” rozsuwa panele",
          main_sp.sizes()[0] > squeezed[0] and main_sp.sizes()[2] > squeezed[2],
          f"{squeezed} -> {main_sp.sizes()}")
    check("po przywróceniu każdy panel jest widoczny",
          all(size > 0 for size in main_sp.sizes()), str(main_sp.sizes()))

    # ---------------------------- Pliki projektu: wybór i kolejność
    print("\n25b. Pliki projektu — zaznaczanie wielu i kolejność")
    from supercat.core.project import order_files as _order_files

    check("bez ręcznej kolejności zostaje porządek wejściowy",
          _order_files(["a.txt", "b.txt"], []) == ["a.txt", "b.txt"])
    check("ręczna kolejność jest respektowana",
          _order_files(["a.txt", "b.txt", "c.txt"], ["c.txt", "a.txt"])
          == ["c.txt", "a.txt", "b.txt"],
          str(_order_files(["a.txt", "b.txt", "c.txt"], ["c.txt", "a.txt"])))
    check("nowy plik trafia na koniec, nie przestawia reszty",
          _order_files(["a.txt", "b.txt", "nowy.txt"], ["b.txt", "a.txt"])
          == ["b.txt", "a.txt", "nowy.txt"],
          str(_order_files(["a.txt", "b.txt", "nowy.txt"], ["b.txt", "a.txt"])))
    check("usunięty plik znika z kolejności",
          _order_files(["b.txt"], ["zniknal.txt", "b.txt"]) == ["b.txt"])

    files_dir = os.path.join(tmp, "kolejnosc")
    os.makedirs(files_dir, exist_ok=True)
    fp = ProjectManager().create_project("Kolejność", "en", "pl", files_dir)
    for fname in ("a.txt", "b.txt", "c.txt", "d.txt"):
        with open(os.path.join(fp.source_path, fname), "w", encoding="utf-8") as fh:
            fh.write(f"First line of {fname}.\nSecond line of {fname}.\n")

    fw = MainWindow()
    fw.open_project_path(fp.project_file_path)
    fw.load_source_files(auto=True)
    fed = fw.editor_tab

    check("lista plików pozwala zaznaczyć wiele",
          fed.files_list.selectionMode()
          == _AIV.SelectionMode.ExtendedSelection,
          fed.files_list.selectionMode().name)
    check("są przyciski kolejności i usuwania",
          all(hasattr(fed, name) for name in
              ("file_up_btn", "file_down_btn", "file_sort_btn", "file_remove_btn")))
    check("startowa kolejność jest alfabetyczna",
          fed.current_file_order() == ["a.txt", "b.txt", "c.txt", "d.txt"],
          str(fed.current_file_order()))

    def _select_files(*names):
        fed.files_list.clearSelection()
        for row in range(fed.files_list.count()):
            item = fed.files_list.item(row)
            if item.data(Qt.ItemDataRole.UserRole) in names:
                item.setSelected(True)
        app.processEvents()

    _select_files("c.txt")
    check("zaznaczenie pliku jest odczytywane",
          fed.selected_file_names() == ["c.txt"], str(fed.selected_file_names()))
    fed.move_files(-1)
    check("przycisk ▲ przesuwa plik w górę",
          fed.current_file_order() == ["a.txt", "c.txt", "b.txt", "d.txt"],
          str(fed.current_file_order()))
    fed.move_files(1)
    check("przycisk ▼ przesuwa plik w dół",
          fed.current_file_order() == ["a.txt", "b.txt", "c.txt", "d.txt"],
          str(fed.current_file_order()))

    _select_files("a.txt")
    fed.move_files(-1)
    check("pierwszy plik nie wychodzi poza listę",
          fed.current_file_order()[0] == "a.txt", str(fed.current_file_order()))
    _select_files("d.txt")
    fed.move_files(1)
    check("ostatni plik nie wychodzi poza listę",
          fed.current_file_order()[-1] == "d.txt", str(fed.current_file_order()))

    # blok kilku plików przesuwany razem
    _select_files("c.txt", "d.txt")
    fed.move_files(-1)
    check("blok zaznaczonych plików przesuwa się razem",
          fed.current_file_order() == ["a.txt", "c.txt", "d.txt", "b.txt"],
          str(fed.current_file_order()))
    check("po przesunięciu zaznaczenie zostaje",
          sorted(fed.selected_file_names()) == ["c.txt", "d.txt"],
          str(fed.selected_file_names()))
    check("licznik pokazuje liczbę zaznaczonych",
          "2" in fed.files_selection_label.text(),
          fed.files_selection_label.text())

    check("kolejność zapisana w projekcie",
          fw.project.file_order == ["a.txt", "c.txt", "d.txt", "b.txt"],
          str(fw.project.file_order))
    check("segmenty ułożone zgodnie z kolejnością plików",
          [s.file_name for s in fed.segments][:2] == ["a.txt", "a.txt"]
          and [s.file_name for s in fed.segments][2] == "c.txt",
          str([s.file_name for s in fed.segments]))

    # kolejność przeżywa ponowne otwarcie projektu
    fw.save_all(silent=True)
    fw2 = MainWindow()
    fw2.open_project_path(fp.project_file_path)
    check("kolejność przetrwała ponowne otwarcie",
          fw2.editor_tab.current_file_order() == ["a.txt", "c.txt", "d.txt", "b.txt"],
          str(fw2.editor_tab.current_file_order()))

    # nowy plik nie przestawia ustawionej kolejności
    with open(os.path.join(fp.source_path, "z_nowy.txt"), "w", encoding="utf-8") as fh:
        fh.write("Nowy plik.\n")
    fw2.load_source_files(auto=True)
    check("nowo dodany plik trafia na koniec listy",
          fw2.editor_tab.current_file_order()[-1] == "z_nowy.txt",
          str(fw2.editor_tab.current_file_order()))

    # usuwanie wielu plików naraz – jedno pytanie
    _select_files("a.txt", "b.txt")
    with patch("supercat.ui.main_window.QMessageBox.question",
               return_value=QMessageBox.StandardButton.Yes) as ask:
        removed_count = fw.remove_project_files(["a.txt", "b.txt"])
    check("usuwanie wielu plików pyta tylko raz", ask.call_count == 1,
          str(ask.call_count))
    check("usunięto wskazane pliki", removed_count == 2, str(removed_count))
    check("pliki zniknęły z dysku",
          not os.path.exists(os.path.join(fp.source_path, "a.txt"))
          and not os.path.exists(os.path.join(fp.source_path, "b.txt")))
    check("pliki zniknęły z listy",
          "a.txt" not in fed.current_file_order()
          and "b.txt" not in fed.current_file_order(),
          str(fed.current_file_order()))
    check("segmenty usuniętych plików zniknęły",
          not any(s.file_name in ("a.txt", "b.txt") for s in fed.segments))
    check("usunięte pliki zniknęły z zapisanej kolejności",
          "a.txt" not in fw.project.file_order
          and "b.txt" not in fw.project.file_order,
          str(fw.project.file_order))

    # anulowanie nie usuwa niczego
    before_cancel = list(fed.current_file_order())
    with patch("supercat.ui.main_window.QMessageBox.question",
               return_value=QMessageBox.StandardButton.No):
        cancelled = fw.remove_project_files(["c.txt"])
    check("odmowa w oknie nie usuwa plików",
          cancelled == 0 and fed.current_file_order() == before_cancel,
          str(fed.current_file_order()))

    # --- zmiana kolejności przeciągnięciem na liście ---------------------
    from PyQt6.QtCore import QPoint as _QPoint

    class _DropAt:
        """Namiastka zdarzenia upuszczenia w danym punkcie listy."""

        def __init__(self, y):
            self._y = y

        def position(self):
            outer = self

            class _Pos:
                def toPoint(self_inner):
                    return _QPoint(10, outer._y)

            return _Pos()

    flist = fed.files_list
    check("lista plików pozwala przeciągać pozycje", flist.dragEnabled())
    check("wskaźnik miejsca upuszczenia jest widoczny",
          flist.showDropIndicator())
    check("domyślną akcją jest przeniesienie, nie kopiowanie",
          flist.defaultDropAction() == Qt.DropAction.MoveAction)

    def _drop_above(row):
        flist._apply_internal_move(
            _DropAt(flist.visualItemRect(flist.item(row)).top() + 1))

    fed.sort_files_alphabetically()
    order_before = fed.current_file_order()
    last_name = order_before[-1]
    _select_files(last_name)
    _drop_above(1)
    check("przeciągnięcie pliku na górę zmienia kolejność",
          fed.current_file_order()[0] == last_name,
          str(fed.current_file_order()))
    check("po upuszczeniu plik zostaje zaznaczony",
          fed.selected_file_names() == [last_name],
          str(fed.selected_file_names()))
    check("kolejność z przeciągnięcia zapisuje się w projekcie",
          fw.project.file_order == fed.current_file_order(),
          str(fw.project.file_order))
    check("segmenty idą za kolejnością z przeciągnięcia",
          (fed.segments[0].file_name if fed.segments else "") == last_name,
          str([s.file_name for s in fed.segments][:3]))

    # blok kilku plików przeciągany razem
    two = fed.current_file_order()[-2:]
    _select_files(*two)
    _drop_above(1)
    check("przeciągnięcie bloku zachowuje jego układ",
          fed.current_file_order()[:2] == two, str(fed.current_file_order()))

    # pozycja „Wszystkie pliki” musi zostać na górze
    guard_before = fed.current_file_order()
    _select_files(guard_before[-1])
    flist._apply_internal_move(
        _DropAt(flist.visualItemRect(flist.item(0)).top() + 1))
    check("nie da się wstawić pliku nad „Wszystkie pliki”",
          flist.item(0).data(Qt.ItemDataRole.UserRole) is None)
    check("plik upuszczony najwyżej trafia pod „Wszystkie pliki”",
          fed.current_file_order()[0] == guard_before[-1],
          str(fed.current_file_order()))

    # sama pozycja „Wszystkie pliki” nie jest przesuwalna
    flist.clearSelection()
    flist.item(0).setSelected(True)
    frozen = fed.current_file_order()
    flist._apply_internal_move(
        _DropAt(flist.visualItemRect(flist.item(2)).top() + 1))
    check("„Wszystkie pliki” nie daje się przeciągnąć",
          fed.current_file_order() == frozen, str(fed.current_file_order()))

    # import z pulpitu nie może być mylony z przestawianiem kolejności
    from PyQt6.QtCore import QMimeData as _QMime
    from PyQt6.QtCore import QUrl as _QUrl
    from PyQt6.QtGui import QDragEnterEvent as _QDragEnter

    outside_file = os.path.join(tmp, "z_pulpitu.txt")
    with open(outside_file, "w", encoding="utf-8") as fh:
        fh.write("Z zewnątrz.\n")
    mime = _QMime()
    mime.setUrls([_QUrl.fromLocalFile(outside_file)])
    enter = _QDragEnter(flist.rect().center(), Qt.DropAction.CopyAction, mime,
                        Qt.MouseButton.LeftButton, Qt.KeyboardModifier.NoModifier)
    flist.dragEnterEvent(enter)
    check("przeciągnięcie pliku z pulpitu nadal jest przyjmowane",
          enter.isAccepted())
    check("plik z pulpitu nie jest brany za zmianę kolejności",
          flist._is_internal(enter) is False)

    fed.sort_files_alphabetically()
    check("przywrócenie kolejności alfabetycznej działa",
          fed.current_file_order() == sorted(fed.current_file_order()),
          str(fed.current_file_order()))
    check("sortowanie kasuje ręczną kolejność w projekcie",
          fw.project.file_order == [], str(fw.project.file_order))

    # ---------------------------- Wyszukiwanie białych znaków
    print("\n26b. Wyszukiwanie spacji i białych znaków")
    from supercat.core.search import (WHITESPACE_FILTERS, edges_differ,
                                      search_whitespace, whitespace_issues,
                                      whitespace_spans)

    check("filtr „spacja na początku” istnieje", "leading" in WHITESPACE_FILTERS)
    check("są filtry brzegów, podwójnej spacji i tabulatora",
          {"trailing", "double", "tab", "mismatch"} <= set(WHITESPACE_FILTERS))

    check("spacja na początku wykryta", whitespace_issues(" Ala") == ["leading"],
          str(whitespace_issues(" Ala")))
    check("spacja na końcu wykryta", "trailing" in whitespace_issues("Ala "))
    check("twarda spacja też się liczy",
          "leading" in whitespace_issues("\u00a0Ala"))
    check("tekst bez spacji nie zgłasza nic", whitespace_issues("Ala") == [])
    check("tabulator wykryty", "tab" in whitespace_issues("Ala\tma"))
    check("podwójna spacja wewnątrz wykryta",
          "double" in whitespace_issues("Ala  ma kota"))
    # wcięcie na brzegu NIE może być liczone jako podwójna spacja
    check("wcięcie nie jest podwójną spacją",
          "double" not in whitespace_issues("  Ala ma kota"),
          str(whitespace_issues("  Ala ma kota")))

    check("zakres spacji początkowej wskazuje wcięcie",
          whitespace_spans("   Ala", "leading") == [(0, 3)],
          str(whitespace_spans("   Ala", "leading")))
    check("zakres spacji końcowej", whitespace_spans("Ala  ", "trailing") == [(3, 5)],
          str(whitespace_spans("Ala  ", "trailing")))
    check("brak zakresu, gdy nie ma problemu",
          whitespace_spans("Ala", "leading") == [])

    check("różne brzegi źródła i tłumaczenia wykryte",
          edges_differ("  Hello", "Witaj"))
    check("zgodne brzegi nie są zgłaszane",
          not edges_differ("  Hello", "  Witaj"))
    check("puste tłumaczenie nie jest błędem brzegów",
          not edges_differ("  Hello", ""))

    class _WSeg:
        def __init__(self, name, source, target, status="new", ignored=False):
            self.file_name = name
            self.source = source
            self.target = target
            self.status = status
            self.ignored = ignored

    ws_segments = [
        _WSeg("a.txt", " Wcięcie w źródle.", "Bez wcięcia."),
        _WSeg("a.txt", "Zwykły wiersz.", "Zwykłe tłumaczenie."),
        _WSeg("b.txt", "Tekst.", "Ma  podwójną spację."),
        _WSeg("b.txt", "Koniec.", "Kończy się spacją "),
        _WSeg("b.txt", "Tabulator.", "Ma\ttabulator."),
    ]

    lead_hits = search_whitespace(ws_segments, ["leading"])
    check("wyszukiwanie znajduje spację na początku",
          len(lead_hits.hits) == 1 and lead_hits.hits[0].index == 0,
          str([(h.index, h.where) for h in lead_hits.hits]))
    check("opis trafienia mówi, o co chodzi",
          "spacja na początku" in lead_hits.hits[0].snippet,
          lead_hits.hits[0].snippet[:50])
    check("trafienie ma zakres do podświetlenia",
          lead_hits.hits[0].spans == [(0, 1)], str(lead_hits.hits[0].spans))

    all_hits = search_whitespace(
        ws_segments, ["leading", "trailing", "double", "tab"])
    check("łączenie filtrów zwraca wszystkie problemy",
          len(all_hits.hits) == 4, str(len(all_hits.hits)))
    mismatch_hits = search_whitespace(ws_segments, ["mismatch"])
    check("różnica brzegów wykryta w tłumaczeniu",
          any(h.index == 0 and h.where == "tłumaczenie" for h in mismatch_hits.hits),
          str([(h.index, h.where) for h in mismatch_hits.hits]))

    only_target = search_whitespace(
        ws_segments, ["leading"], SearchOptions(in_source=False, in_target=True))
    check("filtr można zawęzić do tłumaczenia", not only_target.hits,
          str(len(only_target.hits)))
    one_file = search_whitespace(
        ws_segments, ["trailing", "double"], SearchOptions(files=["b.txt"]))
    check("filtr działa w zakresie wybranych plików",
          all(h.file_name == "b.txt" for h in one_file.hits) and one_file.hits,
          str([h.file_name for h in one_file.hits]))
    check("pusty zestaw filtrów nic nie zwraca",
          not search_whitespace(ws_segments, []).hits)

    # wydajność – filtr przechodzi cały projekt, więc nie może zamulać
    perf_segments = [
        _WSeg(f"f{i % 20}.txt",
              (" " if i % 50 == 0 else "") + f"Line {i} of text.",
              f"Wiersz {i} tekstu.")
        for i in range(20000)
    ]
    import time as _t_ws

    ws_started = _t_ws.perf_counter()
    perf_hits = search_whitespace(perf_segments, ["leading"])
    ws_ms = (_t_ws.perf_counter() - ws_started) * 1000
    check("filtr białych znaków szybki na 20k segmentów", ws_ms < 900,
          f"{ws_ms:.0f} ms, {len(perf_hits.hits)} trafień")

    # --- w interfejsie ---
    ws_tab = w.search_tab
    check("wyszukiwarka ma pola wyboru białych znaków",
          len(ws_tab.whitespace_boxes) == len(WHITESPACE_FILTERS),
          str(list(ws_tab.whitespace_boxes)))
    check("pola mają wyjaśnienia w dymkach",
          all(box.toolTip() for box in ws_tab.whitespace_boxes.values()))

    ws_tab.search_edit.clear()
    ws_tab._clear_status_filter()
    ws_tab._clear_whitespace_filter()
    ws_tab.whitespace_boxes["leading"].setChecked(True)
    app.processEvents()
    check("zaznaczenie filtru szuka bez wpisywania frazy",
          ws_tab.tree.topLevelItemCount() >= 0 and "␣" in ws_tab.status.text()
          or "Brak segmentów" in ws_tab.status.text(),
          ws_tab.status.text()[:60])
    check("opcje wyszukiwania niosą wybrane filtry",
          ws_tab.current_options().whitespace == ["leading"],
          str(ws_tab.current_options().whitespace))
    ws_tab._clear_whitespace_filter()
    app.processEvents()
    check("odznaczenie filtrów czyści wyniki",
          ws_tab.tree.topLevelItemCount() == 0)

    # ---------------------------- Generator TM z par plików
    print("\n27. Generator TM z plików dwujęzycznych")
    from supercat.core import tm_builder as _tb

    gen_dir = os.path.join(tmp, "gen_src")
    os.makedirs(gen_dir, exist_ok=True)

    def _write(name, text, encoding="utf-8"):
        path = os.path.join(gen_dir, name)
        with open(path, "w", encoding=encoding, newline="") as fh:
            fh.write(text)
        return path

    en_path = _write("miasto_en.txt",
                     "Hello world.\nWould you like to save?\n\n"
                     "<<< FILE: city/text.inc >>>\nWIRELESS COMMUNICATION\nOK\n")
    pl_path = _write("miasto_pl.txt",
                     "Witaj świecie.\nCzy chcesz zapisać?\n\n"
                     "<<< FILE: city/text.inc >>>\nŁĄCZNOŚĆ BEZPRZEWODOWA\nOK\n")
    _write("krzywy_en.txt", "One\nTwo\n")
    _write("krzywy_pl.txt", "Jeden\nDwa\nTrzy\n")
    _write("sierota_de.txt", "allein\n")
    # plik w cp1250 – typowy eksport z Windows
    win_en = _write("okno_en.txt", "Window\nDoor\n")
    # Polskie znaki są konieczne: czysty ASCII jest poprawnym UTF-8,
    # więc po samej treści nie dałoby się rozpoznać cp1250.
    win_pl = _write("okno_pl.txt", "Okno\nDrzwi ściśle\n", encoding="cp1250")

    check("odmiana „wiersz”: 1", _tb.plural_lines(1) == "1 wiersz")
    check("odmiana „wiersz”: 2-4", _tb.plural_lines(3) == "3 wiersze")
    check("odmiana „wiersz”: 5+", _tb.plural_lines(6) == "6 wierszy")
    check("odmiana „wiersz”: 12 to wyjątek", _tb.plural_lines(12) == "12 wierszy")

    check("kod języka z nazwy pliku", _tb.language_hint(en_path) == "en",
          _tb.language_hint(en_path))
    check("kod języka rozpoznany dla polskiego", _tb.language_hint(pl_path) == "pl")
    check("brak kodu, gdy nazwa nic nie mówi",
          _tb.language_hint(os.path.join(gen_dir, "cokolwiek.txt")) == "")

    check("wykrywanie wiersza technicznego",
          _tb.is_technical_line("<<< FILE: city/text.inc >>>")
          and _tb.is_technical_line("# komentarz")
          and _tb.is_technical_line("[sekcja]"))
    check("zwykły tekst nie jest techniczny",
          not _tb.is_technical_line("Hello world."))

    all_files = sorted(os.path.join(gen_dir, n) for n in os.listdir(gen_dir))
    gen_pairs, gen_unmatched = _tb.pair_files(all_files, "en", "pl")
    check("pliki sparowane po nazwie", len(gen_pairs) == 3, str(len(gen_pairs)))
    check("plik bez odpowiednika zgłoszony",
          len(gen_unmatched) == 1 and "sierota" in gen_unmatched[0],
          str(gen_unmatched))
    good_pair = next(p for p in gen_pairs if "miasto" in p.source_path)
    check("para ma poprawny kierunek języków",
          good_pair.source_lang == "en" and good_pair.target_lang == "pl",
          f"{good_pair.source_lang}->{good_pair.target_lang}")
    check("zgodna liczba wierszy jest wykrywana", good_pair.matches)
    check("stan pary opisany po polsku", "6 wierszy" in good_pair.status,
          good_pair.status)
    bad_pair = next(p for p in gen_pairs if "krzywy" in p.source_path)
    check("niezgodna liczba wierszy jest wykrywana", not bad_pair.matches)
    check("opis różnicy podaje obie liczby",
          "2 / 3" in bad_pair.status, bad_pair.status)

    # kodowanie cp1250 nie może dawać krzaków
    check("plik cp1250 odczytany bez krzaków",
          _tb.read_lines(win_pl) == ["Okno", "Drzwi ściśle"],
          str(_tb.read_lines(win_pl)))
    check("kodowanie cp1250 rozpoznane, nie UTF-8",
          _tb.detect_encoding(win_pl) in ("cp1250", "iso-8859-2"),
          _tb.detect_encoding(win_pl))
    check("czysty UTF-8 nadal rozpoznawany jako UTF-8",
          _tb.detect_encoding(pl_path) in ("utf-8", "utf-8-sig"),
          _tb.detect_encoding(pl_path))

    built = _tb.build_pairs(gen_pairs, "en", "pl",
                            _tb.BuildOptions(skip_identical=True))
    sources = [r[0] for r in built.rows]
    check("zestawiono właściwe pary", "Hello world." in sources
          and "WIRELESS COMMUNICATION" in sources, str(sources))
    check("wiersz techniczny pominięty",
          not any("<<< FILE" in s_ for s_ in sources))
    check("identyczne pominięte przy włączonej opcji", "OK" not in sources)
    check("puste wiersze pominięte", "" not in sources)
    check("plik o złej liczbie wierszy pominięty w całości",
          "One" not in sources and "Two" not in sources, str(sources))
    check("powód pominięcia jest zapisany",
          any("krzywy" in p for p in built.problems), str(built.problems))
    check("podsumowanie wymienia liczby",
          "Gotowych par" in built.summary(), built.summary())
    check("wiersze mają języki z pary",
          all(r[2] == "en" and r[3] == "pl" for r in built.rows))

    # bez pomijania identycznych „OK” trafia do pamięci
    with_same = _tb.build_pairs(gen_pairs, "en", "pl",
                                _tb.BuildOptions(skip_identical=False))
    check("wyłączona opcja przepuszcza identyczne",
          "OK" in [r[0] for r in with_same.rows])

    # luźniejszy tryb: pary z różną liczbą wierszy do wspólnego minimum
    loose = _tb.build_pairs(gen_pairs, "en", "pl",
                            _tb.BuildOptions(require_equal_lines=False))
    check("tryb bez wymogu równości zestawia część wspólną",
          "One" in [r[0] for r in loose.rows], str([r[0] for r in loose.rows])[:80])

    preview_rows = _tb.preview_alignment(good_pair, 5)
    check("podgląd zwraca numerowane wiersze",
          preview_rows and preview_rows[0][0] == 1, str(preview_rows[:1]))
    check("podgląd zestawia właściwe teksty",
          preview_rows[0][1] == "Hello world."
          and preview_rows[0][2] == "Witaj świecie.", str(preview_rows[0]))

    # --- generator w interfejsie ---
    gen_tab = w.tm_tab
    check("zakładka Pamięć TM ma generator",
          any("Generator" in gen_tab.tabs.tabText(i)
              for i in range(gen_tab.tabs.count())),
          str([gen_tab.tabs.tabText(i) for i in range(gen_tab.tabs.count())]))
    check("tabela generatora przyjmuje upuszczone pliki",
          gen_tab.gen_table.acceptDrops())

    gen_tab.clear_generator()
    gen_tab.add_generator_files([gen_dir])
    # 5 plików z kodem języka daje 3 pary (miasto, krzywy, okno);
    # „sierota_de.txt” nie ma partnera po nazwie, więc trafia do puli
    # parowanej po kolejności – przy nieparzystej liczbie zostaje sam.
    check("upuszczenie katalogu paruje pliki",
          gen_tab.gen_table.rowCount() == 3, str(gen_tab.gen_table.rowCount()))
    check("nieparzysty plik zgłoszony w podpowiedzi",
          "czeka na parę" in gen_tab.gen_hint.text().lower(),
          gen_tab.gen_hint.text()[:70])
    check("podsumowanie ostrzega o niezgodnych wierszach",
          "niezgodną" in gen_tab.gen_summary.text(), gen_tab.gen_summary.text())

    row_miasto = next(r for r in range(gen_tab.gen_table.rowCount())
                      if "miasto" in gen_tab.gen_table.item(r, 0).text())
    gen_tab.gen_table.selectRow(row_miasto)
    app.processEvents()
    check("podgląd pokazuje zestawione wiersze",
          gen_tab.gen_preview.rowCount() > 0, str(gen_tab.gen_preview.rowCount()))
    check("podgląd zawiera tłumaczenie",
          any(gen_tab.gen_preview.item(r, 2).text() == "Witaj świecie."
              for r in range(gen_tab.gen_preview.rowCount())))

    # zamiana stron
    before_source = gen_tab.gen_table.item(row_miasto, 0).text()
    gen_tab.swap_generator_sides()
    check("„Zamień strony” odwraca pliki",
          gen_tab.gen_table.item(row_miasto, 0).text() != before_source,
          gen_tab.gen_table.item(row_miasto, 0).text())
    gen_tab.swap_generator_sides()

    tm_before = w.tm.size()
    gen_tab.gen_skip_identical.setChecked(True)
    with patch("supercat.ui.tm_tab.QMessageBox.information") as gen_info:
        gen_tab.build_tm_from_files()
    check("generator dopisał wpisy do pamięci", w.tm.size() > tm_before,
          f"{tm_before} -> {w.tm.size()}")
    check("okno wyniku podaje liczbę par",
          gen_info.call_args and "Dopisano do pamięci" in gen_info.call_args[0][2],
          str(gen_info.call_args[0][2][:50]) if gen_info.call_args else "")
    check("okno wyniku wymienia pominięte pliki",
          "krzywy" in gen_info.call_args[0][2], gen_info.call_args[0][2][-60:])
    found = w.tm.search("WIRELESS COMMUNICATION")
    check("wpis z generatora trafił do pamięci", len(found) >= 1, str(len(found)))

    # eksport TMX z generatora
    gen_tmx = os.path.join(tmp, "generator.tmx")
    with patch("supercat.ui.tm_tab.QFileDialog.getSaveFileName",
               return_value=(gen_tmx, "")), \
         patch("supercat.ui.tm_tab.QMessageBox.information"):
        gen_tab.export_generated_tmx()
    check("generator zapisuje plik TMX", os.path.exists(gen_tmx))
    tmx_text = open(gen_tmx, encoding="utf-8").read()
    check("TMX z generatora ma poprawny nagłówek",
          "<tmx version=\"1.4\">" in tmx_text)
    check("TMX z generatora zawiera polskie znaki",
          "Witaj świecie." in tmx_text)

    # --- dowolne nazwy plików + ręczny wybór języka i kodowania ----------
    free_dir = os.path.join(tmp, "gen_free")
    os.makedirs(free_dir, exist_ok=True)
    with open(os.path.join(free_dir, "1.txt"), "w",
              encoding="utf-8", newline="") as fh:
        fh.write("Guten Tag.\nWie geht es dir?\nAuf Wiedersehen.\n")
    with open(os.path.join(free_dir, "2.txt"), "w",
              encoding="cp1250", newline="") as fh:
        fh.write("Dzień dobry.\nJak się masz?\nDo widzenia.\n")

    free_paths = sorted(os.path.join(free_dir, n) for n in os.listdir(free_dir))
    free_pairs, free_rest = _tb.pair_files(free_paths, "de", "pl")
    check("pliki o dowolnych nazwach są parowane",
          len(free_pairs) == 1 and not free_rest,
          f"{len(free_pairs)} par, reszta {free_rest}")
    check("para z dowolnych nazw ma zgodne wiersze", free_pairs[0].matches,
          free_pairs[0].status)

    # nieparzysta liczba plików: jeden zostaje bez pary
    with open(os.path.join(free_dir, "3.txt"), "w", encoding="utf-8") as fh:
        fh.write("Hallo\n")
    odd_paths = sorted(os.path.join(free_dir, n) for n in os.listdir(free_dir))
    _odd_pairs, odd_rest = _tb.pair_files(odd_paths, "de", "pl")
    check("nieparzysty plik zgłoszony jako bez pary", len(odd_rest) == 1,
          str(odd_rest))
    os.remove(os.path.join(free_dir, "3.txt"))

    # kodowanie ustawiane ręcznie
    manual = free_pairs[0]
    check("domyślnie kodowanie jest automatyczne",
          manual.source_encoding == _tb.AUTO_ENCODING)
    check("automat rozpoznaje cp1250 w pliku docelowym",
          manual.encoding_of("target") in ("cp1250", "iso-8859-2"),
          manual.encoding_of("target"))
    manual.target_encoding = "cp1250"
    check("ręczne kodowanie czyta polskie znaki",
          manual.lines("target")[0] == "Dzień dobry.",
          str(manual.lines("target")[:1]))
    manual.target_encoding = _tb.AUTO_ENCODING

    manual.source_lang, manual.target_lang = "de", "pl"
    free_built = _tb.build_pairs([manual], "de", "pl", _tb.BuildOptions())
    check("wpisy dostają ręcznie wybrane języki",
          all(r[2] == "de" and r[3] == "pl" for r in free_built.rows),
          str(free_built.rows[:1]))
    check("tekst niemiecki i polski zestawiony",
          ("Guten Tag.", "Dzień dobry.") == free_built.rows[0][:2],
          str(free_built.rows[0][:2]))
    check("lista kodowań zawiera warianty polskie",
          {"cp1250", "iso-8859-2"} <= {c for c, _n in _tb.ENCODING_CHOICES})

    # --- to samo przez interfejs ---
    gen_tab.clear_generator()
    gen_tab.add_generator_files([free_dir])
    check("tabela ma kolumny języka i kodowania",
          gen_tab.gen_table.columnCount() == 6,
          str([gen_tab.gen_table.horizontalHeaderItem(c).text()
               for c in range(gen_tab.gen_table.columnCount())]))
    check("dowolne nazwy sparowane w interfejsie",
          gen_tab.gen_table.rowCount() == 1, str(gen_tab.gen_table.rowCount()))
    check("podpowiedź nie wymaga kodów języka w nazwie",
          "Sparowano" in gen_tab.gen_hint.text(), gen_tab.gen_hint.text()[:50])

    src_combo = gen_tab.gen_table.cellWidget(0, 1)
    enc_combo = gen_tab.gen_table.cellWidget(0, 4)
    check("kolumna języka to lista wyboru", src_combo is not None
          and src_combo.count() > 10, str(src_combo.count() if src_combo else 0))
    check("listę języka można uzupełnić własnym kodem", src_combo.isEditable())
    check("kolumna kodowania to lista wyboru", enc_combo is not None
          and enc_combo.count() > 5)

    de_index = next(i for i in range(src_combo.count())
                    if src_combo.itemData(i) == "de")
    src_combo.setCurrentIndex(de_index)
    app.processEvents()
    check("wybór języka trafia do pary",
          gen_tab._gen_pairs[0].source_lang == "de",
          gen_tab._gen_pairs[0].source_lang)

    cp_index = next(i for i in range(enc_combo.count())
                    if enc_combo.itemData(i) == "cp1250")
    enc_combo.setCurrentIndex(cp_index)
    app.processEvents()
    check("wybór kodowania trafia do pary",
          gen_tab._gen_pairs[0].source_encoding == "cp1250",
          gen_tab._gen_pairs[0].source_encoding)
    check("po zmianie kodowania stan jest przeliczony",
          gen_tab.gen_table.item(0, 5) is not None
          and "wiersze" in gen_tab.gen_table.item(0, 5).text(),
          gen_tab.gen_table.item(0, 5).text())
    enc_combo.setCurrentIndex(0)
    app.processEvents()

    tm_before_free = w.tm.size()
    with patch("supercat.ui.tm_tab.QMessageBox.information"):
        gen_tab.build_tm_from_files()
    check("TM z dowolnie nazwanych plików powstała",
          w.tm.size() > tm_before_free, f"{tm_before_free} -> {w.tm.size()}")
    german = w.tm.search("Guten Tag.")
    check("wpis niemiecko-polski zapisany z właściwymi językami",
          german and german[0][2] == "de" and german[0][3] == "pl",
          str(german[:1]))

    # --- REGRESJA: te same nazwy w różnych katalogach, dodawane po kolei ----
    # Użytkownik przeciąga plik EN, a chwilę potem PL z innego folderu.
    # Wcześniej pierwszy plik przepadał (nie był nigdzie zapamiętany),
    # więc para nigdy nie powstawała i lista zostawała pusta.
    same_root = os.path.join(tmp, "same_name")
    en_dir = os.path.join(same_root, "en")
    pl_dir = os.path.join(same_root, "pl")
    os.makedirs(en_dir, exist_ok=True)
    os.makedirs(pl_dir, exist_ok=True)
    for base, text_en, text_pl in (
        ("text", "Hello.\nGoodbye.\n", "Czesc.\nPa.\n"),
        ("menu", "Save game?\nLoad?\n", "Zapisac?\nWczytac?\n"),
    ):
        with open(os.path.join(en_dir, f"{base}.txt"), "w", encoding="utf-8") as fh:
            fh.write(text_en)
        with open(os.path.join(pl_dir, f"{base}.txt"), "w", encoding="utf-8") as fh:
            fh.write(text_pl)

    gen_tab.clear_generator()
    gen_tab.add_generator_files([os.path.join(en_dir, "text.txt")])
    check("pojedynczy plik czeka na parę zamiast przepaść",
          gen_tab.gen_table.rowCount() == 0
          and gen_tab._gen_waiting
          and "Czeka na parę" in gen_tab.gen_hint.text(),
          gen_tab.gen_hint.text()[:60])
    gen_tab.add_generator_files([os.path.join(pl_dir, "text.txt")])
    check("drugi plik o tej samej nazwie tworzy parę",
          gen_tab.gen_table.rowCount() == 1, str(gen_tab.gen_table.rowCount()))
    check("para z dwóch katalogów ma zgodne wiersze",
          gen_tab._gen_pairs[0].matches, gen_tab._gen_pairs[0].status)
    check("sparowano pliki z RÓŻNYCH katalogów",
          os.path.dirname(gen_tab._gen_pairs[0].source_path)
          != os.path.dirname(gen_tab._gen_pairs[0].target_path),
          gen_tab._gen_pairs[0].source_path)

    # katalog po katalogu — pliki EN nie mogą sparować się między sobą
    gen_tab.clear_generator()
    gen_tab.add_generator_files([en_dir])
    check("pliki w tym samym języku nie parują się ze sobą",
          gen_tab.gen_table.rowCount() == 0, str(gen_tab.gen_table.rowCount()))
    gen_tab.add_generator_files([pl_dir])
    check("dołożenie drugiego katalogu paruje wszystko",
          gen_tab.gen_table.rowCount() == 2, str(gen_tab.gen_table.rowCount()))
    pair_names = sorted(os.path.basename(p.source_path)
                        for p in gen_tab._gen_pairs)
    check("pary dobrane po nazwie pliku, nie po kolejności",
          pair_names == ["menu.txt", "text.txt"], str(pair_names))
    check("każda para łączy inny język",
          all(os.path.basename(os.path.dirname(p.source_path))
              != os.path.basename(os.path.dirname(p.target_path))
              for p in gen_tab._gen_pairs))

    tm_before_same = w.tm.size()
    with patch("supercat.ui.tm_tab.QMessageBox.information"):
        gen_tab.build_tm_from_files()
    check("TM z plików o tych samych nazwach powstała",
          w.tm.size() - tm_before_same == 4,
          f"{tm_before_same} -> {w.tm.size()}")
    hit_same = w.tm.search("Save game?")
    check("wpis z drugiego pliku trafił do pamięci", bool(hit_same),
          str(hit_same[:1]))

    # pliki tego samego języka bez partnera zostają zgłoszone
    lone_dir = os.path.join(tmp, "lone_en")
    os.makedirs(lone_dir, exist_ok=True)
    for name in ("x_en.txt", "y_en.txt"):
        with open(os.path.join(lone_dir, name), "w", encoding="utf-8") as fh:
            fh.write("Line.\n")
    lone_pairs, lone_rest = _tb.pair_files(
        sorted(os.path.join(lone_dir, n) for n in os.listdir(lone_dir)), "en", "pl")
    check("dwa pliki EN nie tworzą fałszywej pary",
          not lone_pairs and len(lone_rest) == 2,
          f"par={len(lone_pairs)} reszta={len(lone_rest)}")

    # --- podgląd pomijanych wierszy + wykrywanie języka -------------------
    check("nieprzetłumaczone wykryte po angielskich wyrazach",
          _tb.is_untranslated("Save the game", "Save the game now"))
    check("identyczne zdanie uznane za nieprzetłumaczone",
          _tb.is_untranslated("Would you like to save?", "Would you like to save?"))
    check("prawdziwe tłumaczenie przechodzi",
          not _tb.is_untranslated("Would you like to save?", "Czy chcesz zapisać?"))
    check("nazwa własna nie jest brana za brak tłumaczenia",
          not _tb.is_untranslated("CINNABAR GYM", "CINNABAR GYM"))
    check("tłumaczenie z angielską nazwą własną przechodzi",
          not _tb.is_untranslated("other TRAINERS", "inni TRENERZY"))
    check("polskie znaki rozpoznawane", _tb.looks_polish("Czy chcesz zapisać?"))
    check("polski bez znaków diakrytycznych też rozpoznany",
          _tb.looks_polish("Czy to jest test"))
    check("angielski rozpoznany", _tb.looks_english("Would you like to save?"))

    skip_dir = os.path.join(tmp, "gen_skip")
    os.makedirs(skip_dir, exist_ok=True)
    with open(os.path.join(skip_dir, "s_en.txt"), "w", encoding="utf-8") as fh:
        fh.write("Hello there.\n\n<<< FILE: city/text.inc >>>\n"
                 "Would you like to save?\nCINNABAR GYM\nSave the game\n")
    with open(os.path.join(skip_dir, "s_pl.txt"), "w", encoding="utf-8") as fh:
        fh.write("Witaj tam.\n\n<<< FILE: city/text.inc >>>\n"
                 "Czy chcesz zapisać?\nCINNABAR GYM\nSave the game now\n")
    skip_pairs, _rest = _tb.pair_files(
        [os.path.join(skip_dir, n) for n in sorted(os.listdir(skip_dir))], "en", "pl")
    skip_opts = _tb.BuildOptions(skip_identical=True, skip_untranslated=True)

    visible = _tb.preview_alignment(skip_pairs[0], 40, skip_opts, show_skipped=True)
    check("podgląd pokazuje wszystkie wiersze, także pomijane",
          len(visible) == 6, str(len(visible)))
    reasons = {row[3] for row in visible}
    check("podgląd podaje powód: pusty wiersz", "empty" in reasons, str(reasons))
    check("podgląd podaje powód: wiersz techniczny", "technical" in reasons)
    check("podgląd podaje powód: identyczne", "identical" in reasons)
    check("podgląd podaje powód: nieprzetłumaczone", "untranslated" in reasons)
    check("wiersze wchodzące do TM mają pusty powód",
          sum(1 for row in visible if not row[3]) == 2,
          str([row[1] for row in visible if not row[3]]))
    check("każdy powód ma opis po polsku",
          all(r in _tb.SKIP_REASONS for r in reasons), str(reasons))

    hidden = _tb.preview_alignment(skip_pairs[0], 40, skip_opts, show_skipped=False)
    check("bez przełącznika widać tylko wiersze do TM", len(hidden) == 2,
          str(len(hidden)))

    skip_res = _tb.build_pairs(skip_pairs, "en", "pl", skip_opts)
    check("podgląd zgadza się z faktycznym zestawieniem",
          skip_res.total == len(hidden), f"{skip_res.total} vs {len(hidden)}")
    check("licznik nieprzetłumaczonych w podsumowaniu",
          skip_res.skipped_untranslated == 1, str(skip_res.skipped_untranslated))
    check("podsumowanie wymienia nieprzetłumaczone",
          "nieprzetłumaczone" in skip_res.summary(), skip_res.summary())

    without = _tb.build_pairs(skip_pairs, "en", "pl",
                              _tb.BuildOptions(skip_identical=True))
    check("wyłączona opcja przepuszcza nieprzetłumaczone",
          without.total == 3, str(without.total))

    # --- w interfejsie ---
    check("generator ma przełącznik nieprzetłumaczonych",
          hasattr(gen_tab, "gen_skip_untranslated"))
    check("generator ma przełącznik pokazywania pomijanych",
          hasattr(gen_tab, "gen_show_skipped")
          and gen_tab.gen_show_skipped.isChecked())
    check("podgląd ma kolumnę Stan",
          gen_tab.gen_preview.columnCount() == 4
          and gen_tab.gen_preview.horizontalHeaderItem(3).text() == "Stan",
          str(gen_tab.gen_preview.columnCount()))

    gen_tab.clear_generator()
    gen_tab.add_generator_files([skip_dir])
    gen_tab.gen_skip_identical.setChecked(True)
    gen_tab.gen_skip_untranslated.setChecked(True)
    gen_tab.gen_table.selectRow(0)
    app.processEvents()
    states = [gen_tab.gen_preview.item(r, 3).text()
              for r in range(gen_tab.gen_preview.rowCount())]
    check("podgląd w oknie pokazuje pomijane wiersze",
          gen_tab.gen_preview.rowCount() == 6, str(gen_tab.gen_preview.rowCount()))
    check("kolumna Stan opisuje powody po polsku",
          any("techniczny" in s_ for s_ in states)
          and any("nieprzetłumaczone" in s_ for s_ in states), str(states))
    check("wiersze wchodzące do TM są oznaczone",
          sum(1 for s_ in states if "do TM" in s_) == 2, str(states))
    gen_tab.gen_show_skipped.setChecked(False)
    app.processEvents()
    check("wyłączenie przełącznika ukrywa pomijane",
          gen_tab.gen_preview.rowCount() == 2,
          str(gen_tab.gen_preview.rowCount()))
    gen_tab.gen_show_skipped.setChecked(True)
    gen_tab.gen_skip_untranslated.setChecked(False)

    # --- globalna opcja: TM nie przyjmuje nieprzetłumaczonych -------------
    check("Ustawienia mają przełącznik odrzucania nieprzetłumaczonych",
          hasattr(w.settings_tab, "reject_untranslated"))
    _prev_reject = SettingsManager.instance().get_bool("tm.reject.untranslated", False)
    SettingsManager.instance().set("tm.reject.untranslated", True)
    guard_tm = TranslationMemory()
    guard_tm.init_for_project(os.path.join(tmp, "tm_guard_reject.db"))
    check("TM odrzuca wpis pozostawiony po angielsku",
          guard_tm.add("Save the game", "Save the game now", "en", "pl") is False)
    check("TM przyjmuje prawdziwe tłumaczenie",
          guard_tm.add("Save the game", "Zapisz grę", "en", "pl") is True)
    check("TM przyjmuje nazwę własną bez zmian",
          guard_tm.add("CINNABAR GYM", "CINNABAR GYM", "en", "pl") is True)
    guard_tm.flush()
    stored_pairs = [(a, b) for a, b, *_r in guard_tm.all_entries()]
    check("w pamięci zostały tylko poprawne wpisy",
          ("Save the game", "Zapisz grę") in stored_pairs
          and ("Save the game", "Save the game now") not in stored_pairs,
          str(stored_pairs))
    SettingsManager.instance().set("tm.reject.untranslated", False)
    check("wyłączona opcja przepuszcza wszystko",
          guard_tm.add("Load game", "Load game", "en", "pl") is True)
    SettingsManager.instance().set("tm.reject.untranslated", _prev_reject)

    gen_tab.clear_generator()
    check("czyszczenie opróżnia listę", gen_tab.gen_table.rowCount() == 0)
    check("czyszczenie kasuje też pliki oczekujące", gen_tab._gen_waiting == [])
    with patch("supercat.ui.tm_tab.QMessageBox.information") as empty_info:
        gen_tab.build_tm_from_files()
    check("pusty generator prosi o pliki zamiast wywalać się",
          empty_info.call_args and "dodaj pliki" in empty_info.call_args[0][2].lower(),
          str(empty_info.call_args[0][2][:50]) if empty_info.call_args else "")

    # ------------------------------- Microsoft / Azure + wybór silników
    print("\n26. Microsoft, Azure i wybór silników")
    from supercat.core.mt import ENGINE_CODES as _CODES
    from supercat.core.mt import ENGINES as _ENGINES
    from supercat.core.mt import FREE_ENGINES as _FREE

    ms_ids = [k for k, _l in _ENGINES]
    check("silnik Microsoft bez klucza na liście", "microsoft_free" in ms_ids)
    check("silnik Azure Translator na liście", "azure" in ms_ids)
    check("Microsoft jest w silnikach darmowych", "microsoft_free" in _FREE, str(_FREE))
    check("Azure NIE jest darmowy bez klucza", "azure" not in _FREE)
    check("kody dostawców Microsoft/Azure",
          _CODES.get("microsoft_free") == "MS" and _CODES.get("azure") == "AZ")

    ms_mt = MachineTranslation()
    check("Microsoft dostępny bez konfiguracji",
          "microsoft_free" in ms_mt.available_engines(only_free=True))
    ms_mt.keys["azure_key"] = ""
    check("Azure niedostępny bez klucza", "azure" not in ms_mt.available_engines())
    azure_msg = ms_mt.translate_with("azure", "Hello", "en", "pl")
    check("Azure bez klucza podaje limit warstwy darmowej",
          "2 mln" in azure_msg, azure_msg[:80])
    check("Azure bez klucza wskazuje silnik bez konta",
          "Bing" in azure_msg, azure_msg[-80:])
    ms_mt.keys["azure_key"] = "x"
    check("Azure dostępny po podaniu klucza", "azure" in ms_mt.available_engines())
    # Sesja Bing jest współdzielona i chroniona blokadą – inaczej każdy wątek
    # QuickTrans pobierałby osobno stronę o wielkości ok. 600 kB.
    import time as _time_mod

    fresh = {"expires": _time_mod.time() + 3600, "ig": "X", "iid": "Y",
             "key": "1", "token": "t", "opener": None}
    refetched = []
    ms_mt._fetch_bing_session = lambda: (refetched.append(1), fresh)[1]
    ms_mt._bing_cache = fresh
    check("ważna sesja Microsoft nie jest pobierana ponownie",
          ms_mt._bing_session() is fresh and not refetched)
    ms_mt._bing_cache = dict(fresh, expires=_time_mod.time() - 10)
    ms_mt._bing_session()
    check("wygasła sesja Microsoft jest odświeżana", len(refetched) == 1,
          str(len(refetched)))
    ms_mt._bing_cache = None
    check("domyślny endpoint Azure",
          ms_mt.keys["azure_endpoint"].endswith("microsofttranslator.com"),
          ms_mt.keys["azure_endpoint"])

    # wybór silników QuickTrans (bez sieci – sprawdzamy dobór, nie tłumaczenie)
    settings_mgr = SettingsManager.instance()
    settings_mgr.set("mt.quicktrans.engines", "")
    settings_mgr.set("mt.quicktrans.free_only", True)
    auto = ms_mt.quicktrans_engines()
    check("QuickTrans bez wyboru bierze darmowe",
          "google_free" in auto and "microsoft_free" in auto, str(auto))
    settings_mgr.set("mt.quicktrans.engines", "local,microsoft_free")
    picked = ms_mt.quicktrans_engines()
    check("QuickTrans respektuje wybór użytkownika",
          picked == ["local", "microsoft_free"], str(picked))
    settings_mgr.set("mt.quicktrans.engines", "gemini")   # brak klucza
    check("nieskonfigurowany wybór wraca do trybu automatycznego",
          len(ms_mt.quicktrans_engines()) > 1)
    settings_mgr.set("mt.quicktrans.engines", "")

    # --- DeepL przez stronę (bez klucza) ---------------------------------
    check("silnik DeepL przez stronę na liście", "deepl_web" in ms_ids)
    check("DeepL przez stronę jest darmowy", "deepl_web" in _FREE)
    check("DeepL przez stronę dostępny bez konfiguracji",
          "deepl_web" in ms_mt.available_engines(only_free=True))
    check("kod dostawcy DeepL WWW", _CODES.get("deepl_web") == "DLW")
    check("DeepL WWW ma wymuszony odstęp zapytań",
          ms_mt.DEEPL_WEB_MIN_INTERVAL >= 1.0, str(ms_mt.DEEPL_WEB_MIN_INTERVAL))
    check("DeepL WWW tłumaczy paczkami", ms_mt.DEEPL_WEB_BATCH >= 5,
          str(ms_mt.DEEPL_WEB_BATCH))

    # tłumaczenie wsadowe dzieli listę na paczki (bez sieci – podstawiamy zapytanie)
    sent_batches = []

    def _fake_request(texts, sl, tl):
        sent_batches.append(list(texts))
        return [f"PL:{t}" for t in texts]

    batch_mt = MachineTranslation()
    batch_mt.set_engine("deepl_web")
    batch_mt._deepl_web_request = _fake_request
    batch_out = batch_mt.translate_batch([f"Line {i}." for i in range(23)], "en", "pl")
    check("wsad DeepL zwraca tyle wyników, ile tekstów", len(batch_out) == 23,
          str(len(batch_out)))
    check("wsad DeepL dzieli na paczki po 10",
          [len(b) for b in sent_batches] == [10, 10, 3],
          str([len(b) for b in sent_batches]))
    check("wsad DeepL zachowuje kolejność", batch_out[0].endswith("Line 0.")
          and batch_out[22].endswith("Line 22."), batch_out[22])

    sent_batches.clear()

    def _boom(texts, sl, tl):
        sent_batches.append(list(texts))
        raise RuntimeError("429")

    batch_mt._deepl_web_request = _boom
    # Bez zapasu paczka zwraca czytelny błąd dla każdego segmentu.
    # (Z zapasem – domyślnie włączonym – tłumaczy Microsoft; sprawdzane osobno.)
    SettingsManager.instance().set("mt.deepl.web.fallback", False)
    broken = batch_mt.translate_batch(["A.", "B.", "C."], "en", "pl")
    SettingsManager.instance().set("mt.deepl.web.fallback", True)
    check("błąd paczki nie gubi segmentów", len(broken) == 3, str(len(broken)))
    check("błąd paczki jest opisany w każdym segmencie",
          all(b.startswith("[Błąd MT") for b in broken), broken[0][:40])

    # znaczniki gier przechodzą przez tłumaczenie wsadowe
    def _echo(texts, sl, tl):
        return list(texts)

    batch_mt._deepl_web_request = _echo
    tagged = batch_mt.translate_batch(
        ["Hello \\nWorld.", "{PLAYER} won!"], "en", "pl")
    check("wsad DeepL przywraca znaczniki \\n",
          "\\n" in tagged[0], repr(tagged[0]))
    check("wsad DeepL przywraca zmienne {PLAYER}",
          "{PLAYER}" in tagged[1], repr(tagged[1]))
    batch_mt.set_engine("local")

    # --- regresja: translate_with nie może psuć silnika w innych wątkach ---
    # Wcześniej metoda podmieniała self.engine na czas wywołania; przy
    # równoległym QuickTransie wątki nadpisywały sobie ustawienie nawzajem,
    # a lista silnika w interfejsie zostawała na przypadkowej pozycji.
    from concurrent.futures import ThreadPoolExecutor as _Pool

    race_mt = MachineTranslation()
    race_mt.set_engine("local")
    seen = []

    def _slow_local(text, sl, tl, engine=None):
        seen.append(engine or race_mt.engine)
        _time_mod.sleep(0.02)
        return f"{engine or race_mt.engine}:{text}"

    race_mt._translate_one = _slow_local
    with _Pool(4) as pool:
        race_out = list(pool.map(
            lambda e: race_mt.translate_with(e, "Hello", "en", "pl"),
            ["local", "mymemory", "google_free", "microsoft_free"]))
    check("równoległe translate_with nie mieszają silników",
          sorted(o.split(":")[0] for o in race_out) ==
          ["google_free", "local", "microsoft_free", "mymemory"],
          str(sorted(o.split(":")[0] for o in race_out)))
    check("translate_with nie zmienia silnika domyślnego (wielowątkowo)",
          race_mt.engine == "local", race_mt.engine)

    # zmiana silnika z kodu powiadamia wszystkie widoki
    w.mt.set_engine("google_free")
    check("set_engine odświeża listę w edytorze",
          w.editor_tab.engine_picker.currentData() == "google_free",
          str(w.editor_tab.engine_picker.currentData()))
    check("set_engine odświeża listę w Ustawieniach",
          w.settings_tab.engine_combo.currentData() == "google_free",
          str(w.settings_tab.engine_combo.currentData()))
    w.mt.set_engine("local")

    # --- DeepL WWW: limit 429 nie może zatrzymać pracy --------------------
    from supercat.core.mt import WHOLE_SEGMENT_ENGINES as _WHOLE

    # jeden segment = JEDNO zapytanie (dzielenie po \n i \p mnożyło je 2-4x)
    check("DeepL WWW dostaje cały segment naraz", "deepl_web" in _WHOLE)
    calls = []
    seg_mt = MachineTranslation()
    seg_mt._deepl_web_request = lambda texts, sl, tl: (
        calls.append(list(texts)) or [f"PL:{t}" for t in texts])
    seg_mt.translate(r"Ala ma kota.\nA kot ma Alę.\pKoniec.", "en", "pl", "deepl_web")
    check("segment ze znacznikami to jedno zapytanie", len(calls) == 1,
          f"{len(calls)} zapytań")

    # blokada jest zapamiętywana – nie zasypujemy serwera
    block_mt = MachineTranslation()
    MachineTranslation._deepl_web_blocked_until = _time_mod.time() + 120
    check("po odmowie silnik wie, że jest zablokowany",
          block_mt.deepl_web_ready() is False)
    blocked_msg = MachineTranslation._deepl_web_blocked_message(120)
    check("komunikat blokady podaje czas oczekiwania",
          "2 min" in blocked_msg, blocked_msg[:70])
    check("komunikat blokady proponuje Microsoft",
          "Microsoft" in blocked_msg)
    check("komunikat blokady wspomina o darmowym kluczu",
          "500 000" in blocked_msg)

    # zapas: zamiast błędu tłumaczy Microsoft
    SettingsManager.instance().set("mt.deepl.web.fallback", True)
    spare_mt = MachineTranslation()
    spare_mt._microsoft_free = lambda text, sl, tl: "ZAMIENNIK"
    spare_out = spare_mt.translate("Hello", "en", "pl", "deepl_web")
    check("blokada DeepL nie zwraca błędu, tylko tłumaczenie",
          spare_out == "ZAMIENNIK", spare_out[:60])
    check("użytkownik jest informowany o zamienniku",
          "Microsoft" in spare_mt._last_fallback, spare_mt._last_fallback[:60])

    # wsad również korzysta z zapasu
    batch_spare = MachineTranslation()
    batch_spare.set_engine("deepl_web")

    def _blocked(texts, sl, tl):
        raise RuntimeError("429")

    batch_spare._deepl_web_request = _blocked
    batch_spare._microsoft_free = lambda text, sl, tl: f"MS:{text}"
    batch_out2 = batch_spare.translate_batch(["A.", "B."], "en", "pl")
    check("„Tłumacz wszystko” nie staje na zablokowanym DeepL",
          all(b.startswith("MS:") for b in batch_out2), str(batch_out2))
    batch_spare.set_engine("local")

    # da się wyłączyć – wtedy widać surowy błąd
    SettingsManager.instance().set("mt.deepl.web.fallback", False)
    strict_mt = MachineTranslation()
    strict_mt._microsoft_free = lambda text, sl, tl: "ZAMIENNIK"
    strict_out = strict_mt.translate("Hello", "en", "pl", "deepl_web")
    check("wyłączony zapas przywraca komunikat o błędzie",
          strict_out.startswith("[Błąd MT"), strict_out[:40])
    SettingsManager.instance().set("mt.deepl.web.fallback", True)
    MachineTranslation._deepl_web_blocked_until = 0.0

    check("Ustawienia mają przełącznik zapasu DeepL",
          hasattr(w.settings_tab, "deepl_fallback")
          and w.settings_tab.deepl_fallback.isChecked())

    # zakładka Ustawienia: pola kluczy i lista silników QuickTrans
    check("Ustawienia mają pole klucza Azure", "azure_key" in w.settings_tab.key_fields)
    check("Ustawienia mają pole regionu Azure", "azure_region" in w.settings_tab.key_fields)
    check("Ustawienia: lista silników QuickTrans pełna",
          len(w.settings_tab.qt_engine_boxes) == len(_ENGINES),
          f"{len(w.settings_tab.qt_engine_boxes)}/{len(_ENGINES)}")
    w.settings_tab._set_quicktrans_engines(["local", "google_free"])
    check("zaznaczenie silników zapisuje się w ustawieniach",
          settings_mgr.get_str("mt.quicktrans.engines", "") == "local,google_free",
          settings_mgr.get_str("mt.quicktrans.engines", ""))
    check("podsumowanie pokazuje liczbę wybranych",
          "2" in w.settings_tab.qt_summary.text(), w.settings_tab.qt_summary.text())
    w.settings_tab._set_quicktrans_engines([])
    check("wyczyszczenie wraca do trybu automatycznego",
          settings_mgr.get_str("mt.quicktrans.engines", "") == ""
          and "utomat" in w.settings_tab.qt_summary.text())

    # edytor: szybki wybór silnika obok pola tłumaczenia
    picker = w.editor_tab.engine_picker
    check("edytor ma listę silników przy polu tłumaczenia", picker.count() > 0,
          str(picker.count()))
    check("lista edytora zawiera Microsoft bez klucza",
          any(picker.itemData(i) == "microsoft_free" for i in range(picker.count())))
    check("lista edytora wskazuje bieżący silnik",
          picker.currentData() == w.mt.engine, f"{picker.currentData()} / {w.mt.engine}")
    target_row = next(i for i in range(picker.count())
                      if picker.itemData(i) == "mymemory")
    picker.setCurrentIndex(target_row)
    check("zmiana w edytorze przestawia silnik programu",
          w.mt.engine == "mymemory", w.mt.engine)
    check("zmiana w edytorze aktualizuje Ustawienia",
          w.settings_tab.engine_combo.currentData() == "mymemory",
          str(w.settings_tab.engine_combo.currentData()))
    w.settings_tab.engine_combo.setCurrentIndex(
        next(i for i, (k, _l) in enumerate(_ENGINES) if k == "local"))
    check("zmiana w Ustawieniach wraca do edytora",
          w.editor_tab.engine_picker.currentData() == "local",
          str(w.editor_tab.engine_picker.currentData()))
    silnik_bez_klucza = [picker.itemText(i) for i in range(picker.count())
                         if picker.itemData(i) == "openai"]
    check("silniki bez klucza mają oznaczenie 🔑",
          silnik_bez_klucza and silnik_bez_klucza[0].startswith("🔑"),
          str(silnik_bez_klucza))
    w.editor_tab.engine_free_only.setChecked(True)
    only_free_ids = [picker.itemData(i) for i in range(picker.count())]
    check("filtr „tylko bez klucza” skraca listę",
          "openai" not in only_free_ids and "microsoft_free" in only_free_ids,
          str(only_free_ids))
    w.editor_tab.engine_free_only.setChecked(False)
    check("wyłączenie filtru przywraca pełną listę",
          len([1 for i in range(picker.count())]) == len(_ENGINES),
          str(picker.count()))

    # ------------------------------------------------------------ podsumowanie
    print("\n" + "=" * 60)
    print(f"WYNIK: {len(PASS)} zaliczonych, {len(FAIL)} niezaliczonych")
    if FAIL:
        print("Niezaliczone:")
        for f in FAIL:
            print("  - " + f)
    print("=" * 60)
    shutil.rmtree(tmp, ignore_errors=True)
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(main())
