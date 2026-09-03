"""Parsowanie i eksport plików (odpowiednik services/FileParserService.java).

Obsługiwane formaty wejściowe: TXT, DOCX, XLSX, XLIFF/XLF, PO/POT, SRT, HTML/HTM, MD, CSV.
Eksport: odtworzenie oryginału z podmienionym tekstem, XLIFF, PO, SRT, HTML, TXT, DOCX.
"""
from __future__ import annotations

import csv
import html as html_module
import os
import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from typing import List, Optional

SUPPORTED_EXTENSIONS = (
    ".txt", ".md", ".csv", ".docx", ".xlsx", ".xliff", ".xlf",
    ".po", ".pot", ".srt", ".html", ".htm", ".xml",
)


@dataclass
class Segment:
    seg_id: str
    source: str
    target: str = ""
    notes: str = ""
    file_type: str = ""
    file_name: str = ""
    ignored: bool = False
    status: str = "new"  # new | draft | translated | approved
    extra: dict = field(default_factory=dict)

    @property
    def is_translated(self) -> bool:
        return bool(self.target and self.target.strip())


# ---------------------------------------------------------------- parsowanie
def parse_file(path: str, seg_settings=None) -> List[Segment]:
    """Parsuje plik na segmenty. seg_settings – ustawienia segmentacji projektu."""
    lower = path.lower()
    if lower.endswith((".xliff", ".xlf")):
        segments = parse_xliff(path)
    elif lower.endswith((".po", ".pot")):
        segments = parse_po(path)
    elif lower.endswith(".srt"):
        segments = parse_srt(path)
    elif lower.endswith(".docx"):
        segments = parse_docx(path, seg_settings)
    elif lower.endswith(".xlsx"):
        segments = parse_xlsx(path)
    elif lower.endswith((".html", ".htm")):
        segments = parse_html(path, seg_settings)
    elif lower.endswith(".csv"):
        segments = parse_csv(path)
    elif lower.endswith((".txt", ".md", ".xml")):
        segments = parse_txt(path, seg_settings)
    else:
        raise ValueError(f"Nieobsługiwany typ pliku: {os.path.basename(path)}")

    name = os.path.basename(path)
    for seg in segments:
        seg.file_name = name
    return segments


def _read_text(path: str) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1250", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _segment_chunks(chunks: List[str], seg_settings) -> List[str]:
    if seg_settings is None:
        return [c for c in chunks if c.strip()]
    from .segmentation import segment_text

    result: List[str] = []
    for chunk in chunks:
        if chunk.strip():
            result.extend(segment_text(chunk, seg_settings))
    return result


def parse_txt(path: str, seg_settings=None) -> List[Segment]:
    content = _read_text(path)
    paragraphs = re.split(r"\r?\n", content)
    parts = _segment_chunks(paragraphs, seg_settings)
    base = os.path.basename(path)
    return [Segment(f"{base}_{i + 1}", text, "", file_type="txt") for i, text in enumerate(parts)]


def parse_csv(path: str) -> List[Segment]:
    segments: List[Segment] = []
    base = os.path.basename(path)
    with open(path, "r", encoding="utf-8", newline="", errors="replace") as fh:
        for row_idx, row in enumerate(csv.reader(fh), start=1):
            for col_idx, cell in enumerate(row):
                if cell.strip():
                    # zachowujemy komórkę dokładnie tak, jak w pliku (także wcięcia)
                    seg = Segment(f"{base}_{row_idx}_{col_idx}", cell, "", file_type="csv")
                    seg.extra = {"row": row_idx, "col": col_idx}
                    segments.append(seg)
    return segments


def parse_xliff(path: str) -> List[Segment]:
    segments: List[Segment] = []
    tree = ET.parse(path)
    root = tree.getroot()
    units = root.findall(".//{*}trans-unit") or root.findall(".//trans-unit")
    if not units:
        units = root.findall(".//{*}unit")
    for i, unit in enumerate(units, start=1):
        uid = unit.get("id") or f"unit_{i}"
        src_el = unit.find("{*}source") or unit.find("source") or unit.find(".//{*}source")
        tgt_el = unit.find("{*}target") or unit.find("target") or unit.find(".//{*}target")
        source = "".join(src_el.itertext()).strip() if src_el is not None else ""
        target = "".join(tgt_el.itertext()).strip() if tgt_el is not None else ""
        if source:
            segments.append(Segment(uid, source, target, file_type="xliff"))
    return segments


def parse_po(path: str) -> List[Segment]:
    segments: List[Segment] = []
    content = _read_text(path)
    blocks = re.split(r"\n\s*\n", content)
    counter = 1
    for block in blocks:
        msgid_parts = re.findall(r'^msgid\s+"(.*)"$', block, re.MULTILINE)
        msgstr_parts = re.findall(r'^msgstr\s+"(.*)"$', block, re.MULTILINE)
        cont = re.findall(r'^"(.*)"$', block, re.MULTILINE)
        notes = " ".join(re.findall(r"^#\.\s*(.*)$", block, re.MULTILINE))
        if not msgid_parts:
            continue
        msgid = msgid_parts[0]
        msgstr = msgstr_parts[0] if msgstr_parts else ""
        if not msgid.strip() and cont:
            msgid = "".join(cont)
        if not msgid.strip():
            continue  # nagłówek PO
        seg = Segment(f"po_{counter}", _unescape_po(msgid), _unescape_po(msgstr), notes, file_type="po")
        segments.append(seg)
        counter += 1
    return segments


def parse_srt(path: str) -> List[Segment]:
    segments: List[Segment] = []
    content = _read_text(path)
    blocks = re.split(r"\r?\n\r?\n", content)
    for block in blocks:
        lines = [l for l in re.split(r"\r?\n", block) if l.strip()]
        if len(lines) >= 3 and "-->" in lines[1]:
            text = " ".join(lines[2:]).strip()
            seg = Segment(f"srt_{lines[0].strip()}", text, "", file_type="srt")
            seg.extra = {"index": lines[0].strip(), "time": lines[1].strip()}
            segments.append(seg)
    return segments


def parse_docx(path: str, seg_settings=None) -> List[Segment]:
    from docx import Document

    doc = Document(path)
    chunks: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    parts = _segment_chunks(chunks, seg_settings)
    base = os.path.basename(path)
    return [Segment(f"{base}_{i + 1}", text, "", file_type="docx") for i, text in enumerate(parts)]


def parse_xlsx(path: str) -> List[Segment]:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    segments: List[Segment] = []
    base = os.path.basename(path)
    counter = 1
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.strip():
                    seg = Segment(f"{base}_{counter}", value, "", file_type="xlsx")
                    seg.extra = {"sheet": ws.title, "coord": cell.coordinate}
                    segments.append(seg)
                    counter += 1
    return segments


def parse_html(path: str, seg_settings=None) -> List[Segment]:
    content = _read_text(path)
    # usuń skrypty i style
    content = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", content, flags=re.S | re.I)
    texts = re.split(r"<[^>]+>", content)
    texts = [html_module.unescape(t).strip() for t in texts]
    parts = _segment_chunks([t for t in texts if t], seg_settings)
    base = os.path.basename(path)
    return [Segment(f"{base}_{i + 1}", t, "", file_type="html") for i, t in enumerate(parts)]


def _unescape_po(text: str) -> str:
    return text.replace('\\"', '"').replace("\\n", "\n").replace("\\t", "\t").replace("\\\\", "\\")


def _escape_po(text: str) -> str:
    return text.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n").replace("\t", "\\t")


# ------------------------------------------------------------------- eksport
def export_by_replacement(source_file: str, target_file: str, segments: List[Segment]) -> None:
    """Odtwarza plik źródłowy podmieniając teksty na tłumaczenia (txt/html/md/xml).

    Spacje na brzegach segmentu (wcięcie wiersza w plikach gier) są przenoszone
    na tłumaczenie, nawet jeśli tłumacz wpisał tekst bez nich.
    """
    from .textutil import copy_edge_whitespace

    content = _read_text(source_file)
    result: List[str] = []
    last_index = 0
    for seg in segments:
        if not seg.source:
            continue
        idx = content.find(seg.source, last_index)
        if idx < 0:
            continue
        result.append(content[last_index:idx])
        result.append(copy_edge_whitespace(seg.source, seg.target)
                      if seg.is_translated else seg.source)
        last_index = idx + len(seg.source)
    result.append(content[last_index:])
    os.makedirs(os.path.dirname(os.path.abspath(target_file)), exist_ok=True)
    with open(target_file, "w", encoding="utf-8") as fh:
        fh.write("".join(result))


def export_docx(source_file: str, target_file: str, segments: List[Segment]) -> None:
    """Zapisuje DOCX z podmienionym tekstem akapitów (zachowuje układ dokumentu)."""
    from docx import Document

    doc = Document(source_file)
    mapping = {seg.source.strip(): seg.target for seg in segments if seg.is_translated}

    def translate_paragraph(paragraph) -> None:
        text = paragraph.text.strip()
        if not text:
            return
        new_text = mapping.get(text)
        if new_text is None:
            # segmentacja zdaniowa – podmień fragmenty
            new_text = text
            for src, tgt in mapping.items():
                if src and src in new_text:
                    new_text = new_text.replace(src, tgt)
            if new_text == text:
                return
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = new_text

    for paragraph in doc.paragraphs:
        translate_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    translate_paragraph(paragraph)

    os.makedirs(os.path.dirname(os.path.abspath(target_file)), exist_ok=True)
    doc.save(target_file)


def export_xlsx(source_file: str, target_file: str, segments: List[Segment]) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(source_file)
    for seg in segments:
        if not seg.is_translated:
            continue
        sheet = seg.extra.get("sheet")
        coord = seg.extra.get("coord")
        if sheet and coord and sheet in wb.sheetnames:
            wb[sheet][coord] = seg.target
    os.makedirs(os.path.dirname(os.path.abspath(target_file)), exist_ok=True)
    wb.save(target_file)


def export_srt(segments: List[Segment], target_file: str) -> None:
    blocks = []
    for i, seg in enumerate(segments, start=1):
        index = seg.extra.get("index", str(i))
        timing = seg.extra.get("time", "00:00:00,000 --> 00:00:02,000")
        text = seg.target if seg.is_translated else seg.source
        blocks.append(f"{index}\n{timing}\n{text}\n")
    _write(target_file, "\n".join(blocks))


def export_po(segments: List[Segment], target_file: str) -> None:
    lines = ['msgid ""', 'msgstr ""', '"Content-Type: text/plain; charset=UTF-8\\n"', ""]
    for seg in segments:
        if seg.notes:
            lines.append(f"#. {seg.notes}")
        lines.append(f'msgid "{_escape_po(seg.source)}"')
        lines.append(f'msgstr "{_escape_po(seg.target)}"')
        lines.append("")
    _write(target_file, "\n".join(lines))


def export_xliff(segments: List[Segment], target_file: str, source_lang: str = "en", target_lang: str = "pl") -> None:
    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<xliff version="1.2" xmlns="urn:oasis:names:tc:xliff:document:1.2">',
        f'  <file original="SuperCAT" source-language="{source_lang}" target-language="{target_lang}" datatype="plaintext">',
        "    <body>",
    ]
    for seg in segments:
        state = "translated" if seg.is_translated else "new"
        lines.append(f'      <trans-unit id="{_xml_escape(seg.seg_id)}">')
        lines.append(f"        <source>{_xml_escape(seg.source)}</source>")
        lines.append(f'        <target state="{state}">{_xml_escape(seg.target)}</target>')
        if seg.notes:
            lines.append(f"        <note>{_xml_escape(seg.notes)}</note>")
        lines.append("      </trans-unit>")
    lines += ["    </body>", "  </file>", "</xliff>"]
    _write(target_file, "\n".join(lines))


def export_html_bilingual(segments: List[Segment], target_file: str) -> None:
    rows = []
    for i, seg in enumerate(segments, start=1):
        color = "#e8f5e9" if seg.is_translated else "#fff8e1"
        rows.append(
            f'<tr style="background:{color}"><td>{i}</td><td>{_xml_escape(seg.source)}</td>'
            f"<td>{_xml_escape(seg.target)}</td></tr>"
        )
    html_doc = (
        "<!DOCTYPE html><html lang='pl'><head><meta charset='utf-8'>"
        "<title>SuperCAT – eksport dwujęzyczny</title>"
        "<style>body{font-family:Segoe UI,Arial,sans-serif;padding:20px}"
        "table{border-collapse:collapse;width:100%}td,th{border:1px solid #ccc;padding:8px;vertical-align:top}"
        "th{background:#1976d2;color:#fff}</style></head><body>"
        "<h2>SuperCAT – eksport dwujęzyczny</h2><table>"
        "<tr><th style='width:50px'>#</th><th>Źródło</th><th>Tłumaczenie</th></tr>"
        + "".join(rows)
        + "</table></body></html>"
    )
    _write(target_file, html_doc)


def export_txt(segments: List[Segment], target_file: str) -> None:
    _write(target_file, "\n".join(seg.target if seg.is_translated else seg.source for seg in segments))


def _write(path: str, content: str) -> None:
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(content)


def _xml_escape(text: Optional[str]) -> str:
    if not text:
        return ""
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    )
